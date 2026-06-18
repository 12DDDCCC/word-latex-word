"""Build a Word References section with bookmarks from parsed .bbl items."""

from __future__ import annotations

import os
import sys
from dataclasses import dataclass
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
if _THIS_DIR not in sys.path:
    sys.path.insert(0, _THIS_DIR)

from cross_ref_builder import (  # noqa: E402
    _add_bookmark_to_paragraph,
    _bib_key_to_bookmark,
    _get_max_bookmark_id,
)

_PLACEHOLDER = "[REFERENCES_PLACEHOLDER]"
_REF_TITLE = "References"
_HANG_LEFT = 720
_HANG_HANGING = 360


@dataclass
class RefSectionResult:
    bookmarks_added: int = 0
    paras_added: int = 0
    doi_links_added: int = 0


def build_references_section(
    doc: Document,
    bbl_items: list,
    cite_map: dict | None = None,
    cite_style: str | None = None,
) -> RefSectionResult:
    """Insert a generated References section at the placeholder paragraph."""
    from docx.shared import Pt

    result = RefSectionResult()
    if not bbl_items:
        print("  [ref_section_builder] no .bbl items; skipping generated References")
        return result

    placeholder_para = _find_placeholder(doc)
    if placeholder_para is None:
        print("  [ref_section_builder] placeholder not found; appending References at document end")

    next_bookmark_id = _get_max_bookmark_id(doc) + 1

    if placeholder_para is not None:
        title_para = _insert_paragraph_before(placeholder_para, _REF_TITLE)
    else:
        title_para = doc.add_paragraph(_REF_TITLE)
    _set_heading_style(title_para)
    result.paras_added += 1

    cited_keys = set(cite_map or {})
    numbered = _is_numbered_style(cite_style, cite_map)
    visible_idx = 0
    for item in bbl_items:
        if cited_keys and item.key not in cited_keys:
            continue
        visible_idx += 1

        para = _insert_paragraph_before(placeholder_para, "") if placeholder_para is not None else doc.add_paragraph("")
        _set_hanging_indent(para)

        if numbered:
            label = _reference_number_label(item.key, cite_map, visible_idx)
            run = para.add_run(f"[{label}] ")
            run.font.size = Pt(10)

        if item.plain_text:
            run = para.add_run(item.plain_text)
            run.font.size = Pt(10)

        bookmark_name = _bib_key_to_bookmark(item.key)
        _add_bookmark_to_paragraph(para, bookmark_name, next_bookmark_id)
        next_bookmark_id += 1
        result.bookmarks_added += 1

        url = item.doi_url or item.external_url
        if url:
            if item.plain_text:
                para.add_run(" ")
            _add_hyperlink_to_paragraph(para, url, doc)
            result.doi_links_added += 1

        result.paras_added += 1

    if placeholder_para is not None:
        placeholder_para._element.getparent().remove(placeholder_para._element)

    print(
        "  [ref_section_builder] generated References: "
        f"bookmarks={result.bookmarks_added}, paragraphs={result.paras_added}, "
        f"links={result.doi_links_added}"
    )
    return result


def _is_numbered_style(cite_style: str | None, cite_map: dict | None) -> bool:
    style = str(cite_style or "").lower()
    if style in {"numbered", "ieee", "numeric", "vancouver"}:
        return True
    values = [str(value) for value in (cite_map or {}).values()]
    return bool(values) and all(value.isdigit() for value in values)


def _reference_number_label(key: str, cite_map: dict | None, fallback: int) -> str:
    label = str((cite_map or {}).get(key) or "").strip()
    return label if label.isdigit() else str(fallback)


def _find_placeholder(doc: Document) -> Paragraph | None:
    for para in doc.paragraphs:
        if _PLACEHOLDER in para.text:
            return para
    return None


def _insert_paragraph_before(target_para: Paragraph, text: str) -> Paragraph:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    if text:
        new_r = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.set(qn("xml:space"), "preserve")
        new_t.text = text
        new_r.append(new_t)
        new_p.append(new_r)
    target_para._element.addprevious(new_p)
    return Paragraph(new_p, target_para._parent)


def _set_heading_style(para: Paragraph) -> None:
    from docx.shared import Pt

    try:
        para.style = para.part.document.styles["Heading 1"]
    except KeyError:
        pass
    for run in para.runs:
        run.bold = True
        run.font.size = Pt(12)


def _set_hanging_indent(para: Paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        para._element.insert(0, ppr)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:left"), str(_HANG_LEFT))
    ind.set(qn("w:hanging"), str(_HANG_HANGING))


def _add_hyperlink_to_paragraph(para: Paragraph, url: str, doc: Document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    rid = doc.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)

    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = url
    run.append(text)
    hyperlink.append(run)
    para._element.append(hyperlink)


if __name__ == "__main__":
    print("Use build_references_section() from tex-to-word postprocessing.")
