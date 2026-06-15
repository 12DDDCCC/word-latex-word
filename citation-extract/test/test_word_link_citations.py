import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SKILL_DIR = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(SKILL_DIR / "citation-extract"))
sys.path.insert(0, str(SKILL_DIR / "text-extract"))
sys.path.insert(0, str(SKILL_DIR / "omml-to-latex"))

from text_extract import extract_docx_text
from word_link_citations import extract_word_link_citations


def _field_run(kind, text=None):
    run = OxmlElement("w:r")
    if kind in ("begin", "separate", "end"):
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), kind)
        run.append(fld)
    elif kind == "instr":
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = text
        run.append(instr)
    elif kind == "text":
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
    return run


def _add_citation_field(paragraph, target, display):
    for kind, text in [
        ("begin", None),
        ("instr", f' HYPERLINK \\l "{target}" '),
        ("separate", None),
        ("text", display),
        ("end", None),
    ]:
        paragraph._element.append(_field_run(kind, text))


class WordLinkCitationTest(unittest.TestCase):
    def _write_bib(self, tmp):
        bib = Path(tmp) / "refs.bib"
        bib.write_text(
            """
@article{1,
  author = {Jiang, F. and Wang, H.},
  year = {2021},
  title = {Regional fluxes}
}
@article{abc-key,
  author = {Smith, A. and Doe, B.},
  year = {2024},
  title = {Linked citation}
}
""",
            encoding="utf-8",
        )
        return bib

    def test_extracts_word_field_citation_by_bookmark_key(self):
        with tempfile.TemporaryDirectory() as tmp:
            bib = self._write_bib(tmp)
            docx = Path(tmp) / "paper.docx"
            doc = Document()
            para = doc.add_paragraph("Before ")
            _add_citation_field(para, "_Bib_1", "(Jiang et al., 2021)")
            para.add_run(" after.")
            doc.save(docx)

            result = extract_word_link_citations(docx, bib)

            self.assertEqual(result["total_link_citations"], 1)
            self.assertEqual(result["citations"][0]["key"], "1")

    def test_text_extract_uses_linked_citation_and_groups_parenthetical_refs(self):
        with tempfile.TemporaryDirectory() as tmp:
            bib = self._write_bib(tmp)
            docx = Path(tmp) / "paper.docx"
            doc = Document()
            para = doc.add_paragraph("Grouped (")
            _add_citation_field(para, "_Bib_1", "Jiang et al., 2021")
            para.add_run("; ")
            _add_citation_field(para, "_OtherBookmark", "Smith et al., 2024")
            para.add_run(").")
            doc.save(docx)

            result = extract_docx_text(docx, bib_path=bib)
            latex = result["paragraphs"][0]["latex"]

            self.assertEqual(latex, "Grouped \\citep{1,abc-key}.")
            self.assertEqual(result["statistics"]["citations"], 2)


if __name__ == "__main__":
    unittest.main()
