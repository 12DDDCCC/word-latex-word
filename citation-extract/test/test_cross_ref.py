#!/usr/bin/env python3
"""文献交叉引用功能测试

测试 cross_ref_builder.py 的文献引用交叉引用功能:
1. insert_bib_cross_references - 在Word文档中插入bookmark和REF域代码
2. 验证参考文献条目有bookmark
3. 验证正文引用变为REF域代码(可点击跳转)
4. 验证CiteStyleConfig模板接口
5. 验证cite_style参数传递
"""
import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from cross_ref_builder import (
    insert_bib_cross_references,
    _bib_key_to_bookmark,
    set_update_fields_on_open,
    CiteStyleConfig,
    CITE_STYLES,
    detect_cite_style_from_tex,
    get_cite_style,
)
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


def test_bib_key_to_bookmark():
    print('[1/5] 测试 bookmark名转换...')
    assert _bib_key_to_bookmark('44') == '_Bib_44'
    assert _bib_key_to_bookmark('Wang2022') == '_Bib_Wang2022'
    assert _bib_key_to_bookmark('kondo2020a') == '_Bib_kondo2020a'
    print('  PASSED')
    return True


def test_cite_style_config():
    print('[2/5] 测试 CiteStyleConfig模板接口...')

    # 预置模板完整性
    assert 'apa' in CITE_STYLES
    assert 'ieee' in CITE_STYLES
    assert 'copernicus' in CITE_STYLES
    assert 'nature' in CITE_STYLES
    assert 'chicago' in CITE_STYLES
    assert 'hyperref_default' in CITE_STYLES
    assert 'hyperref_green' in CITE_STYLES
    assert 'hyperref_red' in CITE_STYLES
    assert 'numbered' in CITE_STYLES

    # get_cite_style
    style = get_cite_style('ieee')
    assert style.superscript is True
    assert style.cite_format == 'numbered'

    style = get_cite_style('apa')
    assert style.color == RGBColor(0, 0, 0)
    assert style.cite_format == 'author_year'

    # 自定义样式
    custom = CiteStyleConfig(name='test', color=RGBColor(0xFF, 0, 0), bold=True)
    assert get_cite_style(custom) is custom

    # None → 默认apa
    assert get_cite_style(None).name == 'apa'

    # 未知key → 默认apa
    assert get_cite_style('unknown').name == 'apa'

    print('  PASSED')
    return True


def test_detect_cite_style():
    print('[3/5] 测试 自动检测引用样式...')

    # IEEE documentclass
    assert detect_cite_style_from_tex(r'\documentclass{IEEEtran}') == 'ieee'

    # Copernicus documentclass
    assert detect_cite_style_from_tex(r'\documentclass{acp}') == 'copernicus'

    # hyperref citecolor
    assert detect_cite_style_from_tex(r'\usepackage[citecolor=green]{hyperref}') == 'hyperref_green'
    assert detect_cite_style_from_tex(r'\usepackage[citecolor=red]{hyperref}') == 'hyperref_red'
    assert detect_cite_style_from_tex(r'\usepackage[citecolor=black]{hyperref}') == 'apa'

    # bibliographystyle
    assert detect_cite_style_from_tex(r'\bibliographystyle{IEEEtran}') == 'ieee'
    assert detect_cite_style_from_tex(r'\bibliographystyle{plain}') == 'numbered'
    assert detect_cite_style_from_tex(r'\bibliographystyle{apacite}') == 'apa'

    # natbib → apa
    assert detect_cite_style_from_tex(r'\usepackage{natbib}') == 'apa'

    # 默认
    assert detect_cite_style_from_tex(r'\documentclass{article}') == 'apa'

    print('  PASSED')
    return True


def test_insert_with_style():
    print('[4/5] 测试 带cite_style的文献交叉引用插入...')

    doc = Document()
    doc.add_paragraph('The method was proposed by Kondo et al. (2020).')
    doc.add_paragraph('This is consistent with (Wang et al., 2022).')
    doc.add_paragraph('References')
    doc.add_paragraph('Kondo, T., ... (2020). Title. Journal, 1, 1-10.')
    doc.add_paragraph('Wang, Ding, and Ma (2022). Title. Journal, 2, 1-10.')

    cite_map = {
        '12': 'Kondo et al. (2020)',
        '44': 'Wang et al. (2022)',
    }

    # 测试不同样式
    for style_key in ['apa', 'ieee', 'copernicus']:
        doc_test = Document()
        doc_test.add_paragraph('The method was proposed by Kondo et al. (2020).')
        doc_test.add_paragraph('This is consistent with (Wang et al., 2022).')
        doc_test.add_paragraph('References')
        doc_test.add_paragraph('Kondo, T., ... (2020). Title. Journal, 1, 1-10.')
        doc_test.add_paragraph('Wang, Ding, and Ma (2022). Title. Journal, 2, 1-10.')

        result = insert_bib_cross_references(doc_test, cite_map, cite_style=style_key)
        print(f'  样式={style_key}: {result}')

        # 验证bookmark
        bm_names = [bm.get(qn('w:name')) for bm in doc_test.element.body.iter(qn('w:bookmarkStart'))]
        assert '_Bib_12' in bm_names, f'[{style_key}] _Bib_12未找到'

        # 验证内部HYPERLINK域代码
        instr_texts = [it.text for it in doc_test.element.body.iter(qn('w:instrText')) if it.text and 'HYPERLINK' in it.text]
        assert len(instr_texts) > 0, f'[{style_key}] 未找到HYPERLINK域代码'

    print('  PASSED')
    return True


def test_save_and_verify():
    print('[5/5] 测试 保存文档验证...')

    doc = Document()
    doc.add_paragraph('The result (Kondo et al., 2020) is significant.')
    doc.add_paragraph('References')
    doc.add_paragraph('Kondo, T. (2020). Title. Journal.')

    cite_map = {'12': 'Kondo et al. (2020)'}
    result = insert_bib_cross_references(doc, cite_map, cite_style='copernicus')

    test_dir = os.path.dirname(__file__)
    output_path = os.path.join(test_dir, 'test_bib_crossref.docx')
    doc.save(output_path)

    doc2 = Document(output_path)
    bm_names = [bm.get(qn('w:name')) for bm in doc2.element.body.iter(qn('w:bookmarkStart'))]
    instr_texts = [it.text for it in doc2.element.body.iter(qn('w:instrText')) if it.text and 'HYPERLINK' in it.text]
    assert len(bm_names) > 0, '保存后bookmark丢失'
    assert len(instr_texts) > 0, '保存后HYPERLINK域丢失'

    try:
        os.remove(output_path)
    except Exception:
        pass

    print('  PASSED')
    return True


if __name__ == '__main__':
    print('=' * 50)
    print('  文献交叉引用功能测试 (v1.1)')
    print('=' * 50)

    all_pass = True
    for fn in [test_bib_key_to_bookmark, test_cite_style_config,
               test_detect_cite_style, test_insert_with_style, test_save_and_verify]:
        try:
            if not fn():
                all_pass = False
        except Exception as e:
            print(f'  FAILED - {e}')
            import traceback
            traceback.print_exc()
            all_pass = False

    print()
    print('=' * 50)
    print(f'  {"ALL TESTS PASSED" if all_pass else "SOME TESTS FAILED"}')
    print('=' * 50)
