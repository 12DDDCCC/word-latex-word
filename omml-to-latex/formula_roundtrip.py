r"""公式双向无损转换测试器

完整链路: Word OMML → LaTeX → 编译验证 → OMML → Word
对比原始公式与转换后公式的等价性，生成测试报告。

使用方法:
    python formula_roundtrip.py <docx路径> [输出目录]
    python formula_roundtrip.py --tex <tex路径> [输出目录]
"""

import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

# 本模块目录
_THIS_DIR = Path(__file__).parent
sys.path.insert(0, str(_THIS_DIR))

from omml_to_latex import extract_formulas, omml_to_latex
from latex_to_omml import (
    latex_to_omml,
    latex_to_omml_element,
    extract_formulas_from_tex,
)

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from lxml import etree


# ============================================================
# 核心转换：单条公式 roundtrip
# ============================================================

def roundtrip_single(latex_str, xslt_path=None):
    """单条公式双向转换：LaTeX → OMML → LaTeX

    Args:
        latex_str: LaTeX 公式代码
        xslt_path: MML2OMML.XSL 路径

    Returns:
        dict: {
            'original_latex': 原始LaTeX,
            'omml': OMML XML字符串,
            'omml_element': lxml Element,
            'roundtrip_latex': 回转后的LaTeX,
            'status': 'ok' / 'failed',
            'error': 错误信息 或 None,
        }
    """
    result = {
        'original_latex': latex_str,
        'omml': None,
        'omml_element': None,
        'roundtrip_latex': None,
        'status': 'failed',
        'error': None,
    }

    # Step 1: LaTeX → OMML
    omml_str = latex_to_omml(latex_str, xslt_path)
    if omml_str is None:
        result['error'] = 'LaTeX→OMML conversion failed'
        return result
    result['omml'] = omml_str

    # Step 2: OMML Element
    omml_elem = latex_to_omml_element(latex_str, xslt_path)
    if omml_elem is None:
        result['error'] = 'OMML element creation failed'
        return result
    result['omml_element'] = omml_elem

    # Step 3: OMML → LaTeX (roundtrip)
    try:
        # omml_to_latex 需要的标签格式：确保命名空间正确
        roundtrip = omml_to_latex(omml_elem)
        if roundtrip:
            result['roundtrip_latex'] = roundtrip
            result['status'] = 'ok'
        else:
            result['error'] = 'OMML→LaTeX returned empty'
    except Exception as e:
        result['error'] = f'OMML→LaTeX error: {e}'

    return result


def compare_latex(original, roundtrip):
    """比较原始LaTeX和回转LaTeX的等价性

    归一化策略（3级比较）:
    1. exact: 原文完全一致
    2. normalized: 去空格 + 统一命令 + 统一括号 + 统一减号 + 去小间距
    3. semantic: 去掉所有花括号后比较核心结构

    Returns:
        dict: {
            'exact_match': bool,
            'normalized_match': bool,
            'semantic_match': bool,
            'original_normalized': str,
            'roundtrip_normalized': str,
            'diff_note': str,
        }
    """
    def normalize(s):
        """归一化：统一格式差异"""
        s = re.sub(r'\s+', '', s)
        # 统一减号: U+2212 (−) → ASCII -
        s = s.replace('−', '-')
        # 去除 \left \right
        s = s.replace('\\left', '').replace('\\right', '')
        # 去除小间距 \, \; \! \quad
        s = re.sub(r'\\[,;!]', '', s)
        s = re.sub(r'\\quad', '', s)
        # 统一命令别名
        s = s.replace('\\mathrm', '\\rm')
        s = s.replace('\\mathbf', '\\bf')
        s = s.replace('\\boldsymbol', '\\bs')
        s = s.replace('\\mathit', '\\it')
        # 统一 \overline → \bar (OMML bar 映射)
        s = s.replace('\\overline', '\\bar')
        return s

    def semantic_normalize(s):
        """语义归一化：去除花括号，只保留核心结构"""
        s = normalize(s)
        # 去除花括号（保留内容）
        s = s.replace('{', '').replace('}', '')
        return s

    orig_n = normalize(original)
    rt_n = normalize(roundtrip)
    orig_s = semantic_normalize(original)
    rt_s = semantic_normalize(roundtrip)

    exact = original.strip() == roundtrip.strip()
    norm = orig_n == rt_n
    sem = orig_s == rt_s

    note = ''
    if exact:
        note = '完全一致'
    elif norm:
        note = '归一化后一致（空格/格式差异）'
    elif sem:
        note = '语义一致（花括号/格式差异）'
    else:
        # 找出归一化后的第一个差异
        for i in range(min(len(orig_n), len(rt_n))):
            if orig_n[i] != rt_n[i]:
                ctx = 8
                note = (f'差异@{i}: '
                        f'原文="...{orig_n[max(0,i-ctx):i+ctx]}...", '
                        f'回转="...{rt_n[max(0,i-ctx):i+ctx]}..."')
                break
        if not note:
            note = f'长度不同: 原文={len(orig_n)}, 回转={len(rt_n)}'

    return {
        'exact_match': exact,
        'normalized_match': norm,
        'semantic_match': sem,
        'original_normalized': orig_n,
        'roundtrip_normalized': rt_n,
        'diff_note': note,
    }


# ============================================================
# 编译验证
# ============================================================

def compile_latex_formulas(formulas_list, output_dir, xelatex='xelatex'):
    """将公式列表编译为PDF进行视觉验证

    Args:
        formulas_list: [(label, latex_str), ...]
        output_dir: 输出目录
        xelatex: xelatex 可执行文件路径

    Returns:
        dict: {'pdf_path': str or None, 'tex_path': str, 'errors': list}
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    tex_path = output_dir / 'roundtrip_test.tex'

    lines = [
        r'\documentclass{article}',
        r'\usepackage[UTF8,fontset=windows]{ctex}',
        r'\usepackage{amsmath,amssymb}',
        r'\usepackage{geometry}',
        r'\geometry{a4paper,margin=2cm}',
        r'\begin{document}',
        '',
        r'\section*{公式双向转换编译测试}',
        '',
    ]

    for i, (label, latex) in enumerate(formulas_list):
        lines.append(f'% --- 公式 {i+1}: {label} ---')
        lines.append(r'\begin{equation}')
        lines.append(f'  {latex}')
        lines.append(r'\end{equation}')
        lines.append('')

    lines.append(r'\end{document}')

    tex_path.write_text('\n'.join(lines), encoding='utf-8')

    # 编译
    result = subprocess.run(
        [xelatex, '-interaction=nonstopmode', str(tex_path)],
        capture_output=True, text=True, encoding='utf-8',
        errors='replace', cwd=str(output_dir),
        timeout=60,
    )

    pdf_path = output_dir / 'roundtrip_test.pdf'
    errors = []

    if not pdf_path.exists():
        # 从日志提取错误
        log_path = output_dir / 'roundtrip_test.log'
        if log_path.exists():
            for line in log_path.read_text(encoding='utf-8', errors='ignore').split('\n'):
                if line.startswith('!'):
                    errors.append(line.strip())

    return {
        'pdf_path': str(pdf_path) if pdf_path.exists() else None,
        'tex_path': str(tex_path),
        'errors': errors,
    }


# ============================================================
# 生成测试Word文档
# ============================================================

def generate_roundtrip_docx(formulas_data, output_path):
    """生成包含原始公式和回转公式的Word文档，方便人工对比

    Args:
        formulas_data: [{
            'original_latex': str,
            'roundtrip_latex': str or None,
            'omml_element': Element or None,
            'status': str,
        }]
        output_path: 输出.docx路径
    """
    doc = Document()

    # 标题
    title = doc.add_heading('公式双向转换测试报告', level=1)
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'公式总数: {len(formulas_data)}')

    ok_count = sum(1 for f in formulas_data if f['status'] == 'ok')
    fail_count = len(formulas_data) - ok_count
    doc.add_paragraph(f'成功: {ok_count}, 失败: {fail_count}')

    # 逐条公式
    for i, f in enumerate(formulas_data):
        doc.add_heading(f'公式 {i+1}', level=2)

        # 状态标记
        status_text = 'OK' if f['status'] == 'ok' else 'FAILED'
        para = doc.add_paragraph()
        run = para.add_run(f'状态: {status_text}')
        run.bold = True

        # 原始LaTeX
        para = doc.add_paragraph()
        para.add_run('原始 LaTeX: ').bold = True
        para.add_run(f['original_latex'])

        # 原始公式（OMML插入）
        if f.get('omml_element') is not None:
            doc.add_paragraph('转换后公式:')
            math_para = doc.add_paragraph()
            try:
                math_para._element.append(f['omml_element'])
            except Exception:
                doc.add_paragraph('[插入失败]')

        # 回转LaTeX
        if f.get('roundtrip_latex'):
            para = doc.add_paragraph()
            para.add_run('回转 LaTeX: ').bold = True
            para.add_run(f['roundtrip_latex'])

        # 错误信息
        if f.get('error'):
            para = doc.add_paragraph()
            run = para.add_run(f'错误: {f["error"]}')
            run.font.color.rgb = RGBColor(192, 0, 0)

        doc.add_paragraph('')  # 空行分隔

    doc.save(str(output_path))
    return str(output_path)


# ============================================================
# 生成测试报告
# ============================================================

def generate_report(formulas_data, comparisons, output_dir):
    """生成Markdown测试报告

    Args:
        formulas_data: roundtrip结果列表
        comparisons: compare结果列表
        output_dir: 输出目录

    Returns:
        str: 报告文件路径
    """
    output_dir = Path(output_dir)
    report_path = output_dir / 'roundtrip_report.md'

    ok_count = sum(1 for f in formulas_data if f['status'] == 'ok')
    exact_count = sum(1 for c in comparisons if c['exact_match'])
    norm_count = sum(1 for c in comparisons if c['normalized_match'])
    sem_count = sum(1 for c in comparisons if c.get('semantic_match', False))
    total = len(formulas_data)

    def pct(count):
        return count / total * 100 if total else 0.0

    lines = [
        '# 公式双向无损转换测试报告',
        '',
        f'**测试时间**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
        f'**公式总数**: {len(formulas_data)}',
        '',
        '## 总结',
        '',
        f'| 指标 | 数量 | 比率 |',
        f'|------|------|------|',
        f'| OMML转换成功 | {ok_count} | {pct(ok_count):.1f}% |',
        f'| 完全一致 | {exact_count} | {pct(exact_count):.1f}% |',
        f'| 归一化一致 | {norm_count} | {pct(norm_count):.1f}% |',
        f'| 语义一致 | {sem_count} | {pct(sem_count):.1f}% |',
        '',
    ]

    # 失败列表
    failed = [(i, f) for i, f in enumerate(formulas_data) if f['status'] != 'ok']
    if failed:
        lines.append('## 转换失败详情')
        lines.append('')
        for idx, f in failed:
            lines.append(f'### 公式 {idx+1}')
            lines.append(f'- **LaTeX**: `{f["original_latex"]}`')
            lines.append(f'- **错误**: {f["error"]}')
            lines.append('')

    # 归一化一致但不完全一致的
    diff_cases = [
        (i, f, c) for i, (f, c) in enumerate(zip(formulas_data, comparisons))
        if c['normalized_match'] and not c['exact_match']
    ]
    if diff_cases:
        lines.append('## 格式差异（归一化一致）')
        lines.append('')
        for idx, f, c in diff_cases:
            lines.append(f'### 公式 {idx+1}')
            lines.append(f'- **原始**: `{f["original_latex"]}`')
            lines.append(f'- **回转**: `{f["roundtrip_latex"]}`')
            lines.append(f'- **说明**: {c["diff_note"]}')
            lines.append('')

    # 不一致列表
    mismatch = [
        (i, f, c) for i, (f, c) in enumerate(zip(formulas_data, comparisons))
        if not c['normalized_match'] and f['status'] == 'ok'
    ]
    if mismatch:
        lines.append('## 内容差异（归一化不一致）')
        lines.append('')
        for idx, f, c in mismatch:
            lines.append(f'### 公式 {idx+1}')
            lines.append(f'- **原始**: `{f["original_latex"]}`')
            lines.append(f'- **回转**: `{f["roundtrip_latex"]}`')
            lines.append(f'- **差异**: {c["diff_note"]}')
            lines.append('')

    report_path.write_text('\n'.join(lines), encoding='utf-8')
    return str(report_path)


# ============================================================
# 主流程
# ============================================================

def run_roundtrip_from_docx(docx_path, output_dir, xslt_path=None):
    """从Word文档执行完整的双向转换测试

    Args:
        docx_path: Word文档路径
        output_dir: 输出目录
        xslt_path: MML2OMML.XSL 路径

    Returns:
        dict: 测试结果摘要
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f'[1/5] 从Word提取公式...')
    formulas = extract_formulas(docx_path)
    print(f'  提取到 {len(formulas)} 个公式')

    print(f'[2/5] 执行双向转换 (LaTeX → OMML → LaTeX)...')
    roundtrip_results = []
    for i, f in enumerate(formulas):
        if (i + 1) % 10 == 0 or i == len(formulas) - 1:
            print(f'  [{i+1}/{len(formulas)}] 处理中...')
        rt = roundtrip_single(f['latex'], xslt_path)
        rt['source_type'] = f['type']
        rt['source_para'] = f['para_index']
        roundtrip_results.append(rt)

    print(f'[3/5] 比较等价性...')
    comparisons = []
    for rt in roundtrip_results:
        if rt['status'] == 'ok' and rt['roundtrip_latex']:
            comp = compare_latex(rt['original_latex'], rt['roundtrip_latex'])
        else:
            comp = {
                'exact_match': False,
                'normalized_match': False,
                'semantic_match': False,
                'original_normalized': '',
                'roundtrip_normalized': '',
                'diff_note': '转换失败，无法比较',
            }
        comparisons.append(comp)

    # 统计
    ok_count = sum(1 for r in roundtrip_results if r['status'] == 'ok')
    exact_count = sum(1 for c in comparisons if c['exact_match'])
    norm_count = sum(1 for c in comparisons if c['normalized_match'])
    sem_count = sum(1 for c in comparisons if c.get('semantic_match', False))
    print(f'  OMML转换成功: {ok_count}/{len(formulas)}')
    print(f'  完全一致: {exact_count}/{len(formulas)}')
    print(f'  归一化一致: {norm_count}/{len(formulas)}')
    print(f'  语义一致: {sem_count}/{len(formulas)}')

    print(f'[4/5] 编译验证...')
    compile_formulas = [
        (f'eq{i+1}', rt['original_latex'])
        for i, rt in enumerate(roundtrip_results) if rt['status'] == 'ok'
    ]
    if compile_formulas:
        compile_result = compile_latex_formulas(compile_formulas, output_dir / 'compile')
        if compile_result['pdf_path']:
            print(f'  PDF生成成功: {compile_result["pdf_path"]}')
        else:
            print(f'  编译失败: {compile_result["errors"][:3]}')
    else:
        compile_result = {'pdf_path': None, 'errors': ['no formulas to compile']}

    print(f'[5/5] 生成报告...')
    report_path = generate_report(roundtrip_results, comparisons, output_dir)
    print(f'  报告: {report_path}')

    docx_out = output_dir / 'roundtrip_test.docx'
    generate_roundtrip_docx(roundtrip_results, docx_out)
    print(f'  测试Word: {docx_out}')

    # 保存JSON数据
    json_data = []
    for rt, comp in zip(roundtrip_results, comparisons):
        entry = {
            'original_latex': rt['original_latex'],
            'roundtrip_latex': rt.get('roundtrip_latex'),
            'status': rt['status'],
            'error': rt.get('error'),
            'exact_match': comp['exact_match'],
            'normalized_match': comp['normalized_match'],
            'semantic_match': comp.get('semantic_match', False),
            'diff_note': comp['diff_note'],
        }
        json_data.append(entry)
    json_path = output_dir / 'roundtrip_data.json'
    json_path.write_text(
        json.dumps(json_data, ensure_ascii=False, indent=2),
        encoding='utf-8',
    )

    return {
        'total': len(formulas),
        'omml_ok': ok_count,
        'exact_match': exact_count,
        'normalized_match': norm_count,
        'semantic_match': sem_count,
        'report_path': report_path,
        'docx_path': str(docx_out),
        'json_path': str(json_path),
        'compile_pdf': compile_result.get('pdf_path'),
    }


def run_roundtrip_from_tex(tex_path, output_dir, xslt_path=None):
    """从LaTeX文件执行双向转换测试

    Args:
        tex_path: .tex 文件路径
        output_dir: 输出目录
        xslt_path: MML2OMML.XSL 路径

    Returns:
        dict: 测试结果摘要
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f'[1/4] 从LaTeX文件提取公式...')
    formulas = extract_formulas_from_tex(tex_path)
    print(f'  提取到 {len(formulas)} 个公式')

    print(f'[2/4] 执行双向转换...')
    roundtrip_results = []
    for i, f in enumerate(formulas):
        if (i + 1) % 10 == 0 or i == len(formulas) - 1:
            print(f'  [{i+1}/{len(formulas)}] 处理中...')
        rt = roundtrip_single(f['latex'], xslt_path)
        rt['source_type'] = f['type']
        rt['source_env'] = f['env']
        roundtrip_results.append(rt)

    print(f'[3/4] 比较等价性...')
    comparisons = []
    for rt in roundtrip_results:
        if rt['status'] == 'ok' and rt['roundtrip_latex']:
            comp = compare_latex(rt['original_latex'], rt['roundtrip_latex'])
        else:
            comp = {
                'exact_match': False, 'normalized_match': False,
                'semantic_match': False,
                'original_normalized': '', 'roundtrip_normalized': '',
                'diff_note': '转换失败，无法比较',
            }
        comparisons.append(comp)

    ok_count = sum(1 for r in roundtrip_results if r['status'] == 'ok')
    exact_count = sum(1 for c in comparisons if c['exact_match'])
    norm_count = sum(1 for c in comparisons if c['normalized_match'])
    sem_count = sum(1 for c in comparisons if c.get('semantic_match', False))

    print(f'[4/4] 生成报告...')
    report_path = generate_report(roundtrip_results, comparisons, output_dir)

    docx_out = output_dir / 'roundtrip_test.docx'
    generate_roundtrip_docx(roundtrip_results, docx_out)

    print(f'  OMML转换成功: {ok_count}/{len(formulas)}')
    print(f'  完全一致: {exact_count}/{len(formulas)}')
    print(f'  归一化一致: {norm_count}/{len(formulas)}')
    print(f'  语义一致: {sem_count}/{len(formulas)}')

    return {
        'total': len(formulas),
        'omml_ok': ok_count,
        'exact_match': exact_count,
        'normalized_match': norm_count,
        'semantic_match': sem_count,
        'report_path': report_path,
        'docx_path': str(docx_out),
    }


# ============================================================
# CLI
# ============================================================

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("公式双向无损转换测试器")
        print()
        print("用法:")
        print("  python formula_roundtrip.py <docx路径> [输出目录]")
        print("  python formula_roundtrip.py --tex <tex路径> [输出目录]")
        print()
        print("示例:")
        print("  python formula_roundtrip.py paper.docx test/formula_roundtrip")
        print("  python formula_roundtrip.py --tex main.tex test/tex_roundtrip")
        sys.exit(1)

    xslt = None
    if '--xslt' in sys.argv:
        xi = sys.argv.index('--xslt')
        if xi + 1 < len(sys.argv):
            xslt = sys.argv[xi + 1]

    if sys.argv[1] == '--tex':
        if len(sys.argv) < 3:
            print("错误: --tex 需要指定 .tex 文件路径")
            sys.exit(1)
        tex_path = sys.argv[2]
        out_dir = sys.argv[3] if len(sys.argv) > 3 else 'test/tex_roundtrip'
        summary = run_roundtrip_from_tex(tex_path, out_dir, xslt)
    else:
        docx_path = sys.argv[1]
        out_dir = sys.argv[2] if len(sys.argv) > 2 else 'test/formula_roundtrip'
        summary = run_roundtrip_from_docx(docx_path, out_dir, xslt)

    print()
    print('=' * 50)
    print('测试完成！')
    print(f'  总公式数: {summary["total"]}')
    print(f'  OMML转换成功: {summary["omml_ok"]}')
    print(f'  完全一致: {summary["exact_match"]}')
    print(f'  归一化一致: {summary["normalized_match"]}')
