import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SKILL_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(SKILL_DIR))
sys.path.insert(0, str(SKILL_DIR / 'shared'))
sys.path.insert(0, str(SKILL_DIR / 'citation-extract'))
sys.path.insert(0, str(SKILL_DIR / 'tex-to-word'))

from cross_ref_builder import insert_bib_cross_references
from _docx_insert import _append_table_caption_paragraph_props
from latex_text_utils import citation_marker, clean_latex_text


class CitationAndCaptionFixesTest(unittest.TestCase):
    def test_citation_marker_is_pandoc_safe_plain_text(self):
        marker = citation_marker('key_with-symbols', 'M')
        self.assertRegex(marker, r'^[A-Za-z0-9]+$')

    def test_latex_percent_is_cleaned_for_caption(self):
        self.assertEqual(clean_latex_text(r'数据占比大于75\%'), '数据占比大于75%')

    def test_table_caption_does_not_force_page_break(self):
        pPr = OxmlElement('w:pPr')

        _append_table_caption_paragraph_props(pPr)

        self.assertIsNotNone(pPr.find(qn('w:keepNext')))
        self.assertIsNotNone(pPr.find(qn('w:keepLines')))
        self.assertIsNone(pPr.find(qn('w:pageBreakBefore')))

    def test_multi_citation_markers_create_distinct_links(self):
        doc = Document()
        first_group = '; '.join([
            citation_marker('22', 'M'),
            citation_marker('23', 'M'),
            citation_marker('24', 'M'),
        ])
        second_group = '; '.join([
            citation_marker('20', 'M'),
            citation_marker('21', 'M'),
        ])
        doc.add_paragraph(f'Result ({first_group}). Middle ({second_group}). End.')
        doc.add_paragraph('References')
        doc.add_paragraph('Terhaar et al. (2025). First.')
        doc.add_paragraph('Yao et al. (2026). Second.')
        doc.add_paragraph('Kelley et al. (2025). Third.')
        doc.add_paragraph('Gui et al. (2024). Fourth.')
        doc.add_paragraph('Ke et al. (2024). Fifth.')
        cite_map = {
            '22': 'Terhaar et al. (2025)',
            '23': 'Yao et al. (2026)',
            '24': 'Kelley et al. (2025)',
            '20': 'Gui et al. (2024)',
            '21': 'Ke et al. (2024)',
        }

        result = insert_bib_cross_references(doc, cite_map, cite_style='copernicus')
        fields = [
            node.text
            for node in doc.element.body.iter(qn('w:instrText'))
            if node.text and 'HYPERLINK' in node.text
        ]

        self.assertEqual(result['refs_replaced'], 5)
        self.assertEqual(len(fields), 5)
        self.assertEqual(
            doc.paragraphs[0].text,
            'Result (Terhaar et al., 2025; Yao et al., 2026; Kelley et al., 2025). '
            'Middle (Gui et al., 2024; Ke et al., 2024). End.',
        )


if __name__ == '__main__':
    unittest.main()
