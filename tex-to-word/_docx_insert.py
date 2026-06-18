#!/usr/bin/env python3
r"""DOCX后处理模块 — Word文档的公式插入、样式修复、图片嵌入、表格重建

包含Pandoc输出后的所有Word文档修改操作。
"""
import copy
import math
import json, re, os, sys
from pathlib import Path

# 导入共享工具
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from shared.latex_text_utils import clean_latex_text, to_subscript, to_superscript

# 导入交叉引用构建器
_CROSSREF_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'citation-extract')
if os.path.isdir(_CROSSREF_DIR):
    sys.path.insert(0, _CROSSREF_DIR)
try:
    from cross_ref_builder import insert_bib_cross_references
    from _bbl_item_parser import parse_bbl_items
    from _ref_section_builder import build_references_section
    _HAS_CROSSREF = True
except ImportError:
    _HAS_CROSSREF = False

# 导入 latex_to_omml skill
_SKILL_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'omml-to-latex')
if os.path.isdir(_SKILL_DIR):
    sys.path.insert(0, _SKILL_DIR)
try:
    from latex_to_omml import (
        latex_to_omml as _latex_to_omml_skill,
        gather_to_display as _gather_skill,
        equation_to_display as _equation_skill,
    )
    _HAS_SKILL = True
except ImportError:
    _HAS_SKILL = False


def _clean_latex_text(text):
    """清理LaTeX文本中的命令, 转为纯文本显示 (委托给 shared)"""
    return clean_latex_text(text)


def _to_subscript(s):
    """Unicode下标 (委托给 shared)"""
    return to_subscript(s)


def _to_superscript(s):
    """Unicode上标 (委托给 shared)"""
    return to_superscript(s)


def _append_zero_paragraph_spacing(pPr):
    """Set direct paragraph spacing to LaTeX-like compact float spacing."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def _clear_image_para_indent_markers(pPr):
    """清除图片段落的首行缩进和列表标记,防止图片左侧出现黑色小方块。

    图片段继承正文 firstLine 缩进时,inline drawing 会在缩进位置渲染出
    段落标记符/边距填充,表现为黑色小方块。图片段应为无缩进居中。
    """
    from docx.oxml.ns import qn

    # 1. 移除 numPr(列表标记 → 黑方块)
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)
    # 2. 清零首行缩进(图片不需要正文缩进)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        from docx.oxml import OxmlElement
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    if ind.get(qn("w:hanging")):
        ind.attrib.pop(qn("w:hanging"), None)


def _set_paragraph_spacing_twips(pPr, before=None, after=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn("w:before"), str(before))
    if after is not None:
        spacing.set(qn("w:after"), str(after))


def _append_caption_paragraph_props(pPr, caption_style=None, keep_next=False):
    """Apply template-derived paragraph properties for figure/table captions."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    caption_style = caption_style or {}
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), caption_style.get("alignment") or "left")
    pPr.append(jc)
    _set_paragraph_spacing_twips(
        pPr,
        before=caption_style.get("space_before_twips"),
        after=caption_style.get("space_after_twips"),
    )
    if keep_next:
        pPr.append(OxmlElement("w:keepNext"))
    pPr.append(OxmlElement("w:keepLines"))


def _append_table_caption_paragraph_props(pPr, caption_style=None):
    _append_caption_paragraph_props(pPr, caption_style=caption_style, keep_next=True)


def latex_to_omml(latex_str):
    """LaTeX → OMML (委托给 latex_to_omml skill)"""
    if _HAS_SKILL:
        return _latex_to_omml_skill(latex_str)
    return None


def _page_content_width_twips(page_geometry, single_column=False):
    """Return template-derived text or column width in twips."""
    if not page_geometry:
        return 0
    try:
        width_mm = float(page_geometry.get('textwidth_mm') or 0)
        column_count = int(page_geometry.get('column_count', 1) or 1)
        column_sep_mm = float(page_geometry.get('column_sep_mm', 0) or 0)
    except (TypeError, ValueError):
        return 0
    if single_column and column_count > 1:
        width_mm = (width_mm - (column_count - 1) * column_sep_mm) / column_count
    return int(round(width_mm * 56.6929)) if width_mm > 0 else 0


def _section_content_width_twips(section, single_column=False):
    """Return the active Word section's text or column width in twips."""
    from docx.oxml.ns import qn

    dims = (section.page_width, section.left_margin, section.right_margin)
    if not all(dim is not None for dim in dims):
        return 0
    width = int((dims[0] - dims[1] - dims[2]) / 635)
    cols = section._sectPr.find(qn('w:cols'))
    if not single_column or cols is None:
        return width
    try:
        column_count = int(cols.get(qn('w:num')) or 1)
        column_space = int(cols.get(qn('w:space')) or 0)
    except (TypeError, ValueError):
        return width
    return int((width - (column_count - 1) * column_space) / column_count)


def _compiled_layout_twips(tex_dir):
    """Read effective dimensions emitted by the compiled LaTeX document."""
    values = {}
    for log_path in Path(tex_dir).glob("*.log"):
        content = _read_text_file(log_path) or ''
        for name, value in re.findall(
                r"SKILL-LAYOUT-(TEXTWIDTH|COLUMNWIDTH|COLUMNSEP)=([\d.]+)pt",
                content):
            values[name.lower()] = int(round(float(value) * 20))
    return values


def _base_section_columns(doc):
    from docx.oxml.ns import qn

    sectPr = doc.element.body.find(qn("w:sectPr"))
    cols = sectPr.find(qn("w:cols")) if sectPr is not None else None
    if cols is None:
        return 1
    try:
        return int(cols.get(qn("w:num")) or 1)
    except (TypeError, ValueError):
        return 1


def _section_break_para(doc, column_count):
    """Create a continuous section break paragraph with template page settings."""
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    base = doc.element.body.find(qn("w:sectPr"))
    sectPr = deepcopy(base) if base is not None else OxmlElement("w:sectPr")
    sect_type = sectPr.find(qn("w:type"))
    if sect_type is None:
        sect_type = OxmlElement("w:type")
        sectPr.insert(0, sect_type)
    sect_type.set(qn("w:val"), "continuous")

    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(max(int(column_count or 1), 1)))

    para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pPr.append(sectPr)
    para.append(pPr)
    _compact_section_break_para(para)
    return para


def _section_break_column_count(p_elem):
    from docx.oxml.ns import qn

    sectPr = p_elem.find(f'{qn("w:pPr")}/{qn("w:sectPr")}')
    cols = sectPr.find(qn("w:cols")) if sectPr is not None else None
    if cols is None:
        return None
    try:
        return int(cols.get(qn("w:num")) or 1)
    except (TypeError, ValueError):
        return None


def _paragraph_xml_text(p_elem):
    from docx.oxml.ns import qn

    return ''.join(node.text or '' for node in p_elem.iter(qn("w:t")))


def _is_empty_section_break_para(p_elem):
    return (
        _section_break_column_count(p_elem) is not None
        and not _paragraph_xml_text(p_elem).strip()
    )


def _compact_section_break_para(p_elem):
    """Keep empty continuous section-break paragraphs visually zero-height."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:ind", "w:jc"):
        for elem in list(pPr.findall(qn(tag))):
            pPr.remove(elem)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        sectPr = pPr.find(qn("w:sectPr"))
        if sectPr is not None:
            sectPr.addprevious(spacing)
        else:
            pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "1")
    spacing.set(qn("w:lineRule"), "exact")


def _table_description(tbl_elem):
    from docx.oxml.ns import qn

    desc = tbl_elem.find(f'{qn("w:tblPr")}/{qn("w:tblDescription")}')
    return (desc.get(qn("w:val")) or "") if desc is not None else ""


def _is_layout_floating_table_elem(tbl_elem):
    from docx.oxml.ns import qn

    if tbl_elem.find(f'{qn("w:tblPr")}/{qn("w:tblpPr")}') is None:
        return False
    return _table_description(tbl_elem).startswith("skill-")


def _is_front_matter_float_elem(tbl_elem):
    return _table_description(tbl_elem) == "skill-front-matter"


def _heading_anchor_element(target_elem):
    from docx.oxml.ns import qn

    anchor = target_elem
    prev = anchor.getprevious()
    while prev is not None and prev.tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
        anchor = prev
        prev = anchor.getprevious()
    return anchor


_TBL_PR_ORDER = {
    name: idx for idx, name in enumerate((
        "tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
        "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
        "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
        "tblCellMar", "tblLook", "tblCaption", "tblDescription",
    ))
}

_TC_PR_ORDER = {
    name: idx for idx, name in enumerate((
        "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge",
        "tcBorders", "shd", "noWrap", "tcMar", "textDirection",
        "tcFitText", "vAlign", "hideMark", "headers", "cellIns",
        "cellDel", "cellMerge", "tcPrChange",
    ))
}


def _local_name(elem):
    return elem.tag.rsplit("}", 1)[-1]


def _append_ordered_child(parent, child, order):
    child_rank = order.get(_local_name(child))
    if child_rank is None:
        parent.append(child)
        return child
    for idx, existing in enumerate(list(parent)):
        rank = order.get(_local_name(existing))
        if rank is not None and rank > child_rank:
            parent.insert(idx, child)
            return child
    parent.append(child)
    return child


def _append_tbl_pr_child(tbl_pr, child):
    return _append_ordered_child(tbl_pr, child, _TBL_PR_ORDER)


def _append_tc_pr_child(tc_pr, child):
    return _append_ordered_child(tc_pr, child, _TC_PR_ORDER)


def _append_nil_borders(parent, edge_tag="w:tblBorders"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement(edge_tag)
    edges = ("top", "left", "bottom", "right")
    if edge_tag == "w:tblBorders":
        edges += ("insideH", "insideV")
    for edge in edges:
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    if edge_tag == "w:tblBorders" and _local_name(parent) == "tblPr":
        _append_tbl_pr_child(parent, borders)
    elif edge_tag == "w:tcBorders" and _local_name(parent) == "tcPr":
        _append_tc_pr_child(parent, borders)
    else:
        parent.append(borders)


def _clear_table_borders(tbl_elem):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    for borders in list(tblPr.findall(qn("w:tblBorders"))):
        tblPr.remove(borders)
    _append_nil_borders(tblPr)

    for tr in tbl_elem.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            for borders in list(tcPr.findall(qn("w:tcBorders"))):
                tcPr.remove(borders)
            _append_nil_borders(tcPr, edge_tag="w:tcBorders")


def _clear_layout_floating_table_borders(doc):
    from docx.oxml.ns import qn

    for tbl in doc.element.body.iter(qn("w:tbl")):
        if _is_layout_floating_table_elem(tbl):
            _clear_table_borders(tbl)


def _append_zero_cell_margins(parent):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    margins = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:w"), "0")
        elem.set(qn("w:type"), "dxa")
        margins.append(elem)
    if _local_name(parent) == "tblPr":
        _append_tbl_pr_child(parent, margins)
    else:
        parent.append(margins)


def _full_width_block_elements(doc, elements):
    """Wrap double-column floats with native continuous section breaks."""
    base_columns = _base_section_columns(doc)
    elements = [elem for elem in elements if elem is not None]
    if base_columns < 2 or not elements:
        return elements
    # In Word XML, paragraph sectPr closes the section that precedes it.
    # So the leading break preserves the previous two-column section, while
    # the trailing break marks the inserted block itself as single-column.
    return [_section_break_para(doc, base_columns), *elements, _section_break_para(doc, 1)]


def _ensure_page_break_before(p_elem):
    """给段元素加 pageBreakBefore，强制该段从新页顶部开始。

    若 pPr 含 sectPr（分节符段），插在 sectPr 之前以保 OOXML schema 顺序
    （sectPr 必须是 pPr 的最后一个子元素）。
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    if pPr.find(qn('w:pageBreakBefore')) is None:
        pbb = OxmlElement('w:pageBreakBefore')
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            sectPr.addprevious(pbb)
        else:
            pPr.append(pbb)


def _prev_sibling_is_float(elem):
    """elem 的前一个非空兄弟是否为浮动体（图片段/caption/表格）。

    用于判断当前浮动体是否开启新组：前一个是图表→连续堆叠同页；
    前一个是正文/标题→新组，需 pageBreakBefore 页顶对齐。

    caption 段需结构性确认（前面紧跟 drawing/tbl）才视为浮动体，
    排除"Figure 4.1 shows…"这类叙事引用（编号+空格+动词）。
    """
    from docx.oxml.ns import qn

    prev = elem.getprevious()
    while prev is not None:
        if prev.tag == qn('w:tbl'):
            return True
        if prev.tag == qn('w:p'):
            if _is_empty_section_break_para(prev):
                prev = prev.getprevious()
                continue
            if prev.find(f'.//{qn("w:drawing")}') is not None:
                return True
            text = _paragraph_xml_text(prev).strip()
            if not text:
                prev = prev.getprevious()
                continue
            if _is_caption_like_text(text):
                return _caption_follows_float(prev)
            return False
        prev = prev.getprevious()
    return False


def _caption_follows_float(caption_elem):
    """caption 段前面（跳过空段/分节符段）是否紧跟 drawing/tbl。

    真 caption 紧跟图/表（连续浮动体）；叙事引用"Figure 4.1 shows…"前是正文
    → 非浮动体 → 当前浮动体应开新组加 pageBreakBefore。
    """
    from docx.oxml.ns import qn

    prev = caption_elem.getprevious()
    while prev is not None:
        if prev.tag == qn('w:tbl'):
            return True
        if prev.tag == qn('w:p'):
            if _is_empty_section_break_para(prev):
                prev = prev.getprevious()
                continue
            if prev.find(f'.//{qn("w:drawing")}') is not None:
                return True
            if _paragraph_xml_text(prev).strip():
                return False  # 非空正文 → caption 不属于浮动体
        prev = prev.getprevious()
    return False


def _pdf_guidance_allows_inline(pdf_guidance):
    if not isinstance(pdf_guidance, dict):
        return False
    position = str(pdf_guidance.get("position") or "").lower()
    if position in {"inline", "middle", "bottom"}:
        return True
    if position == "top":
        return False
    try:
        return float(pdf_guidance.get("y0_ratio")) > 0.22
    except (TypeError, ValueError):
        return False


def _pdf_guidance_delay_text(pdf_guidance):
    if not isinstance(pdf_guidance, dict):
        return ""
    delay_after = pdf_guidance.get("delay_after")
    if not isinstance(delay_after, dict):
        return ""
    return str(delay_after.get("text") or delay_after.get("query") or "").strip()


def _normalize_guidance_match_text(text):
    return re.sub(r"[\W_]+", "", text or "", flags=re.UNICODE).lower()


def _paragraph_matches_guidance_delay(p_elem, delay_text):
    needle = _normalize_guidance_match_text(delay_text)
    haystack = _normalize_guidance_match_text(_paragraph_xml_text(p_elem))
    if len(needle) < 12 or len(haystack) < 12:
        return False
    for size in (60, 45, 32, 24, 18, 12):
        if len(needle) >= size and needle[:size] in haystack:
            return True
    return len(haystack) >= 24 and haystack[:24] in needle


def _pdf_guidance_insert_anchor(anchor_elem, pdf_guidance):
    delay_text = _pdf_guidance_delay_text(pdf_guidance)
    if not delay_text:
        return None
    from docx.oxml.ns import qn

    candidate = anchor_elem.getnext()
    while candidate is not None:
        if candidate.tag == qn("w:tbl"):
            break
        if candidate.tag == qn("w:p"):
            text = _paragraph_xml_text(candidate).strip()
            if not text or _is_empty_section_break_para(candidate):
                candidate = candidate.getnext()
                continue
            if _is_heading_xml(candidate):
                if _pdf_guidance_is_top_float(pdf_guidance):
                    candidate = candidate.getnext()
                    continue
                break
            if _is_float_placeholder_text(text) or _is_caption_like_text(text):
                break
            if _paragraph_matches_guidance_delay(candidate, delay_text):
                return candidate
        candidate = candidate.getnext()
    return None


def _apply_float_page_top(insert_anchor, block_elems, pdf_guidance=None):
    """浮动体页顶对齐：新组第一个块加 pageBreakBefore 到新页顶。

    连续浮动体（前一个兄弟也是图表）不加 pageBreakBefore，由 keepWithNext
    自然堆叠同页；遇到正文/标题则视为新组，强制下一张图表从新页顶开始，
    让上一页被正文填满（消除留白）。
    """
    if not block_elems:
        return
    if _prev_sibling_is_float(insert_anchor):
        return
    if _pdf_guidance_allows_inline(pdf_guidance):
        return
    _ensure_page_break_before(block_elems[0])


def _front_matter_block_elements(doc, elements):
    """Keep title/abstract full width at document start."""
    elements = [elem for elem in elements if elem is not None]
    if _base_section_columns(doc) < 2 or not elements:
        return elements
    return [*elements, _section_break_para(doc, 1)]


def _ensure_body_section_starts_continuous(doc):
    """Let the two-column body continue on the title/abstract page."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sect_pr = doc.element.body.find(qn("w:sectPr"))
    if sect_pr is None:
        return
    sect_type = sect_pr.find(qn("w:type"))
    if sect_type is None:
        sect_type = OxmlElement("w:type")
        sect_pr.insert(0, sect_type)
    sect_type.set(qn("w:val"), "continuous")


def _paragraph_style_id(p_elem):
    from docx.oxml.ns import qn

    p_style = p_elem.find(f'{qn("w:pPr")}/{qn("w:pStyle")}')
    return (p_style.get(qn("w:val")) or "") if p_style is not None else ""


def _is_heading_xml(p_elem):
    style_id = _paragraph_style_id(p_elem).lower()
    if style_id.startswith("heading"):
        return True
    text = _paragraph_xml_text(p_elem).strip()
    return re.match(r"^\d+(?:\.\d+)*\s+\S+", text) is not None and len(text) <= 80


def _is_float_placeholder_text(text):
    return re.fullmatch(r'\s*\[(?:FIGURE|TIKZ_TABLE)_\d+\]\s*', text or "") is not None


def _is_caption_like_text(text):
    """判断文本是否像图表标题（caption）。

    严格匹配：编号后必须紧跟标点（. ： 、 等）或行尾，
    排除"Figure 4.1 shows..."/"图4.1展示了..."这类叙事性引用（编号后空格+动词）。
    """
    text = text or ""
    english = r'^\s*(?:Figure|Fig\.|Table)\s*\d+(?:\.\d+)?\s*(?:[.:：．;；]|$)'
    chinese = r'^\s*(?:图|表)\s*\d+(?:\.\d+)?\s*(?:[.．、:：；;]|$)'
    return (
        re.match(english, text, re.I) is not None
        or re.match(chinese, text) is not None
    )


def _paragraph_text_units(text):
    units = 0.0
    for ch in text or "":
        if '\u4e00' <= ch <= '\u9fff':
            units += 1.0
        elif ch.isspace():
            units += 0.25
        else:
            units += 0.55
    return units


def _float_page_height_mm(doc, style_spec):
    page_geo = (style_spec or {}).get("page_geometry") or {}
    if page_geo:
        height = page_geo.get("paper_height_mm")
        top = page_geo.get("top_margin_mm", 0)
        bottom = page_geo.get("bottom_margin_mm", 0)
        if height:
            return max(float(height) - float(top or 0) - float(bottom or 0), 120.0)
    section = doc.sections[0]
    page_height = getattr(section, "page_height", None)
    top_margin = getattr(section, "top_margin", None)
    bottom_margin = getattr(section, "bottom_margin", None)
    if page_height is not None and top_margin is not None and bottom_margin is not None:
        return max((int(page_height) - int(top_margin) - int(bottom_margin)) / 36000.0, 120.0)
    return 247.0


def _float_column_width_mm(doc, style_spec):
    page_geo = (style_spec or {}).get("page_geometry") or {}
    if page_geo.get("text_width_mm"):
        cols = max(int(page_geo.get("column_count") or 1), 1)
        sep = float(page_geo.get("column_sep_mm") or 0)
        return (float(page_geo["text_width_mm"]) - sep * (cols - 1)) / cols
    section = doc.sections[0]
    page_width = getattr(section, "page_width", None)
    left_margin = getattr(section, "left_margin", None)
    right_margin = getattr(section, "right_margin", None)
    if page_width is not None and left_margin is not None and right_margin is not None:
        return max((int(page_width) - int(left_margin) - int(right_margin)) / 36000.0, 40.0)
    return 82.0


def _estimate_paragraph_height_mm(p_elem, doc, style_spec):
    body_size = float((style_spec or {}).get("body_size") or 10)
    line_spacing = float((style_spec or {}).get("line_spacing") or 1.2)
    line_height = body_size * 0.3528 * line_spacing
    width_mm = max(_float_column_width_mm(doc, style_spec), 20.0)
    cjk_per_line = max(width_mm / max(body_size * 0.3528, 1.0), 8.0)
    text = _paragraph_xml_text(p_elem).strip()
    lines = max(1, math.ceil(_paragraph_text_units(text) / cjk_per_line))
    return lines * line_height


def _pdf_guidance_is_top_float(pdf_guidance):
    return (
        isinstance(pdf_guidance, dict)
        and str(pdf_guidance.get("position") or "").lower() == "top"
    )


def _float_delay_target_mm(doc, style_spec, required_space_mm=None, pdf_guidance=None):
    page_height = _float_page_height_mm(doc, style_spec)
    required = float(required_space_mm or page_height * 0.35)
    if _pdf_guidance_is_top_float(pdf_guidance):
        return min(page_height * 0.90, max(page_height * 0.55, required * 0.95))
    return min(page_height * 0.72, max(page_height * 0.34, required * 0.7))


def _graphic_option_length_mm(option_text, key, body_pt):
    match = re.search(
        rf'\b{re.escape(key)}\s*=\s*([0-9.]+)\s*(mm|cm|pt|in|em|ex)',
        option_text or "",
    )
    if not match:
        return None
    return _latex_vertical_space_pt(''.join(match.groups()), body_pt) / 2.83465


def _image_required_space_mm(img_info, image_path, width_emu, style_spec):
    body_pt = float((style_spec or {}).get("body_size") or 10)
    options = img_info.get("width") or ""
    height_mm = _graphic_option_length_mm(options, "height", body_pt)
    width_mm = width_emu / 36000.0 if width_emu else None
    if height_mm is None and width_mm:
        try:
            from docx.image.image import Image
            image = Image.from_file(str(image_path))
            if image.px_width:
                height_mm = width_mm * image.px_height / image.px_width
        except Exception:
            height_mm = None
    caption_space = 0.0
    if img_info.get("caption") or img_info.get("caption_full"):
        caption_space = body_pt * 0.3528 * 3
    return (height_mm or 80.0) + caption_space


def _table_required_space_mm(table_elem, style_spec):
    from docx.oxml.ns import qn

    body_pt = float((style_spec or {}).get("body_size") or 10)
    row_count = max(len(list(table_elem.iter(qn("w:tr")))), 1)
    row_height = body_pt * 0.3528 * 1.65
    return row_count * row_height + body_pt * 0.3528 * 3


def _full_width_table_anchor(anchor_elem, doc, table_elem, style_spec, pdf_guidance=None):
    return _full_width_float_insert_anchor(
        anchor_elem,
        doc=doc,
        style_spec=style_spec,
        required_space_mm=_table_required_space_mm(table_elem, style_spec),
        pdf_guidance=pdf_guidance,
    )


def _full_width_float_insert_anchor(anchor_elem, doc=None, style_spec=None,
                                    required_space_mm=None, pdf_guidance=None):
    """Delay a full-width float so preceding body can fill the prior page."""
    from docx.oxml.ns import qn

    if _pdf_guidance_allows_inline(pdf_guidance):
        return anchor_elem

    pdf_anchor = _pdf_guidance_insert_anchor(anchor_elem, pdf_guidance)
    if pdf_anchor is not None:
        return pdf_anchor

    legacy_one_paragraph = doc is None
    target_mm = 0.0 if legacy_one_paragraph else _float_delay_target_mm(
        doc, style_spec, required_space_mm, pdf_guidance=pdf_guidance)
    filled_mm = 0.0
    last_body = anchor_elem
    candidate = anchor_elem.getnext()
    while candidate is not None:
        if candidate.tag == qn("w:p"):
            text = _paragraph_xml_text(candidate).strip()
            if not text or _is_empty_section_break_para(candidate):
                candidate = candidate.getnext()
                continue
            if _is_heading_xml(candidate):
                if _pdf_guidance_is_top_float(pdf_guidance):
                    candidate = candidate.getnext()
                    continue
                break
            if _is_float_placeholder_text(text):
                break
            if _is_caption_like_text(text):
                break
            if legacy_one_paragraph:
                return candidate
            filled_mm += _estimate_paragraph_height_mm(candidate, doc, style_spec)
            last_body = candidate
            if filled_mm >= target_mm:
                return last_body
        if candidate.tag == qn("w:tbl"):
            break
        candidate = candidate.getnext()
    return last_body


def _is_body_start_paragraph(para):
    text = (para.text or '').strip()
    style_name = (para.style.name or '') if para.style else ''
    return (
        style_name.startswith('Heading')
        or re.match(r'^\d+(?:\.\d+)*\s+\S+', text) is not None
    )


def _ensure_front_matter_single_column(doc):
    """Keep title/author/abstract full width without splitting the body section."""
    from docx.oxml.ns import qn

    base_columns = _base_section_columns(doc)
    if base_columns < 2:
        return False
    target = next((para for para in doc.paragraphs if _is_body_start_paragraph(para)), None)
    if target is None:
        return False

    body = doc.element.body
    anchor = _heading_anchor_element(target._element)
    front_elems = []

    for elem in list(body):
        if elem is anchor:
            break
        if elem.tag == qn("w:p") and _is_empty_section_break_para(elem):
            body.remove(elem)
            continue
        if elem.tag == qn("w:tbl") and _is_front_matter_float_elem(elem):
            tc = elem.find(f'.//{qn("w:tc")}')
            if tc is not None:
                for child in list(tc):
                    if child.tag != qn("w:tcPr"):
                        tc.remove(child)
                        front_elems.append(child)
            body.remove(elem)
            continue
        if elem.tag not in (qn("w:p"), qn("w:tbl")):
            continue
        if elem.tag == qn("w:p") and not _paragraph_xml_text(elem).strip():
            body.remove(elem)
            continue
        body.remove(elem)
        front_elems.append(elem)

    if not front_elems:
        return False
    _insert_before_in_order(anchor, _front_matter_block_elements(doc, front_elems))
    _ensure_body_section_starts_continuous(doc)
    return True


def _insert_before_in_order(anchor_elem, elements):
    previous = None
    for elem in (elem for elem in elements if elem is not None):
        anchor_elem.addprevious(elem)
        previous = elem
    return previous


def _insert_after_in_order(anchor_elem, elements):
    current = anchor_elem
    for elem in (elem for elem in elements if elem is not None):
        current.addnext(elem)
        current = elem
    return current


def _insert_display_formulas(doc, display_formula_data, layout_spec=None):
    """将 [DISPLAY_FORMULA_N] 占位符替换为 OMML 公式

    Args:
        doc: python-docx Document 对象
        display_formula_data: list, 每个元素为 {latex, eq_num, env}
            ph_id 就是列表索引, 保证一一对应
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree

    def _append_zero_indent(pPr):
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:right'), '0')
        ind.set(qn('w:firstLine'), '0')
        pPr.append(ind)

    def _remove_blank_math_runs(omml_root):
        """递归清理OMML中空白/仅占位符的 math runs"""
        parent_map = {child: parent for parent in omml_root.iter() for child in parent}
        placeholder_chars = {'□', '■', '▪', '▫', '￼', '�', '​', '﻿'}
        changed = True
        while changed:
            changed = False
            parent_map = {child: parent for parent in omml_root.iter() for child in parent}
            for mr in list(omml_root.iter(qn('m:r'))):
                texts = [mt.text or '' for mt in mr.iter(qn('m:t'))]
                if not texts:
                    continue
                cleaned = ''.join(texts)
                for pc in placeholder_chars:
                    cleaned = cleaned.replace(pc, '')
                if not cleaned.strip():
                    parent = parent_map.get(mr)
                    if parent is not None:
                        parent.remove(mr)
                        changed = True

    if not display_formula_data:
        return

    # 收集所有含占位符的段落
    placeholder_map = {}  # ph_id -> paragraph
    for para in doc.paragraphs:
        m = re.search(r'\[DISPLAY_FORMULA_(\d+)\]', para.text)
        if m:
            ph_id = int(m.group(1))
            placeholder_map[ph_id] = para

    inserted = 0
    failed = 0

    for ph_id in sorted(placeholder_map.keys()):
        target_para = placeholder_map[ph_id]

        if ph_id >= len(display_formula_data):
            failed += 1
            continue

        f = display_formula_data[ph_id]
        latex_str = f['latex']
        eq_num = f.get('eq_num', '')

        # 调用 skill 的 latex_to_omml 转换
        omml_str = latex_to_omml(latex_str)
        if omml_str is None:
            failed += 1
            continue

        try:
            omml_elem = etree.fromstring(omml_str.encode('utf-8'))
        except Exception:
            failed += 1
            continue

        # 如果根元素是 m:oMath, 需要用 m:oMathPara 包裹 (display公式要求)
        omml_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        if omml_elem.tag == f'{{{omml_ns}}}oMath':
            omml_para = OxmlElement('m:oMathPara')
            omml_para.append(omml_elem)
            omml_block = omml_para
        else:
            omml_block = omml_elem

        # 创建新段落: OMML公式 + 制表符右对齐编号
        new_para = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')  # fleqn: 公式左对齐, 匹配LaTeX编译结果
        pPr.append(jc)
        _append_zero_indent(pPr)

        # 右对齐制表位(用于公式编号) — 位置等于页面文本宽度
        # 从style_spec的page_geometry动态获取textwidth_mm
        tab_right_pos = '10080'  # 默认值(177mm)
        tab_right_mm = 177.0  # 默认文本宽度mm
        if layout_spec:
            pg = layout_spec.get('page_geometry') or {}
            tab_right_mm = pg.get('textwidth_mm', 177.0)
            tab_right_pos = str(int(tab_right_mm * 56.7))  # mm → twips
        if eq_num:
            tabs = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            page_geometry = (layout_spec or {}).get('page_geometry') or {}
            compiled_layout = (layout_spec or {}).get('compiled_layout') or {}
            column_width = compiled_layout.get('columnwidth', 0)
            column_width = column_width or _page_content_width_twips(
                page_geometry, single_column=True)
            column_width = column_width or _section_content_width_twips(
                doc.sections[0], single_column=True)
            if column_width:
                tab.set(qn('w:pos'), str(column_width))
                tabs.append(tab)
                pPr.append(tabs)

        new_para.append(pPr)

        # 插入 OMML 元素
        new_para.append(omml_block)

        # 插入公式编号 (制表符+编号，右对齐)
        if eq_num:
            tab_run = OxmlElement('w:r')
            tab_elem = OxmlElement('w:tab')
            tab_run.append(tab_elem)
            new_para.append(tab_run)
            num_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '20')
            rPr.append(szCs)
            num_run.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = eq_num
            num_run.append(t)
            new_para.append(num_run)

        # 替换原占位符段落
        # 如果段落包含正文文本（占位符与文本合并在同一段落），
        # 需要拆分：占位符前文本 → 公式 → 占位符后文本
        target_elem = target_para._element
        para_text = target_para.text or ''
        ph_match = re.search(r'\[DISPLAY_FORMULA_\d+\]', para_text)
        if not ph_match:
            failed += 1
            continue
        before_text = para_text[:ph_match.start()].strip()
        after_text = para_text[ph_match.end():].strip()

        if after_text:
            # 创建后续文本段落（保留原段落的格式）
            after_para = OxmlElement('w:p')
            # 复制原段落的pPr（格式属性）
            orig_pPr = target_elem.find(qn('w:pPr'))
            if orig_pPr is not None:
                after_para.append(copy.deepcopy(orig_pPr))
            after_run = OxmlElement('w:r')
            after_t = OxmlElement('w:t')
            after_t.set(qn('xml:space'), 'preserve')
            after_t.text = after_text
            after_run.append(after_t)
            after_para.append(after_run)
            target_elem.addnext(after_para)

        # 插入公式段落
        target_elem.addnext(new_para)

        if before_text:
            # 创建前置文本段落
            before_para = OxmlElement('w:p')
            orig_pPr = target_elem.find(qn('w:pPr'))
            if orig_pPr is not None:
                before_para.append(copy.deepcopy(orig_pPr))
            before_run = OxmlElement('w:r')
            before_t = OxmlElement('w:t')
            before_t.set(qn('xml:space'), 'preserve')
            before_t.text = before_text
            before_run.append(before_t)
            before_para.append(before_run)
            target_elem.addnext(before_para)

        # 删除原占位符段落
        target_elem.getparent().remove(target_elem)
        inserted += 1

    # 清理公式段落中的空白占位符run
    for para in doc.paragraphs:
        for omml in para._element.iter(qn('m:oMath')):
            _remove_blank_math_runs(omml)

    print(f'  [公式插入] 成功: {inserted}, 失败: {failed}, 总计: {len(display_formula_data)}')


def _subscript_to_unicode(text):
    """Convert subscript text to Unicode subscript characters."""
    sub_map = str.maketrans('0123456789+-=()aeghijklmnoprstuvx',
                             '₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ')
    return text.translate(sub_map)


def _clean_formula_placeholders(doc):
    """清理OMML公式中的占位符残留

    Pandoc --mathml 转换时，某些LaTeX下标/上标被转为占位符形式：
    - 〖XCO〗_2 (Unicode数学方括号 U+3016/U+3017)
    - 【XCO】_2 (全角方括号 U+3010/U+3011)
    - □/■/▪ 空白占位方块
    - FFFC/FFFD 对象替换字符
    需要清理掉这些包裹，保留内部文本。
    """
    from docx.oxml.ns import qn
    # 需要清理的占位符对（左括号, 右括号）
    placeholder_pairs = [
        ('〖', '〗'),  # 〖〗 数学左/右白色方括号
        ('【', '】'),  # 【】 左/右黑方括号
        ('︵', '︶'),  # ︹︺ 数学方括号变体
        ('┇', '┈'),  # ﹇﹈ 小方括号变体
    ]
    all_left = set(p[0] for p in placeholder_pairs)
    all_right = set(p[1] for p in placeholder_pairs)
    bracket_chars = all_left | all_right
    # 独立清理字符（不成对）
    standalone_chars = {
        '□',  # □ 白色方块
        '■',  # ■ 黑色方块
        '▪',  # ▪ 小黑色方块
        '▫',  # ▫ 小白色方块
        '￼',  # OBJ 对象替换字符
        '�',  #  替换字符
        '​',  # 零宽空格(ZWSP) — Pandoc OMML常见残留
        '‌',  # 零宽非连接符
        '﻿',  # BOM
    }
    all_chars = bracket_chars | standalone_chars

    # \textsubscript 残留模式: CO\textsubscript{2} → CO₂
    textsub_pattern = re.compile(r'\\textsubscript\{([^}]+)\}')

    cleaned = 0
    for para in doc.paragraphs:
        # 清理OMML公式中的占位符
        for omml in para._element.iter(qn('m:oMath')):
            for mt in list(omml.iter(qn('m:t'))):
                if mt.text and any(c in mt.text for c in all_chars):
                    for left, right in placeholder_pairs:
                        mt.text = mt.text.replace(left, '').replace(right, '')
                    for sc in standalone_chars:
                        mt.text = mt.text.replace(sc, '')
                    cleaned += 1
                # 清理 \textsubscript 残留
                if mt.text and '\\textsubscript' in mt.text:
                    mt.text = textsub_pattern.sub(lambda m: _subscript_to_unicode(m.group(1)), mt.text)
                    cleaned += 1
        # 清理段落级别run文本中的占位符
        for run in para.runs:
            if any(c in (run.text or '') for c in all_chars):
                for left, right in placeholder_pairs:
                    run.text = run.text.replace(left, '').replace(right, '')
                for sc in standalone_chars:
                    run.text = run.text.replace(sc, '')
                cleaned += 1
            # 清理 \textsubscript 残留
            if run.text and '\\textsubscript' in run.text:
                run.text = textsub_pattern.sub(lambda m: _subscript_to_unicode(m.group(1)), run.text)
                cleaned += 1
    if cleaned:
        print(f'  [公式清理] 清理占位符: {cleaned}处')


def _prune_empty_math_placeholders(doc):
    """Thoroughly prune empty OMML placeholder slots from math elements.

    Removes empty m:r, m:nary (with empty m:e), and other structural
    containers that display as small boxes in Word.
    Also cleans empty m:sub/m:sup inside m:nary (e.g. \\sum with empty sub/sup).
    """
    from docx.oxml.ns import qn

    removable_tags = {
        qn("m:r"), qn("m:nary"), qn("m:d"), qn("m:sSub"), qn("m:sSup"),
        qn("m:sSubSup"), qn("m:acc"), qn("m:bar"), qn("m:box"),
        qn("m:borderBox"), qn("m:func"), qn("m:groupChr"),
        qn("m:limLow"), qn("m:limUpp"),
    }

    # 清理nary结构中的空下标/上标（∑后面的空白占位符）
    nary_sub_tags = {qn("m:sub"), qn("m:sup")}
    nary_cleaned = 0
    for para in doc.paragraphs:
        for omml in para._element.iter(qn('m:oMath')):
            for nary in list(omml.iter(qn('m:nary'))):
                for sub_tag in nary_sub_tags:
                    sub_elem = nary.find(sub_tag)
                    if sub_elem is not None:
                        # 检查是否为空（无可见文本）
                        has_text = False
                        for mt in sub_elem.iter(qn('m:t')):
                            if mt.text and mt.text.strip():
                                has_text = True
                                break
                        if not has_text:
                            # 空的sub/sup：移除其内容但保留元素（OMML要求结构完整）
                            for child in list(sub_elem):
                                sub_elem.remove(child)
                            nary_cleaned += 1
    if nary_cleaned:
        print(f"  [formula cleanup] cleaned empty nary sub/sup: {nary_cleaned}")

    def has_visible_math_text(elem):
        for mt in elem.iter(qn("m:t")):
            if mt.text and mt.text.strip():
                return True
        return False

    removed = 0
    for para in doc.paragraphs:
        for omml in para._element.iter(qn("m:oMath")):
            parent_map = {child: parent for parent in omml.iter() for child in parent}
            for elem in list(omml.iter())[::-1]:
                if elem.tag not in removable_tags:
                    continue
                if has_visible_math_text(elem):
                    continue
                if len(elem) and any(has_visible_math_text(c) for c in elem):
                    continue
                parent = parent_map.get(elem)
                if parent is not None:
                    parent.remove(elem)
                    removed += 1
    if removed:
        print(f"  [formula cleanup] removed empty OMML placeholder slot(s): {removed}")


def _fix_overline_acc_to_bar(doc):
    r"""将OMML中错误的 m:acc(overline) 转为正确的 m:bar

    latex2mathml 把 \overline{} 转为 MathML mover + 水平线(U+2015)，
    MML2OMML.XSL 再映射为 m:acc(chr=―)，但 Word 中正确显示上划线
    应该用 m:bar(pos=top) 结构。

    修复: m:acc + chr=U+2015 → m:bar(pos=top)
    """
    from docx.oxml.ns import qn
    from lxml import etree

    # U+2015 = ― (Horizontal Bar / Quotation Dash)
    OVERLINE_CHAR = '―'
    fixed = 0

    for para in doc.paragraphs:
        for omml in para._element.iter(qn('m:oMath')):
            for acc in list(omml.iter(qn('m:acc'))):
                accPr = acc.find(qn('m:accPr'))
                if accPr is None:
                    continue
                chr_elem = accPr.find(qn('m:chr'))
                if chr_elem is None:
                    continue
                chr_val = chr_elem.get(qn('m:val'), '')
                if chr_val == OVERLINE_CHAR:
                    # 将 m:acc 转为 m:bar
                    # 提取子元素
                    children = list(acc)
                    acc_parent = acc.getparent()
                    if acc_parent is None:
                        continue

                    # 构建 m:bar
                    bar = etree.SubElement(acc_parent, qn('m:bar'))
                    # 移动位置到原acc之后
                    acc_parent.remove(bar)
                    idx = list(acc_parent).index(acc)
                    acc_parent.insert(idx, bar)

                    # m:barPr: pos=top
                    barPr = etree.SubElement(bar, qn('m:barPr'))
                    pos = etree.SubElement(barPr, qn('m:pos'))
                    pos.set(qn('m:val'), 'top')
                    ctrlPr = etree.SubElement(barPr, qn('m:ctrlPr'))

                    # m:e (内容) — 从acc中提取
                    acc_e = acc.find(qn('m:e'))
                    if acc_e is not None:
                        bar.append(acc_e)

                    # 移除原acc
                    acc_parent.remove(acc)
                    fixed += 1

    if fixed:
        print(f"  [formula fix] converted {fixed} m:acc(overline) → m:bar")


def _fix_empty_nary_body(doc):
    r"""修复OMML中 nary(∑等) 结构缺失 m:e (求和体) 的问题

    latex2mathml 将 \sum_{h} P_h a_h 转换时，把 \sum_{h} 生成为一个
    空的 m:nary (只有 m:sub/m:sup，没有 m:e)，而把 P_h a_h 放在 nary 外面
    作为兄弟元素。Word 显示空的 nary 为小方块占位符。

    修复：将空的 nary 后面的兄弟元素移入 nary 的 m:e 中。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    fixed = 0

    for para in doc.paragraphs:
        for omml in para._element.iter(qn('m:oMath')):
            for nary in list(omml.iter(qn('m:nary'))):
                # 检查 nary 是否已有非空的 m:e
                e_elem = nary.find(qn('m:e'))
                if e_elem is not None:
                    # m:e 存在，检查是否有内容
                    has_content = False
                    for child in e_elem:
                        has_content = True
                        break
                    if has_content:
                        continue  # m:e 已有内容，跳过

                # 空的 nary — 需要把后面的兄弟元素移入 m:e
                # 找到 nary 在父元素中的位置
                parent = nary.getparent()
                if parent is None:
                    continue

                parent_children = list(parent)
                nary_idx = parent_children.index(nary)

                # 收集 nary 后面的兄弟元素（直到遇到结构边界）
                # 边界条件：遇到另一个 nary/omathpara，或到末尾
                stop_tags = {qn('m:nary'), qn('m:oMathPara')}
                siblings_to_move = []
                for si in range(nary_idx + 1, len(parent_children)):
                    sibling = parent_children[si]
                    if sibling.tag in stop_tags:
                        break
                    siblings_to_move.append(sibling)

                if not siblings_to_move:
                    continue

                # 创建 m:e 并移入兄弟元素
                if e_elem is None:
                    e_elem = etree.SubElement(nary, qn('m:e'))
                else:
                    # 清理空的 m:e
                    for child in list(e_elem):
                        e_elem.remove(child)

                for sib in siblings_to_move:
                    parent.remove(sib)
                    e_elem.append(sib)

                fixed += 1

    if fixed:
        print(f"  [formula fix] filled {fixed} empty nary m:e with sibling elements")


def _clean_leading_table_bullets(doc):
    """Remove bullet/numbering markers (numPr) from table cell paragraphs.

    Prevents black squares or bullet symbols appearing next to table content.
    """
    from docx.oxml.ns import qn
    cleaned = 0
    for tbl in doc.tables:
        if _is_layout_floating_table_elem(tbl._element):
            continue
        for row_idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    pPr = para._element.find(qn("w:pPr"))
                    if pPr is not None:
                        numPr = pPr.find(qn("w:numPr"))
                        if numPr is not None:
                            pPr.remove(numPr)
                            cleaned += 1
    if cleaned:
        print(f"  [table cleanup] removed {cleaned} numPr markers from table cells")


def _deep_clean_table_cells(doc):
    """Thoroughly clean table cell paragraphs to prevent black squares.

    Removes: numPr (list markers), orphan indents, list-style references,
    and any numbering-related XML that Word may interpret as bullet markers.
    Also cleans style-level numbering inheritance.
    """
    from docx.oxml.ns import qn
    cleaned = 0
    for tbl in doc.tables:
        if _is_layout_floating_table_elem(tbl._element):
            continue
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    elem = para._element
                    pPr = elem.find(qn("w:pPr"))
                    if pPr is None:
                        # 即使没有pPr，也检查run中是否有隐藏的bullet符号
                        continue
                    # 1. 移除 numPr（列表标记 → 黑色方块）
                    numPr = pPr.find(qn("w:numPr"))
                    if numPr is not None:
                        pPr.remove(numPr)
                        cleaned += 1
                    # 2. 移除列表式缩进（hanging > 0）
                    ind = pPr.find(qn("w:ind"))
                    if ind is not None:
                        hanging = ind.get(qn("w:hanging"))
                        if hanging and int(hanging) > 0:
                            pPr.remove(ind)
                            cleaned += 1
                    # 3. 清除样式中的列表引用（pStyle指向ListParagraph等）
                    pStyle = pPr.find(qn("w:pStyle"))
                    if pStyle is not None:
                        style_val = pStyle.get(qn("w:val"), "")
                        if "list" in style_val.lower() or "List" in style_val:
                            pPr.remove(pStyle)
                            cleaned += 1

                # 4. 检查每个cell的第一个run，移除开头的bullet字符
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            stripped = run.text.lstrip()
                            # bullet字符：•, (U+00B7), ■, ●, -, *, 以及零宽字符
                            if stripped and stripped[0] in '•·■●○-*':
                                run.text = stripped[1:].lstrip()
                                cleaned += 1
                            break  # 只检查第一个run
    if cleaned:
        print(f"  [table deep-clean] removed {cleaned} markers from table cells")


def _purge_numbering_xml(doc):
    """Remove all numbering definitions from numbering.xml to prevent black squares.

    Also removes TableGrid style from tables to prevent left-border rendering
    that looks like black squares on each row.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return
        numbering_xml = numbering_part._element
        removed = 0
        for tag in ('w:abstractNum', 'w:num'):
            for elem in list(numbering_xml.findall(qn(tag))):
                numbering_xml.remove(elem)
                removed += 1
        if removed:
            print(f"  [numbering purge] removed {removed} numbering definitions to prevent black squares")
    except Exception:
        pass

    # 移除 TableGrid 样式引用 + tblLook + 清零 tblCellMar
    # 三层清理彻底消除表格左侧黑色方块:
    # 1) tblStyle: TableGrid 自带全边框
    # 2) tblLook: firstColumn/firstRow 条件格式触发默认样式渲染
    # 3) tblCellMar: 默认"Table"样式有 left=108(~1.9mm) 内边距，
    #    Word 在该空间内渲染段落标记符，表现为黑色小方块
    for tbl in doc.tables:
        if _is_layout_floating_table_elem(tbl._element):
            _clear_table_borders(tbl._element)
            continue
        tbl_pr = tbl._element.find(qn('w:tblPr'))
        if tbl_pr is not None:
            tbl_style = tbl_pr.find(qn('w:tblStyle'))
            if tbl_style is not None:
                tbl_pr.remove(tbl_style)
            tbl_look = tbl_pr.find(qn('w:tblLook'))
            if tbl_look is not None:
                tbl_pr.remove(tbl_look)
            # 清零单元格内边距，覆盖默认"Table"样式的 tblCellMar left=108
            cell_mar = tbl_pr.find(qn('w:tblCellMar'))
            if cell_mar is None:
                cell_mar = OxmlElement('w:tblCellMar')
                tbl_pr.append(cell_mar)
            for edge in ('top', 'left', 'bottom', 'right', 'start', 'end'):
                e = cell_mar.find(qn(f'w:{edge}'))
                if e is not None:
                    cell_mar.remove(e)
            for edge in ('top', 'start', 'bottom', 'end'):
                e = OxmlElement(f'w:{edge}')
                e.set(qn('w:w'), '0')
                e.set(qn('w:type'), 'dxa')
                cell_mar.append(e)

            _complete_table_bottom_line(tbl)
    _clear_layout_floating_table_borders(doc)


def _remove_blank_paragraphs_around_tables(doc):
    """清理表格前后的多余空白段落，压缩间距

    LaTeX中表格是浮动体(floating)，文字可以跨越表格连续排版。
    但Word中表格是固定块级元素，Pandoc会在表格前后生成多个空段落，
    导致文字被强行断开，出现大片空白。

    修复策略:
    1. 删除表格前后所有空段落（LaTeX浮动体不需要额外间距）
    2. 压缩表格前后有内容段落的间距为0（模拟LaTeX连续排版）
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    removed = 0
    body = doc.element.body
    tbl_elements = list(body.findall(qn('w:tbl')))

    for tbl_elem in tbl_elements:
        # === 清理表格前面的空段落 ===
        prev = tbl_elem.getprevious()
        blank_before = []
        while prev is not None and prev.tag == qn('w:p'):
            if _is_empty_section_break_para(prev):
                break
            text = ''
            for t_elem in prev.iter(qn('w:t')):
                text += (t_elem.text or '')
            if text.strip():
                break  # 遇到有内容的段落，停止
            blank_before.append(prev)
            prev = prev.getprevious()

        # 删除所有前面的空段落（LaTeX浮动体不需要间距）
        for p in blank_before:
            body.remove(p)
            removed += 1

        # 压缩表格前有内容段落的间距为0
        if prev is not None and prev.tag == qn('w:p'):
            pPr = prev.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                prev.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:after'), '0')
            for attr in (qn('w:line'), qn('w:lineRule')):
                if spacing.get(attr) is not None:
                    del spacing.attrib[attr]

        # === 清理表格后面的空段落 ===
        nxt = tbl_elem.getnext()
        blank_after = []
        while nxt is not None and nxt.tag == qn('w:p'):
            if _is_empty_section_break_para(nxt):
                break
            text = ''
            for t_elem in nxt.iter(qn('w:t')):
                text += (t_elem.text or '')
            if text.strip():
                break
            blank_after.append(nxt)
            nxt = nxt.getnext()

        # 删除所有后面的空段落
        for p in blank_after:
            body.remove(p)
            removed += 1

        # 压缩表格后有内容段落的间距为0
        if nxt is not None and nxt.tag == qn('w:p'):
            pPr = nxt.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                nxt.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            for attr in (qn('w:line'), qn('w:lineRule')):
                if spacing.get(attr) is not None:
                    del spacing.attrib[attr]

    if removed:
        print(f"  [table cleanup] removed {removed} blank paragraphs, compressed spacing around tables")


def _remove_blank_paragraphs_around_images(doc):
    """Remove body-level blank paragraphs around inserted image paragraphs."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    removed = 0
    body = doc.element.body
    image_paras = [
        p for p in body.findall(qn("w:p"))
        if p.findall(".//" + qn("w:drawing"))
    ]

    for img_p in image_paras:
        pPr = img_p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            img_p.insert(0, pPr)
        _append_zero_paragraph_spacing(pPr)
        _clear_image_para_indent_markers(pPr)  # 清首行缩进+numPr,防图片左侧黑方块

        for direction in ("previous", "next"):
            neighbor = img_p.getprevious() if direction == "previous" else img_p.getnext()
            blanks = []
            while neighbor is not None and neighbor.tag == qn("w:p"):
                if _is_empty_section_break_para(neighbor) or _paragraph_xml_text(neighbor).strip():
                    break
                blanks.append(neighbor)
                neighbor = neighbor.getprevious() if direction == "previous" else neighbor.getnext()
            for blank in blanks:
                body.remove(blank)
                removed += 1

    if removed:
        print(f"  [image cleanup] removed {removed} blank paragraphs around images")


def _fix_table_continuation_indent(doc):
    """修复被表格截断的续接段落的缩进

    当表格前后的文字应该是同一段落时（前段不以句末标点结尾），
    表格后的段落不应该有首行缩进（它是前段的续接）。

    LaTeX浮动体特性：表格在PDF中浮动，前后的文字是连续的一段。
    但Word中表格是固定块级元素，会打断段落。如果前段确实未结束
    （不以 。.!?！？ 结尾），则后段应视为续接，移除首行缩进。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    sentence_end = set('.!?。！？')
    fixed = 0
    body = doc.element.body
    tbl_elements = list(body.iter(qn('w:tbl')))

    def _normalize_following_body_paragraph(p_elem):
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            for attr in (qn('w:line'), qn('w:lineRule')):
                if spacing.get(attr) is not None:
                    del spacing.attrib[attr]
        for tag in ('w:keepNext', 'w:keepLines', 'w:pageBreakBefore'):
            elem = pPr.find(qn(tag))
            if elem is not None:
                pPr.remove(elem)
        return pPr

    def _starts_float_reference(text):
        return bool(re.match(
            r'^\s*(?:表|图)\s*\d+(?:\.\d+)?|^\s*(?:Table|Figure|Fig\.?)\s*\d+(?:\.\d+)?\b',
            text,
            flags=re.IGNORECASE,
        ))

    for tbl_elem in tbl_elements:
        # 向前找最近的有内容段落
        prev = tbl_elem.getprevious()
        while prev is not None and prev.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in prev.iter(qn('w:t')))
            if text.strip():
                break
            prev = prev.getprevious()

        # 向后找最近的有内容段落
        nxt = tbl_elem.getnext()
        while nxt is not None and nxt.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in nxt.iter(qn('w:t')))
            if text.strip():
                break
            nxt = nxt.getnext()

        if (prev is None or nxt is None
                or prev.tag != qn('w:p') or nxt.tag != qn('w:p')):
            continue

        prev_text = ''.join(t.text or '' for t in prev.iter(qn('w:t'))).rstrip()
        next_text = ''.join(t.text or '' for t in nxt.iter(qn('w:t'))).lstrip()

        if not prev_text or not next_text:
            continue

        prev_last = prev_text[-1]
        next_first = next_text[0]
        prev_is_caption = _starts_float_reference(prev_text)
        next_is_reference = _starts_float_reference(next_text)

        # 前段不以句末标点结尾 → 未结束
        # 后段不以大写字母/数字开头 → 非新段落
        should_continue = (
            prev_last not in sentence_end
            and not next_first.isupper()
            and not next_first.isdigit()
            and not prev_is_caption
            and not next_is_reference
        )
        pPr = _normalize_following_body_paragraph(nxt)

        if should_continue:
            # 移除续接段落的首行缩进
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)

            ind.set(qn('w:firstLine'), '0')
            ind.set(qn('w:firstLineChars'), '0')
            fixed += 1

    if fixed:
        print(f'  [续接修复] 修复了 {fixed} 处表格后续接段落缩进')


def _set_cell_border(cell, **kwargs):
    """设置单元格边框, kwargs: top/bottom/left/start/right/end=(val, sz, color)"""
    _set_physical_tc_border(cell._element, **kwargs)


def _set_physical_tc_border(tc, **kwargs):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        _append_tc_pr_child(tcPr, tcBorders)
    for edge, (val, sz, color) in kwargs.items():
        elem = tcBorders.find(qn(f'w:{edge}'))
        if elem is None:
            elem = OxmlElement(f'w:{edge}')
            tcBorders.append(elem)
        elem.set(qn('w:val'), val)
        elem.set(qn('w:sz'), str(sz))
        elem.set(qn('w:color'), color)
        elem.set(qn('w:space'), '0')


def _add_table_borders(tbl, tex_table=None, layout_spec=None):
    """给Word表格添加边框线, 匹配LaTeX的tophline/middlehline/bottomhline/|

    layout_spec: 排版规格字典, 含 table.no_vertical_rules, table.rule_style 等
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = tbl.rows
    if not rows:
        return

    tblPr = tbl._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tblPr)

    # 从layout_spec获取表格规则
    table_spec = layout_spec.get('table', {}) if layout_spec else {}
    no_vertical = table_spec.get('no_vertical_rules', False)
    rule_style = table_spec.get('rule_style', 'default')

    # 解析LaTeX列格式中的竖线位置
    vline_cols = set()
    if not no_vertical and tex_table and 'col_format' in tex_table:
        col_fmt = tex_table['col_format']
        col_idx = 0
        for ch in col_fmt:
            if ch == '|':
                vline_cols.add(col_idx)
            elif ch in 'lcr':
                col_idx += 1

    total_rows = len(rows)
    for ri, row in enumerate(rows):
        is_first = (ri == 0)
        is_last = (ri == total_rows - 1)

        # 线宽映射: 根据rule_style决定
        if rule_style in ('template_hlines', 'booktabs'):
            # 3线表: tophline/bottomhline粗线, middlehline细线
            top_val = 'single'
            top_sz = 8 if is_first else 4      # tophline=8, internal=4
            bottom_val = 'single'
            bottom_sz = 8 if is_last else 4    # bottomhline=8, internal=4
            if ri == 1 and total_rows > 2:
                top_sz = 4  # middlehline
        else:
            # default: 粗线顶/底, 细线内
            top_val = 'single'
            top_sz = 12 if is_first else 4
            bottom_val = 'single'
            bottom_sz = 12 if is_last else 4
            if ri == 1 and total_rows > 2:
                top_sz = 8  # middlehline

        for ci, cell in enumerate(row.cells):
            border_kwargs = {
                'top': (top_val, top_sz, '000000'),
                'bottom': (bottom_val, bottom_sz, '000000'),
            }
            # 竖线控制: no_vertical_rules 时全部nil
            if no_vertical:
                border_kwargs['left'] = ('nil', '0', 'auto')
                border_kwargs['right'] = ('nil', '0', 'auto')
            else:
                if ci == 0:
                    border_kwargs['left'] = ('single', 4, '000000')
                if ci == len(row.cells) - 1:
                    border_kwargs['right'] = ('single', 4, '000000')
                if ci in vline_cols and ci > 0:
                    border_kwargs['left'] = ('single', 4, '000000')

            _set_cell_border(cell, **border_kwargs)


def _set_font(run, cn_font=None, en_font=None, size=None):
    """设置run的中英文字体"""
    if cn_font or en_font:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)

        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        if en_font:
            rFonts.set(qn('w:ascii'), en_font)
            rFonts.set(qn('w:hAnsi'), en_font)
            rFonts.set(qn('w:cs'), en_font)
        if cn_font:
            rFonts.set(qn('w:eastAsia'), cn_font)

    if size:
        from docx.shared import Pt
        run.font.size = Pt(size)


def _apply_fonts_to_doc(doc, cn_font=None, en_font=None):
    """给整个文档设置中英文字体"""
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    if en_font:
        style.font.name = en_font
    if cn_font:
        rPr = style.element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                rFonts.set(qn('w:eastAsia'), cn_font)

    for para in doc.paragraphs:
        for run in para.runs:
            _set_font(run, cn_font, en_font)

    for tbl in doc.tables:
        for row_idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_font(run, cn_font, en_font)


_LATEX_SIZE_MACROS = {
    "tiny": 5,
    "scriptsize": 7,
    "footnotesize": 8,
    "small": 9,
    "normalsize": 10,
    "large": 12,
    "Large": 14.4,
    "LARGE": 17.28,
    "huge": 20.74,
    "Huge": 24.88,
    "@vpt": 5,
    "@vipt": 6,
    "@viipt": 7,
    "@viiipt": 8,
    "@ixpt": 9,
    "@xpt": 10,
    "@xipt": 11,
    "@xiipt": 12,
    "@xivpt": 14.4,
    "@xviipt": 17.28,
    "@xxpt": 20.74,
    "@xxvpt": 24.88,
}


def _latex_size_token_to_pt(token, default):
    if not token:
        return default
    token = token.strip().strip("{}")
    if token.startswith("\\"):
        token = token[1:]
    if token in _LATEX_SIZE_MACROS:
        return _LATEX_SIZE_MACROS[token]
    if f"@{token}" in _LATEX_SIZE_MACROS:
        return _LATEX_SIZE_MACROS[f"@{token}"]
    try:
        return float(token)
    except ValueError:
        return default


def _latex_length_to_twips(value, body_pt=10):
    if value in (None, ""):
        return None
    match = re.search(r"([-+]?\d+(?:\.\d+)?)\s*(pt|bp|mm|cm|in|em|ex)?", str(value))
    if not match:
        return None
    amount = float(match.group(1))
    unit = match.group(2) or "pt"
    pt = {
        "pt": amount,
        "bp": amount * 72 / 72.27,
        "mm": amount * 72 / 25.4,
        "cm": amount * 72 / 2.54,
        "in": amount * 72,
        "em": amount * body_pt,
        "ex": amount * body_pt * 0.43,
    }.get(unit)
    return int(round(pt * 20)) if pt is not None else None


def _extract_latex_font_size(cls_content, size_name, default_size, default_baseline):
    if not cls_content:
        return default_size, default_baseline
    anchors = [
        f"\\renewcommand\\{size_name}",
        f"\\def\\{size_name}",
        f"\\{size_name}",
    ]
    block = cls_content
    for anchor in anchors:
        idx = cls_content.find(anchor)
        if idx >= 0:
            block = cls_content[idx:idx + 1200]
            break

    escaped = re.escape(size_name)
    patterns = [
        rf"\\@setfontsize\\{escaped}\\(?P<size>@?[A-Za-z]+|\d+(?:\.\d+)?)\s*\{{(?P<lead>\\?@?[A-Za-z]+|\d+(?:\.\d+)?)\}}",
        rf"\\@setfontsize\\{escaped}\\(?P<size>@?[A-Za-z]+|\d+(?:\.\d+)?)\\(?P<lead>@?[A-Za-z]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, block)
        if match:
            size = _latex_size_token_to_pt(match.group("size"), default_size)
            lead = _latex_size_token_to_pt(match.group("lead"), default_baseline)
            return size, lead
    return default_size, default_baseline


def _extract_first_tex_command_value(tex_content, command):
    match = re.search(rf"\\{command}(?:\[[^\]]*\])?\{{([^}}]+)\}}", tex_content)
    return match.group(1).strip() if match else None


def _extract_tex_dimension_mm(cls_content, cmd_name, default_mm=None):
    """Extract a LaTeX dimension from cls content and return mm.

    Handles simple values like '177mm', '16.4mm', and dimexpr like
    '\\dimexpr210mm+2\\bleed\\relax'. Returns ALL found values
    (multiple values may exist for different config branches).
    """
    conv = {'mm': 1.0, 'cm': 10.0, 'pt': 0.3528, 'in': 25.4}
    candidates = []

    # Find ALL simple dimension values (e.g. \\textwidth177mm)
    for m in re.finditer(rf"\\{cmd_name}\s*([\d.]+)\s*(mm|cm|pt|in)", cls_content):
        value = float(m.group(1))
        unit = m.group(2)
        candidates.append(round(value * conv.get(unit, 1.0), 1))

    # Find ALL dimexpr values (e.g. \\paperheight\\dimexpr277mm+2\\bleed\\relax)
    # Note: there may be no space between \\cmd_name and \\dimexpr
    for m in re.finditer(
        rf"\\{cmd_name}\s*\\dimexpr\s*([\d.]+)\s*(mm|cm|pt|in)\s*([+-])\s*([\d.]+)\s*\\\\(\\w+)",
        cls_content
    ):
        base_val = float(m.group(1))
        unit = m.group(2)
        sign = m.group(3)
        mult = float(m.group(4))
        bleed_cmd = m.group(5)
        base_mm = base_val * conv.get(unit, 1.0)

        bleed_val = re.search(rf"\\\\def\\\\{bleed_cmd}\s*([\d.]+)\s*(mm|cm|pt|in)", cls_content)
        if bleed_val:
            bleed_mm = float(bleed_val.group(1)) * conv.get(bleed_val.group(2), 1.0)
        else:
            bleed_mm = 3.0  # default bleed

        if sign == '+':
            candidates.append(round(base_mm + bleed_mm * mult, 1))
        else:
            candidates.append(round(base_mm - bleed_mm * mult, 1))

    return candidates


def _extract_page_geometry(cls_content):
    """Extract page geometry from a LaTeX .cls file, with config awareness.

    Many templates have multiple paper sizes for different modes
    (manuscript vs final, preprint vs published). This extracts ALL found
    values and returns the recommended (largest) set for final output.

    Returns dict with keys: paperwidth_mm, paperheight_mm, textwidth_mm,
    textheight_mm, oddsidemargin_mm, topmargin_mm.
    """
    if not cls_content:
        return None

    paper_w_candidates = _extract_tex_dimension_mm(cls_content, 'paperwidth')
    paper_h_candidates = _extract_tex_dimension_mm(cls_content, 'paperheight')
    text_w_candidates = _extract_tex_dimension_mm(cls_content, 'textwidth')
    oddsidemargin_candidates = _extract_tex_dimension_mm(cls_content, 'oddsidemargin')
    topmargin_candidates = _extract_tex_dimension_mm(cls_content, 'topmargin', 10)

    if not text_w_candidates:
        return None

    # Strategy: use the LARGEST paper dimensions (final/published version)
    # and the largest textwidth (usually the same for all configs)
    paper_w = max(paper_w_candidates) if paper_w_candidates else 210.0
    paper_h = max(paper_h_candidates) if paper_h_candidates else 277.0
    text_w = max(text_w_candidates)

    # When paper_w has bleed, subtract 2*bleed to get actual page width
    # A4 is 210mm; if paper_w is >220mm, bleed is included
    bleed_candidates = _extract_tex_dimension_mm(cls_content, 'bleed')
    if paper_w > 220:
        bleed = bleed_candidates[0] if bleed_candidates else 3.0
        paper_w = round(paper_w - bleed * 2, 1)
    if paper_h > 280:
        bleed = bleed_candidates[0] if bleed_candidates else 3.0
        paper_h = round(paper_h - bleed * 2, 1)

    # Use the last (most recent/best) oddsidemargin
    oddsidemargin = oddsidemargin_candidates[-1] if oddsidemargin_candidates else 16.4
    topmargin = topmargin_candidates[0] if topmargin_candidates else 10.0

    # Right margin = paper_w - text_w - oddsidemargin
    right_margin = round(paper_w - text_w - oddsidemargin, 1)

    return {
        'paperwidth_mm': float(paper_w),
        'paperheight_mm': float(paper_h),
        'textwidth_mm': float(text_w),
        'oddsidemargin_mm': float(oddsidemargin),
        'right_margin_mm': float(right_margin),
        'topmargin_mm': float(topmargin),
    }


def _apply_page_setup_to_doc(doc, page_geo, footer_dims=None):
    """Apply extracted page geometry to Word document XML.

    Uses direct lxml manipulation because python-docx section API is limited.
    栏数从page_geo.column_count动态获取（由template_config.py从CLS提取）。

    footer_dims: dict from extract_page_footer_dims(), contains bottom_margin_mm
    """
    if not page_geo:
        return

    page_geo = _word_page_geometry(page_geo, footer_dims=footer_dims)

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sectPr = doc.element.body.find(qn('w:sectPr'))
    if sectPr is None:
        return

    # mm to twips: 1mm = 56.6929 twips
    mm_to_twips = lambda v: str(int(round(v * 56.6929)))

    pw = mm_to_twips(page_geo['paperwidth_mm'])
    ph = mm_to_twips(page_geo['paperheight_mm'])

    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.insert(0, pgSz)
    pgSz.set(qn('w:w'), pw)
    pgSz.set(qn('w:h'), ph)

    # 动态计算底边距
    bottom_mm = _word_bottom_margin(page_geo, footer_dims)

    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        sectPr.append(pgMar)

    pgMar.set(qn('w:top'), mm_to_twips(page_geo.get('topmargin_mm', 10)))
    pgMar.set(qn('w:bottom'), mm_to_twips(bottom_mm))
    pgMar.set(qn('w:left'), mm_to_twips(page_geo['oddsidemargin_mm']))
    pgMar.set(qn('w:right'), mm_to_twips(page_geo['right_margin_mm']))
    pgMar.set(qn('w:header'), mm_to_twips(page_geo.get('topmargin_mm', 10)))
    pgMar.set(qn('w:footer'), mm_to_twips(footer_dims.get('footskip_mm', 10.6) if footer_dims else 15))

    # 栏数设置: 从模板CLS动态提取（非硬编码）
    # page_geo 是 dict（来自 get_page_geometry_for_mode）
    col_count = page_geo.get('column_count', 1) if isinstance(page_geo, dict) else getattr(page_geo, 'column_count', 1)
    if col_count >= 2:
        cols = sectPr.find(qn('w:cols'))
        if cols is None:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        cols.set(qn('w:num'), str(col_count))
        # 栏间距: 从模板提取columnsep，默认6mm ≈ 340 twips
        col_sep_mm = page_geo.get('column_sep_mm', 6.0) or 6.0 if isinstance(page_geo, dict) else (getattr(page_geo, 'column_sep_mm', 6.0) or 6.0)
        cols.set(qn('w:space'), mm_to_twips(col_sep_mm))

    print(
        f"  [page setup] {page_geo['paperwidth_mm']}x{page_geo['paperheight_mm']}mm, "
        f"margins: L={page_geo['oddsidemargin_mm']}mm R={page_geo['right_margin_mm']}mm "
        f"T={page_geo['topmargin_mm']}mm B={bottom_mm:.1f}mm, tw={page_geo['textwidth_mm']}mm"
        f", cols={col_count}"
    )


def _safe_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _word_bottom_margin(page_geo, footer_dims=None):
    fallback = 20.0
    if footer_dims and footer_dims.get('bottom_margin_mm') is not None:
        fallback = _safe_float(footer_dims.get('bottom_margin_mm'), fallback)
    paper = _safe_float(page_geo.get('paperheight_mm') if page_geo else 0)
    text = _safe_float(page_geo.get('textheight_mm') if page_geo else 0)
    top = _safe_float(page_geo.get('topmargin_mm') if page_geo else 0)
    inferred = paper - text - top
    if text and 0 < inferred < 60:
        return round(inferred, 2)
    return fallback


def _word_page_geometry(page_geo, footer_dims=None):
    """Map LaTeX page boxes to a Word-representable page geometry."""
    if not isinstance(page_geo, dict):
        return page_geo
    geo = dict(page_geo)
    paper = _safe_float(geo.get('paperwidth_mm'))
    text = _safe_float(geo.get('textwidth_mm'))
    left = _safe_float(geo.get('oddsidemargin_mm'))
    right = _safe_float(geo.get('right_margin_mm'))
    if text and left and (paper < text or right < 0):
        right = right if right > 0 else left
        geo['right_margin_mm'] = round(right, 2)
        geo['paperwidth_mm'] = round(left + text + right, 2)
        paper = _safe_float(geo.get('paperwidth_mm'))
    side_total = paper - text
    if side_total > 0 and left <= 1 and right > 20:
        left = min(25.4 + left, side_total)
        right = side_total - left
        geo['oddsidemargin_mm'] = round(left, 2)
        geo['right_margin_mm'] = round(right, 2)
    top = _safe_float(geo.get('topmargin_mm'))
    if top <= 1 and footer_dims:
        head = _safe_float(footer_dims.get('headheight_mm'))
        sep = _safe_float(footer_dims.get('headsep_mm'))
        if head + sep > top:
            geo['topmargin_mm'] = round(top + head + sep, 2)
    return geo


def _template_line_numbering(cls_content, doc_options):
    """Derive Word line numbering from template/class signals."""
    opts = set(doc_options or [])
    if 'noline' in opts or 'nolineno' in opts:
        return None
    if 'lineno' not in (cls_content or ''):
        return None
    enabled = 'manuscript' in opts or re.search(r'\\linenumbers\b', cls_content or '')
    if not enabled:
        return None
    modulo = 'modulo' in (cls_content or '')
    sep_match = re.search(r'\\linenumbersep\s*([\d.]+)\s*\\?p@', cls_content or '')
    sep_pt = float(sep_match.group(1)) if sep_match else 3.0
    return {'count_by': 5 if modulo else 1, 'distance_twips': int(round(sep_pt * 20))}


def _template_page_numbering(cls_content):
    """Detect whether template footer emits page numbers."""
    if not cls_content:
        return None
    if re.search(r'\\def\\@oddfoot\{[^{}]*\\thepage[^{}]*\}', cls_content):
        return {'footer': 'center'}
    return None


def _latex_vertical_space_pt(expr, baseline_pt):
    text = (expr or '').replace(' ', '')
    if not text:
        return 0.0
    m = re.match(r'([+-]?\d*\.?\d*)?\\baselineskip', text)
    if m:
        factor = float(m.group(1)) if m.group(1) not in ('', '+', '-') else 1.0
        return factor * baseline_pt
    m = re.match(r'([+-]?\d*\.?\d+)\\p@', text)
    if m:
        return float(m.group(1))
    m = re.match(r'([+-]?\d*\.?\d+)(pt|mm|cm|in|pc|em|ex)', text)
    if not m:
        return 0.0
    value = float(m.group(1))
    unit = m.group(2)
    factors = {
        'pt': 1.0,
        'mm': 72.0 / 25.4,
        'cm': 72.0 / 2.54,
        'in': 72.0,
        'pc': 12.0,
        'em': baseline_pt,
        'ex': baseline_pt / 2.0,
    }
    return value * factors[unit]


def _title_prefix_spacing_pt(prefix, baseline_pt):
    total = 0.0
    if r'\parbox' in prefix and '@manuscriptInfo' in prefix:
        total += 2 * baseline_pt
    for expr in re.findall(r'\\vspace\*?\s*\{\s*([^{}]+)\s*\}', prefix):
        total += _latex_vertical_space_pt(expr, baseline_pt)
    vskip = (
        r'\\vskip\s*('
        r'(?:[+-]?\d*\.?\d+\s*)?\\(?:baselineskip|p@)|'
        r'[+-]?\d*\.?\d+\s*(?:pt|mm|cm|in|pc|em|ex))'
    )
    for expr in re.findall(vskip, prefix):
        total += _latex_vertical_space_pt(expr, baseline_pt)
    return max(total, 0.0)


def _template_title_before_space_pt(cls_content, doc_options, baseline_pt):
    """Derive first-page title top spacing from the template maketitle block."""
    if not cls_content:
        return 0.0
    opts = set(doc_options or [])
    patterns = []
    if 'manuscript' in opts:
        patterns.append(r'\\def\\@@maketitlemanuscript\{(?P<body>.*?)\\@title')
    else:
        patterns.extend([
            r'\\def\\@@maketitlefinal\{(?P<body>.*?)\\@title',
            r'\\def\\@maketitle\{(?P<body>.*?)\\@title',
            r'\\long\\def\\pprintMaketitle\{(?P<body>.*?)\\@title',
            r'\\long\\def\\MaketitleBox\{(?P<body>.*?)\\@title',
        ])
    for pattern in patterns:
        for block in re.finditer(pattern, cls_content, re.S):
            value = _title_prefix_spacing_pt(block.group('body'), baseline_pt)
            if value:
                return round(value, 2)
    return 0.0


def _extract_template_length_pt(cls_content, name, body_pt):
    if not cls_content:
        return None
    pattern = (
        r'\\setlength\s*\\?' + re.escape(name) +
        r'\s*\{\s*([^{}]+?)\s*\}'
    )
    matches = re.findall(pattern, cls_content)
    if not matches:
        return None
    value = _latex_vertical_space_pt(matches[-1], body_pt)
    return value if value or str(matches[-1]).strip().startswith("0") else None


def _extract_caption_spacing(cls_content, body_pt):
    """Return template caption/float skips in points."""
    above = _extract_template_length_pt(cls_content, "abovecaptionskip", body_pt)
    below = _extract_template_length_pt(cls_content, "belowcaptionskip", body_pt)
    textfloat = _extract_template_length_pt(cls_content, "dbltextfloatsep", body_pt)
    if textfloat is None:
        textfloat = _extract_template_length_pt(cls_content, "textfloatsep", body_pt)
    cls_text = cls_content or ""
    if above is None and "abovecaptionskip" in cls_text:
        above = body_pt
    if below is None and "belowcaptionskip" in cls_text:
        below = 0.0
    if textfloat is None and "textfloatsep" in cls_text:
        textfloat = body_pt * 2
    if textfloat is None and ("abovecaptionskip" in cls_text or "@makecaption" in cls_text):
        textfloat = body_pt * 2
    return {
        "above_pt": above,
        "below_pt": below,
        "textfloat_pt": textfloat,
    }


def _extract_macro_font_spec(cls_content, macro_name, default_size, default_baseline):
    if not cls_content:
        return None
    start = cls_content.find(r'\def\\' + macro_name)
    if start < 0:
        start = cls_content.find(r'\def' + '\\' + macro_name)
    if start < 0:
        return None
    brace = cls_content.find('{', start)
    if brace < 0:
        return None
    depth = 0
    end = None
    for idx in range(brace, len(cls_content)):
        ch = cls_content[idx]
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end = idx
                break
    if end is None:
        return None
    body = cls_content[brace + 1:end]
    size = default_size
    baseline = default_baseline
    font_match = re.search(
        r'\\fontsize\s*\{\s*([^{}]+)\s*\}\s*\{\s*([^{}]+)\s*\}',
        body,
    )
    if font_match:
        size = _latex_size_token_to_pt(font_match.group(1), default_size)
        baseline = _latex_size_token_to_pt(font_match.group(2), default_baseline)
    align = None
    if r'\raggedright' in body:
        align = 'LEFT'
    elif r'\centering' in body:
        align = 'CENTER'
    elif r'\raggedleft' in body:
        align = 'RIGHT'
    return {
        "size": size,
        "baseline": baseline,
        "bold": r'\bfseries' in body,
        "align": align,
    }


def _extract_abstract_style(cls_content, small_size, small_baseline):
    spec = _extract_macro_font_spec(
        cls_content, "abstractfont", small_size, small_baseline)
    if spec is None:
        spec = {"size": small_size, "baseline": small_baseline, "align": None}
    return spec


def _template_front_matter_indent_pt(cls_content, body_pt, config_mode=None, doc_options=None):
    """_abstract_indent_mode_v2: 从对应模式的 maketitle 块提取摘要 leftmargin。

    manuscript→\@@maketitlemanuscript, final→\@@maketitlefinal。
    避免错提其他模式的缩进(manuscript 摘要本无 leftmargin)。
    """
    block = _maketitle_body_for_mode(cls_content, config_mode, doc_options)
    if not block:
        return None
    match = re.search(
        r'\\leftmargin\s*([0-9.]+)\s*(mm|cm|pt|pc|in|em|ex)',
        block,
    )
    if not match:
        return None
    return _latex_vertical_space_pt(''.join(match.groups()), body_pt)


def _template_abstract_after_space_pt(cls_content, baseline_pt):
    block = cls_content or ""
    patterns = [
        r'\\@abstr.*?\\par\\vspace\s*\{\s*([^{}]+)\s*\}',
        r'\\endtrivlist\s*\\vspace\s*\{\s*([^{}]+)\s*\}',
        r'\\par\\vspace\s*\{\s*([^{}]+)\s*\}\s*\}?\s*\}?\s*$',
    ]
    for pattern in patterns:
        match = re.search(pattern, block, re.S)
        if match:
            value = _latex_vertical_space_pt(match.group(1), baseline_pt)
            if value:
                return round(value, 2)
    return round(baseline_pt, 2)


def _apply_line_numbering_to_doc(doc, line_spec):
    """Write Word section line-numbering settings from template-derived spec."""
    if not line_spec:
        return
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sectPr = doc.element.body.find(qn('w:sectPr'))
    if sectPr is None:
        return
    old = sectPr.find(qn('w:lnNumType'))
    if old is not None:
        sectPr.remove(old)
    ln = OxmlElement('w:lnNumType')
    ln.set(qn('w:countBy'), str(line_spec.get('count_by', 1)))
    ln.set(qn('w:distance'), str(line_spec.get('distance_twips', 60)))
    ln.set(qn('w:restart'), 'newPage')
    sectPr.append(ln)
    print(f"  [line numbers] countBy={ln.get(qn('w:countBy'))}, restart=newPage")


def _apply_page_number_footer_to_doc(doc, page_spec):
    """Insert a centered PAGE field when the template footer uses page numbers."""
    if not page_spec:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for child in (begin, instr, separate, text, end):
        run._r.append(child)
    run.bold = True
    print('  [page footer] inserted centered PAGE field')


def _apply_column_setup_to_doc(doc, column_count, column_sep_mm=None):
    """Apply template-derived columns without changing page geometry."""
    if column_count < 2:
        return
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sectPr = doc.element.body.find(qn('w:sectPr'))
    if sectPr is None:
        return
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(column_count))
    if column_sep_mm:
        cols.set(qn('w:space'), str(int(round(column_sep_mm * 56.6929))))


def _extract_column_sep_mm(cls_content):
    """Extract the last explicit column separation from the class."""
    unit_to_mm = {'mm': 1.0, 'cm': 10.0, 'pt': 0.3528, 'pc': 4.2336, 'in': 25.4}
    matches = re.findall(
        r'\\(?:setlength\s*)?\\?columnsep(?:\s*\{)?\s*([\d.]+)\s*(mm|cm|pt|pc|in)',
        cls_content or '',
    )
    if not matches:
        return None
    value, unit = matches[-1]
    return float(value) * unit_to_mm[unit]


def _extract_parindent_pt(cls_content, body_size):
    """Extract first-line indent from the class file without policy overrides."""
    if not cls_content:
        return None
    match = re.search(r"\\parindent\s*([0-9.]+)?\s*(em|pt|cm|mm|in)", cls_content)
    if not match:
        return None
    value = float(match.group(1) or "1")
    unit = match.group(2)
    if unit == "em":
        return value * body_size
    if unit == "pt":
        return value
    if unit == "cm":
        return value * 28.3465
    if unit == "mm":
        return value * 2.83465
    if unit == "in":
        return value * 72
    return None


_LATEX_COLOR_RGB = {
    'black': (0, 0, 0), 'white': (255, 255, 255),
    'red': (255, 0, 0), 'green': (0, 128, 0), 'blue': (0, 0, 255),
    'yellow': (255, 255, 0), 'magenta': (255, 0, 255), 'cyan': (0, 255, 255),
    'gray': (128, 128, 128), 'grey': (128, 128, 128),
    'darkgray': (64, 64, 64), 'lightgray': (211, 211, 211),
    'orange': (255, 165, 0), 'brown': (165, 42, 42),
    'purple': (128, 0, 128), 'pink': (255, 192, 203),
}


def _slice_maketitle_block(cls_content):
    """切出 \\@maketitle / \\@@maketitle* 定义块（到下一个 \\def 或文件尾）。"""
    m = re.search(r'\\def\\@+maketitle\w*\b.*', cls_content, re.S)
    if not m:
        return cls_content
    chunk = m.group(0)
    nxt = re.search(r'\\def\\', chunk[10:])
    return chunk[: nxt.start() + 10] if nxt else chunk


def _resolve_definecolor_rgb(cls_content, name):
    """解析 \\definecolor{name}{rgb}{r,g,b} 为 0-255 tuple，失败 None。"""
    matches = re.findall(
        r'\\definecolor\s*\{\s*' + re.escape(name) + r'\s*\}\s*\{rgb\}\s*\{([0-9.,\s]+)\}',
        cls_content,
    )
    if not matches:
        return None
    try:
        parts = [float(x.strip()) for x in matches[-1].split(',')]
    except ValueError:
        return None
    if len(parts) != 3:
        return None
    return tuple(int(round(min(max(v, 0.0), 1.0) * 255)) for v in parts)


def _extract_heading_color(cls_content, maketitle_block=None):
    """从 CLS maketitle 块提标题颜色（未指定返回 None=默认黑）。

    顺序：命名颜色 → \\definecolor{rgb} 自定义 → None。
    """
    if not cls_content:
        return None
    block = maketitle_block or _slice_maketitle_block(cls_content)
    color_m = re.search(r'\\(?:color|textcolor)\s*\{\s*([A-Za-z]+)\s*\}', block)
    if not color_m:
        return None
    name = color_m.group(1).lower()
    return _LATEX_COLOR_RGB.get(name) or _resolve_definecolor_rgb(cls_content, name)


def _cfg_block_for_options(cfg_content, doc_options, config_mode=None):
    """_cfg_block_v2: 按模式选 cfg 块,支持 \ifacp 内嵌 \if@stage@final/\else。

    copernicus.cfg 结构: \ifacp ... \if@stage@final (final色) \else (manuscript色) \fi ... \fi
    final 模式取 @stage@final 分支,其他模式取 \else 分支(或无嵌套时取整个块)。
    """
    options = [str(item).strip().lower() for item in (doc_options or [])]
    mode = _word_config_mode(config_mode, options)
    for opt in options:
        match = re.search(
            r'^\\if' + re.escape(opt) + r'\b(?P<body>.*?)(?=^\\if[a-zA-Z]+|\\endinput|\Z)',
            cfg_content or "",
            re.S | re.M,
        )
        if not match:
            continue
        body = match.group('body')
        # 块内若有 \if@stage@final/\else 嵌套,按 mode 选分支
        stage_match = re.search(
            r'\\if@stage@final(?P<then>.*?)(?:\\else(?P<else_>.*?))?\\fi',
            body, re.S)
        if stage_match:
            if mode == "final":
                return stage_match.group('then') or body
            # manuscript/discussions 等取 else 分支(颜色通常是黑色)
            return stage_match.group('else_') or body
        return body
    return ""


def _cfg_text_color(cfg_content, doc_options, config_mode=None):
    block = _cfg_block_for_options(cfg_content, doc_options, config_mode)
    return _resolve_definecolor_rgb(block, "textcol") if block else None


def _read_template_cfg(tex_dir, cls_name, cls_path):
    candidates = []
    if cls_name:
        candidates.extend([
            tex_dir / f"{cls_name}.cfg",
            tex_dir / f"{cls_name}_paper" / f"{cls_name}.cfg",
        ])
    if cls_path:
        candidates.append(Path(cls_path).with_suffix(".cfg"))
    candidates.extend(tex_dir.glob("*_paper/*.cfg"))
    for path in candidates:
        if Path(path).exists():
            return _read_text_file(path)
    return ""


def _macro_def_body(cls_content, macro_name):
    for prefix in (r'\def', r'\long\def'):
        marker = prefix + macro_name
        start = (cls_content or "").find(marker)
        if start < 0:
            continue
        brace = cls_content.find('{', start + len(marker))
        if brace >= 0:
            return _balanced_group(cls_content, brace)
    return ""


def _maketitle_body_for_mode(cls_content, config_mode=None, doc_options=None):
    mode = _word_config_mode(config_mode, doc_options or [])
    if mode == "manuscript":
        names = [r'\@@maketitlemanuscript', r'\@maketitle']
    else:
        names = [
            r'\@@maketitlefinal',
            r'\pprintMaketitle',
            r'\MaketitleBox',
            r'\@maketitle',
            r'\@@maketitlemanuscript',
        ]
    for name in names:
        body = _macro_def_body(cls_content, name)
        if body:
            return body
    return _slice_maketitle_block(cls_content or "")


def _latex_alignment_from_block(block):
    positions = [
        ("LEFT", block.rfind(r'\raggedright')),
        ("LEFT", block.rfind(r'\flushleft')),
        ("CENTER", block.rfind(r'\centering')),
        ("CENTER", block.rfind(r'\begin{center}')),
        ("CENTER", block.rfind(r'\center{')),
        ("RIGHT", block.rfind(r'\raggedleft')),
        ("RIGHT", block.rfind(r'\flushright')),
    ]
    align, pos = max(positions, key=lambda item: item[1])
    return align if pos >= 0 else None


def _extract_title_author_styles(cls_content, body_size, config_mode=None, doc_options=None):
    r"""从CLS的\maketitle定义中提取标题和作者的字号+对齐方式

    从 \@@maketitlemanuscript 或 \@maketitle 中提取：
    - 字号：LaTeX字号命令(\LARGE, \large等)映射到pt值
    - 对齐：\raggedright/\flushleft → LEFT, \centering → CENTER

    Args:
        cls_content: .cls文件内容
        body_size: 正文字号(pt)，用于计算相对值

    Returns:
        dict: {title_size, author_size, title_align, author_align}
    """
    # LaTeX标准字号映射（基于11pt正文）
    latex_size_map = {
        'tiny': 6, 'scriptsize': 8, 'footnotesize': 9,
        'small': 10, 'normalsize': 11, 'large': 12,
        'Large': 14, 'LARGE': 17, 'huge': 20, 'Huge': 25,
    }
    result = {
        'title_size': None,
        'author_size': None,
        'title_align': None,
        'author_align': None,
        'title_bold': None,
        'author_bold': None,
        'heading_color': None,
    }

    if not cls_content:
        return result

    title_font = _extract_macro_font_spec(cls_content, "titlefont", None, None)
    if title_font:
        result['title_size'] = title_font.get("size")
        result['title_bold'] = title_font.get("bold")
        result['title_align'] = title_font.get("align")

    author_font = _extract_macro_font_spec(cls_content, "authorfont", None, None)
    if author_font:
        result['author_size'] = author_font.get("size")
        result['author_bold'] = author_font.get("bold")
        result['author_align'] = author_font.get("align")

    # 匹配标题区域: {\SIZECMD\bfseries\@title
    title_m = re.search(
        r'\{\\(LARGE|Large|large|huge|Huge)\s+\\bfseries\s*\\@title',
        cls_content
    )
    if title_m and result['title_size'] is None:
        cmd = title_m.group(1)
        result['title_size'] = latex_size_map.get(cmd)
        result['title_bold'] = True

    # 匹配作者区域: {\SIZECMD ... \@author
    author_m = re.search(
        r'\{\\(LARGE|Large|large|huge|Huge|normalsize)\s+[^}]*\\@author',
        cls_content
    )
    if author_m and result['author_size'] is None:
        cmd = author_m.group(1)
        result['author_size'] = latex_size_map.get(cmd)
        result['author_bold'] = 'bfseries' in author_m.group(0)

    # 从maketitle块中提取对齐方式
    # 查找 \@@maketitlemanuscript 块内容
    maketitle_m = re.search(
        r'\\def\\@@maketitlemanuscript\b.*?'
        r'(\\(?:raggedright|centering|flushleft|begin\{center\}))',
        cls_content, re.S
    )
    if maketitle_m:
        align_cmd = maketitle_m.group(1)
        if 'raggedright' in align_cmd:
            result['title_align'] = 'LEFT'
            result['author_align'] = 'LEFT'  # raggedright影响整个块
        elif 'centering' in align_cmd:
            result['title_align'] = 'CENTER'
            result['author_align'] = 'CENTER'

    maketitle_body = _maketitle_body_for_mode(cls_content, config_mode, doc_options)
    title_align = _latex_alignment_from_block(maketitle_body.split(r'\@title', 1)[0])
    if title_align:
        result['title_align'] = title_align
        result['author_align'] = title_align

    # 单独检查author块中的对齐（可能覆盖全局设置）
    author_block_m = re.search(
        r'\\if@noauthor.*?'
        r'(\\(?:flushleft|raggedright|centering|begin\{center\}))',
        cls_content, re.S
    )
    if author_block_m:
        align_cmd = author_block_m.group(1)
        if 'flushleft' in align_cmd or 'raggedright' in align_cmd:
            result['author_align'] = 'LEFT'
        elif 'centering' in align_cmd:
            result['author_align'] = 'CENTER'

    author_align = _latex_alignment_from_block(maketitle_body.split(r'\@author', 1)[0])
    if author_align:
        result['author_align'] = author_align

    result['heading_color'] = _extract_heading_color(cls_content, maketitle_body)
    return result


def _load_template_word_style(tex_path=None, config_mode=None):
    """Derive Word style choices from the generated TeX, class file, and template spec.

    config_mode: 模板配置模式名 ('manuscript', 'final', 'discussions')
                 None 时默认 'manuscript'
    """
    if not tex_path:
        return None
    tex_path = Path(tex_path)
    tex_dir = tex_path.parent
    compiled_layout = _compiled_layout_twips(tex_dir)
    tex_content = _read_text_file(tex_path)
    if not tex_content:
        return None

    spec = {}
    spec_paths = list(tex_dir.glob("*_paper/*_template_spec.json")) + list(tex_dir.glob("*template_spec.json"))
    if spec_paths:
        try:
            spec = json.loads(_read_text_file(spec_paths[0]))
        except Exception:
            spec = {}

    docclass_match = re.search(r"\\documentclass(?:\[([^\]]*)\])?\{([^}]+)\}", tex_content)
    option_text = docclass_match.group(1) if docclass_match and docclass_match.group(1) else ""
    doc_options = [item.strip().lower() for item in option_text.split(",") if item.strip()]
    cls_name = docclass_match.group(2).strip() if docclass_match else ""
    effective_mode = _word_config_mode(config_mode, doc_options)

    cls_candidates = []
    if cls_name:
        cls_candidates.extend([tex_dir / f"{cls_name}.cls", tex_dir / f"{cls_name}_paper" / f"{cls_name}.cls"])
    cls_candidates.extend(tex_dir.glob("*_paper/*.cls"))
    cls_path = next((Path(path) for path in cls_candidates if Path(path).exists()), None)
    cls_content = _read_text_file(cls_path) if cls_path else ""
    cfg_content = _read_template_cfg(tex_dir, cls_name, cls_path)

    main_font = _extract_first_tex_command_value(tex_content, "setmainfont")
    sans_font = _extract_first_tex_command_value(tex_content, "setsansfont")
    mono_font = _extract_first_tex_command_value(tex_content, "setmonofont")
    cjk_font = _extract_first_tex_command_value(tex_content, "setCJKmainfont")

    required_packages = [str(pkg).lower() for pkg in spec.get("required_packages", [])]
    if not main_font and any(pkg in required_packages for pkg in ("times", "txfonts", "mathptmx", "newtxtext")):
        main_font = "Times New Roman"
    if not sans_font and any(pkg in required_packages for pkg in ("helvet", "hvmath", "hvams")):
        sans_font = "Arial"
    if not mono_font and any(pkg in required_packages for pkg in ("courier", "newtxtt")):
        mono_font = "Courier New"
    body_size, body_baseline = _extract_latex_font_size(cls_content, "normalsize", 10, 12)
    small_size, small_baseline = _extract_latex_font_size(cls_content, "small", max(body_size - 1, 8), max(body_baseline - 1, 10))
    abstract_style = _extract_abstract_style(cls_content, small_size, small_baseline)
    caption_spacing = _extract_caption_spacing(cls_content, body_size)
    parindent = _extract_parindent_pt(cls_content, body_size)

    # 从CLS中提取标题和作者字号+对齐方式（不硬编码）
    ta_styles = _extract_title_author_styles(
        cls_content,
        body_size,
        config_mode=effective_mode,
        doc_options=doc_options,
    )
    title_size = ta_styles['title_size']
    author_size = ta_styles['author_size']
    title_align = ta_styles['title_align']
    author_align = ta_styles['author_align']
    title_bold = ta_styles['title_bold']
    author_bold = ta_styles['author_bold']
    heading_color = _cfg_text_color(cfg_content, doc_options, effective_mode) or ta_styles['heading_color']

    stretch_match = re.search(r"\\gdef\\baselinestretch\{([\d.]+)\}", cls_content)
    line_spacing = float(stretch_match.group(1)) if stretch_match and "manuscript" in doc_options else None
    if line_spacing is None:
        line_spacing = max(body_baseline / body_size, 1.0) if body_size else 1.0

    section_block_match = re.search(r"\\def\\section\b.*?(?=\\def\\subsection\b|\\def\\@sect\b|$)", cls_content, flags=re.S)
    section_block = section_block_match.group(0) if section_block_match else cls_content
    heading_bold = "bfseries" in section_block
    heading_uses_sans = "sffamily" in section_block

    page_geometry = None
    footer_dims = None
    section_spacing = {}
    line_numbering = None
    page_numbering = None
    title_before_pt = 0.0
    if cls_content:
        from shared.template_config import (
            get_page_geometry_for_mode, extract_section_spacing,
            extract_page_footer_dims, detect_effective_column_count,
        )
        page_geometry = get_page_geometry_for_mode(cls_content, config_mode=effective_mode)
        column_count = detect_effective_column_count(cls_content, doc_options)
        column_sep_mm = _extract_column_sep_mm(cls_content)
        if page_geometry:
            page_geometry['column_count'] = column_count
            if column_sep_mm:
                page_geometry['column_sep_mm'] = column_sep_mm
        section_spacing = extract_section_spacing(cls_content, body_size_pt=body_size, config_mode=effective_mode)
        footer_dims = extract_page_footer_dims(cls_content, config_mode=effective_mode)
        line_numbering = _template_line_numbering(cls_content, doc_options)
        page_numbering = _template_page_numbering(cls_content)
        title_before_pt = _template_title_before_space_pt(cls_content, doc_options, body_baseline)
        front_indent_pt = _template_front_matter_indent_pt(cls_content, body_size, config_mode=effective_mode, doc_options=doc_options)
        abstract_after_pt = _template_abstract_after_space_pt(
            cls_content, abstract_style.get("baseline") or small_baseline)
    else:
        column_count = 1
        column_sep_mm = None
        front_indent_pt = None
        abstract_after_pt = small_baseline

    return {
        "main_font": main_font,
        "sans_font": sans_font,
        "mono_font": mono_font,
        "cjk_font": cjk_font or "宋体",
        "body_size": body_size,
        "small_size": small_size,
        "abstract_size": abstract_style.get("size"),
        "abstract_baseline": abstract_style.get("baseline"),
        "abstract_align": abstract_style.get("align"),
        "abstract_left_indent_pt": front_indent_pt,
        "abstract_right_indent_pt": front_indent_pt,
        "abstract_after_pt": abstract_after_pt,
        "caption_spacing": caption_spacing,
        "parindent": parindent,
        "line_spacing": line_spacing,
        "heading_bold": heading_bold,
        "heading_font": sans_font if heading_uses_sans else main_font,
        "heading_uses_sans": heading_uses_sans,
        "title_size": title_size,
        "author_size": author_size,
        "title_align": title_align,
        "author_align": author_align,
        "title_bold": title_bold,
        "author_bold": author_bold,
        "heading_color": heading_color,
        "page_geometry": page_geometry,
        "compiled_layout": compiled_layout,
        "column_count": column_count,
        "column_sep_mm": column_sep_mm,
        "footer_dims": footer_dims,
        "section_spacing": section_spacing,
        "line_numbering": line_numbering,
        "page_numbering": page_numbering,
        "title_before_pt": title_before_pt,
    }


def _word_config_mode(config_mode, doc_options):
    """Use the generated TeX mode when a standalone Word conversion omits it."""
    if config_mode not in (None, ''):
        return config_mode
    options = {str(item).strip().lower() for item in (doc_options or [])}
    if 'manuscript' in options or 'classic' in options:
        return 'manuscript'
    if 'final' in options:
        return 'final'
    if 'discussions' in options or 'discussion' in options:
        return 'discussions'
    return config_mode


def _set_style_fonts(style, en_font, cjk_font, size=None, bold=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    if en_font:
        style.font.name = en_font
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if en_font or cjk_font:
        rPr = style.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.element.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        if en_font:
            rFonts.set(qn("w:ascii"), en_font)
            rFonts.set(qn("w:hAnsi"), en_font)
            rFonts.set(qn("w:cs"), en_font)
        if cjk_font:
            rFonts.set(qn("w:eastAsia"), cjk_font)


def _get_doc_style(doc, style_name):
    for style in doc.styles:
        if style.name == style_name:
            return style
    try:
        return doc.styles[style_name]
    except KeyError:
        return None


def _is_caption_paragraph(para):
    """判断段落是否为图例/表例说明（caption）

    区分真正的图例说明和叙事引用：
    - 图例: "图4.3 中国通量..." (以图号+分隔符开头)
    - 叙事引用: "图4.1中展示了..." (图号后紧跟"中/是/为"等动词)
    - 正文引用: "图4.1中展示了卫星和模拟的评估偏差情况..." → 正文
    """
    text = para.text.strip()
    if not text:
        return False
    style_name = (para.style.name or "").lower() if para.style else ""
    if "caption" in style_name:
        return True
    # 英文: Figure 1. / Table 2. (后面必须跟分隔符或空格+描述文字)
    if re.match(r"^(Figure|Table)\s+[\w.\-]+\.?\s+\S", text):
        return True
    # 中文图例: 图2.1 / 图 4.1 / 表3 / 表 2
    # 严格模式: 图号后必须是空格+描述，或标点+描述
    cap_m = re.match(r"^(图|表)\s*[\d\.]+\s*", text)
    if cap_m:
        after_num = text[cap_m.end():]
        # 空 → 不是图例
        if not after_num:
            return False
        # 叙事引用检测：图号后紧跟"中/是/为/说明/展示/可以看到/所示/的" → 正文
        if re.match(r'^(中|是|为|说明|展示|可以看到|所示|的|和|与|对|在|从|到|及|等|将|被|有|不|也)', after_num):
            return False
        # 中文图例要求：图号后紧跟空格或标点，然后是描述文字
        if re.match(r'^[\s，。、；：]?\S', after_num):
            return True
        return False
    return False


def _has_cjk_text(text):
    return bool(re.search(r'[\u3400-\u4dbf\u4e00-\u9fff]', text or ''))


def _doc_has_cjk_body_text(doc):
    references_active = False
    no_indent_styles = {"Title", "Author", "Date", "Abstract Title", "Subtitle"}
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        if text.lower() == "references":
            references_active = True
            continue
        if (
            not text or references_active or style_name.startswith("Heading")
            or style_name in no_indent_styles or _is_caption_paragraph(para)
        ):
            continue
        if _has_cjk_text(text):
            return True
    return False


def _body_parindent_pt(style_spec, has_cjk_body):
    body_size = float(style_spec.get("body_size") or 10)
    template_indent = float(style_spec.get("parindent") or body_size)
    return max(template_indent, body_size * 2) if has_cjk_body else template_indent


def _prepend_bold_run(para, text):
    """在段落最前面（pPr 之后）插入一个 bold run，保留段落其余内容与结构。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(r)
    else:
        para._element.insert(0, r)


def _normalize_abstract_paragraphs(doc, cls_content=None):
    """Merge Abstract title and body into one paragraph.

    CLS模板格式: \\bfseries Abstract. (加粗标题+句点) + 正文内容
    标题和正文在同一行，标题加粗，正文不加粗。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _style_name(para):
        return (para.style.name or "") if para.style else ""

    def _is_heading_or_boundary(para):
        style = _style_name(para)
        text = para.text.strip()
        if style.startswith("Heading"):
            return True
        if re.match(r"^\d+(?:\.\d+)*\s+\S+", text):
            return True
        return False

    paragraphs = list(doc.paragraphs)
    title_idx = None
    for idx, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        style = _style_name(para).replace(" ", "").lower()
        if text in {"abstract", "abstract:", "abstract."} or style == "abstracttitle":
            title_idx = idx
            break

    if title_idx is None:
        return

    title_para = paragraphs[title_idx]

    # Collect body paragraphs after the title
    body = []
    for para in paragraphs[title_idx + 1:]:
        if not para.text.strip():
            if body:
                break
            continue
        if _is_heading_or_boundary(para):
            break
        style = _style_name(para).lower()
        if not body and "abstract" not in style and len(para.text.strip()) < 20:
            break
        body.append(para)

    if not body:
        return

    # 解析 abstract 标签分隔符（CLS \bfseries\abstractname 后的 . 或 :），默认句点
    abstract_label = "Abstract. "
    if cls_content:
        abs_m = re.search(r'\\bfseries\\abstractname([.:\s])', cls_content)
        if abs_m:
            abstract_label = f"Abstract{abs_m.group(1)} "

    # 保留摘要原始段落结构（不合并）：首段开头插入 bold 标签，所有段继承标题样式（无缩进）
    first_para = body[0]
    _prepend_bold_run(first_para, abstract_label)
    target_style = title_para.style
    for para in body:
        para.style = target_style

    # 原标题段（"Abstract" 单独行）已被首段 bold 标签取代，移除
    title_para._element.getparent().remove(title_para._element)
    print(f"  [abstract] kept {len(body)} paragraph(s); bold label prepended to first")


def _row_border_spec(row, edge):
    from docx.oxml.ns import qn

    best = None
    for tc in row._tr.findall(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        borders = tcPr.find(qn("w:tcBorders")) if tcPr is not None else None
        elem = borders.find(qn(f"w:{edge}")) if borders is not None else None
        if elem is None:
            continue
        val = elem.get(qn("w:val"))
        if val in (None, "nil", "none"):
            continue
        sz = elem.get(qn("w:sz")) or "4"
        color = elem.get(qn("w:color")) or "000000"
        if best is None or int(sz) > int(best[1]):
            best = (val, sz, color)
    return best


def _table_border_spec(tbl, edge):
    from docx.oxml.ns import qn

    tbl_pr = tbl._element.find(qn("w:tblPr"))
    borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    elem = borders.find(qn(f"w:{edge}")) if borders is not None else None
    if elem is None:
        return None
    val = elem.get(qn("w:val"))
    if val in (None, "nil", "none"):
        return None
    return (val, elem.get(qn("w:sz")) or "4", elem.get(qn("w:color")) or "000000")


def _infer_bottom_rule_spec(tbl):
    if not tbl.rows:
        return None
    candidates = [
        _row_border_spec(tbl.rows[-1], "bottom"),
        _table_border_spec(tbl, "bottom"),
        _row_border_spec(tbl.rows[-1], "top"),
        _row_border_spec(tbl.rows[0], "top"),
    ]
    for row in reversed(tbl.rows[:-1]):
        candidates.append(_row_border_spec(row, "bottom"))
        candidates.append(_row_border_spec(row, "top"))
    for item in candidates:
        if item is not None:
            return item
    return ("single", "8", "000000")


def _complete_horizontal_table_borders(tbl):
    if _is_layout_floating_table_elem(tbl._element):
        _clear_table_borders(tbl._element)
        return
    if len(tbl.rows) <= 1 and len(tbl.columns) <= 1:
        return

    row_count = len(tbl.rows)
    for ri, row in enumerate(tbl.rows):
        top = _row_border_spec(row, "top")
        bottom = _row_border_spec(row, "bottom")
        if ri == 0 and top is None:
            top = ("single", "8", "000000")
        if ri == row_count - 1 and bottom is None:
            bottom = ("single", "8", "000000")

        for cell in row.cells:
            kwargs = {
                "left": ("nil", "0", "auto"),
                "right": ("nil", "0", "auto"),
            }
            if top is not None:
                kwargs["top"] = top
            if bottom is not None:
                kwargs["bottom"] = bottom
            _set_cell_border(cell, **kwargs)


def _complete_table_bottom_line(tbl):
    """仅补全表格最后一行的底线，不补中间横线"""
    from docx.oxml.ns import qn

    if _is_layout_floating_table_elem(tbl._element):
        _clear_table_borders(tbl._element)
        return
    rows = tbl._element.findall(qn("w:tr"))
    if not rows:
        return
    bottom = _infer_bottom_rule_spec(tbl)
    for tc in rows[-1].findall(qn("w:tc")):
        _set_physical_tc_border(tc, bottom=bottom)


def _row_text_stats(tr):
    from docx.oxml.ns import qn

    texts = [node.text or "" for node in tr.iter(qn("w:t"))]
    para_count = sum(1 for _ in tr.iter(qn("w:p")))
    return len("".join(texts)), para_count


def _should_keep_row_together(tr):
    return True


def _set_row_keep_together(tr, keep):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        if not keep:
            return
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    cant_split = trPr.find(qn("w:cantSplit"))
    if keep and cant_split is None:
        trPr.append(OxmlElement("w:cantSplit"))
    elif not keep and cant_split is not None:
        trPr.remove(cant_split)


def _set_paragraph_bool_prop(p_elem, tag, keep=True):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        if not keep:
            return
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    prop = pPr.find(qn(tag))
    if keep and prop is None:
        pPr.append(OxmlElement(tag))
    elif not keep and prop is not None:
        pPr.remove(prop)


def _keep_table_on_one_page(table_elem):
    from docx.oxml.ns import qn

    rows = list(table_elem.iter(qn("w:tr")))
    for row_idx, tr in enumerate(rows):
        _set_row_keep_together(tr, True)
        keep_next = row_idx < len(rows) - 1
        for p in tr.iter(qn("w:p")):
            _set_paragraph_bool_prop(p, "w:keepLines", True)
            _set_paragraph_bool_prop(p, "w:keepNext", keep_next)


def _keep_float_block_together(block_elems):
    from docx.oxml.ns import qn

    meaningful = [
        elem for elem in block_elems
        if elem is not None and (
            elem.tag == qn("w:tbl")
            or (elem.tag == qn("w:p") and not _is_empty_section_break_para(elem))
        )
    ]
    for idx, elem in enumerate(meaningful):
        keep_next = idx < len(meaningful) - 1
        if elem.tag == qn("w:p"):
            _set_paragraph_bool_prop(elem, "w:keepLines", True)
            _set_paragraph_bool_prop(elem, "w:keepNext", keep_next)
            continue
        if elem.tag == qn("w:tbl"):
            _keep_table_on_one_page(elem)
            if keep_next:
                for p in elem.iter(qn("w:p")):
                    _set_paragraph_bool_prop(p, "w:keepNext", True)


def _enforce_caption_table_keep_together(doc):
    from docx.oxml.ns import qn

    body = doc.element.body
    elems = list(body)
    for idx, elem in enumerate(elems):
        if elem.tag != qn("w:p"):
            continue
        text = _paragraph_xml_text(elem).strip()
        if not _is_caption_like_text(text) or not re.match(r"^\s*Table\b", text, re.I):
            continue
        blanks = []
        nxt = elem.getnext()
        while nxt is not None and nxt.tag == qn("w:p") and not _paragraph_xml_text(nxt).strip():
            blanks.append(nxt)
            nxt = nxt.getnext()
        if nxt is None or nxt.tag != qn("w:tbl"):
            continue
        _set_paragraph_bool_prop(elem, "w:keepLines", True)
        _set_paragraph_bool_prop(elem, "w:keepNext", True)
        for blank in blanks:
            _set_paragraph_bool_prop(blank, "w:keepNext", True)
        _keep_table_on_one_page(nxt)


def _apply_template_styles_to_doc(doc, style_spec):
    from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    # 0. 应用页面几何设置 (从.cls提取，非硬编码)
    page_geo = _merge_page_geometry_with_layout(style_spec)
    footer_dims = style_spec.get("footer_dims")
    if page_geo:
        _apply_page_setup_to_doc(doc, page_geo, footer_dims=footer_dims)
    else:
        _apply_column_setup_to_doc(
            doc,
            style_spec.get("column_count", 1),
            style_spec.get("column_sep_mm"),
        )
    _apply_line_numbering_to_doc(doc, style_spec.get("line_numbering"))
    _apply_page_number_footer_to_doc(doc, style_spec.get("page_numbering"))

    main_font = style_spec["main_font"]
    heading_font = style_spec["heading_font"]
    cjk_font = style_spec["cjk_font"]
    body_size = style_spec["body_size"]
    small_size = style_spec["small_size"]
    abstract_size = style_spec.get("abstract_size") or small_size
    abstract_baseline = style_spec.get("abstract_baseline") or abstract_size
    line_spacing = style_spec["line_spacing"]
    abstract_line_spacing = (
        max(float(abstract_baseline) / float(abstract_size), 1.0)
        if abstract_size else line_spacing
    )
    caption_spacing = style_spec.get("caption_spacing") or {}
    # 首行缩进2个字符（匹配PDF编译结果）
    parindent = _body_parindent_pt(style_spec, _doc_has_cjk_body_text(doc))
    style_spec["applied_parindent"] = parindent
    # 标题/作者字号+对齐：从CLS动态提取，非硬编码
    title_size = style_spec.get("title_size")
    author_size = style_spec.get("author_size")
    title_bold = style_spec.get("title_bold")
    author_bold = style_spec.get("author_bold")
    title_align_str = style_spec.get("title_align")
    author_align_str = style_spec.get("author_align")
    # 字符串→WD_ALIGN_PARAGRAPH枚举映射
    _align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    title_align = _align_map.get(title_align_str) if title_align_str else None
    author_align = _align_map.get(author_align_str) if author_align_str else None
    section_spacing = style_spec.get("section_spacing", {})

    def _set_direct_indent(fmt, first_line_pt=None, left_pt=0, right_pt=0):
        if left_pt is not None:
            fmt.left_indent = Pt(left_pt)
        if right_pt is not None:
            fmt.right_indent = Pt(right_pt)
        if first_line_pt is not None:
            fmt.first_line_indent = Pt(first_line_pt)

    def _next_nonempty_is_heading(paragraphs, start_idx):
        for item in paragraphs[start_idx + 1:]:
            item_text = (item.text or "").strip()
            if not item_text:
                continue
            item_style = item.style.name if item.style else ""
            return (
                item_style.startswith("Heading")
                or re.match(r"^\d+(?:\.\d+)*\s+\S+", item_text) is not None
            )
        return False

    def _before_first_heading(paragraphs, para_idx):
        for item in paragraphs[:para_idx]:
            item_text = (item.text or "").strip()
            item_style = item.style.name if item.style else ""
            if item_style.startswith("Heading") or re.match(r"^\d+(?:\.\d+)*\s+\S+", item_text):
                return False
        return True

    def _looks_like_abstract(text):
        return re.match(r"^(abstract|摘要)\b[.:：.]?\s+", text, flags=re.I) is not None

    # Heading间距: 从CLS提取的section/subsection/subsubsection参数
    heading_spacing = {
        "Heading 1": section_spacing.get("section", {}),
        "Heading 2": section_spacing.get("subsection", {}),
        "Heading 3": section_spacing.get("subsubsection", {}),
    }

    for style_name, en_font, size, bold in [
        ("Normal", main_font, body_size, False),
        ("Body Text", main_font, body_size, False),
        ("Caption", main_font, small_size, False),
        ("Bibliography", main_font, small_size, False),
        ("Heading 1", heading_font, body_size, style_spec["heading_bold"]),
        ("Heading 2", heading_font, body_size, style_spec["heading_bold"]),
        ("Heading 3", heading_font, body_size, style_spec["heading_bold"]),
    ]:
        style = _get_doc_style(doc, style_name)
        if style is not None:
            _set_style_fonts(style, en_font, cjk_font, size=size, bold=bold)
            sp = heading_spacing.get(style_name, {})
            style.paragraph_format.space_before = Pt(sp.get("before_pt", 0))
            style.paragraph_format.space_after = Pt(sp.get("after_pt", 0))
            style.paragraph_format.line_spacing = line_spacing
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            if style_name in ("Normal", "Body Text"):
                style.paragraph_format.first_line_indent = Pt(0)
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif style_name.startswith("Heading") or style_name in ("Caption", "Bibliography"):
                style.paragraph_format.first_line_indent = Pt(0)

    for style_name in ("First Paragraph", "Compact", "Plain Text"):
        style = _get_doc_style(doc, style_name)
        if style is not None:
            _set_style_fonts(style, main_font, cjk_font, size=body_size, bold=False)
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = line_spacing
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # Title/Author样式：从CLS模板动态提取字号+对齐方式
    title_before_pt = style_spec.get("title_before_pt", 0)
    title_style = _get_doc_style(doc, "Title")
    if title_style:
        title_font = cjk_font or main_font
        _set_style_fonts(title_style, title_font, cjk_font, size=title_size, bold=title_bold)
        title_style.paragraph_format.first_line_indent = Pt(0)
        if title_align is not None:
            title_style.paragraph_format.alignment = title_align  # CLS: \raggedright → LEFT
    author_style = _get_doc_style(doc, "Author")
    if author_style:
        _set_style_fonts(author_style, main_font, cjk_font, size=author_size, bold=author_bold)
        author_style.paragraph_format.first_line_indent = Pt(0)
        if author_align is not None:
            author_style.paragraph_format.alignment = author_align  # CLS: \flushleft → LEFT

    references_active = False
    first_title_done = False
    no_indent_styles = {"Title", "Author", "Date", "Abstract Title", "Subtitle"}
    paragraphs = list(doc.paragraphs)
    for para_idx, para in enumerate(paragraphs):
        if _is_empty_section_break_para(para._element):
            _compact_section_break_para(para._element)
            continue
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")
        is_caption = _is_caption_paragraph(para)
        is_abstract = style_name == "Abstract Title" or (
            _before_first_heading(paragraphs, para_idx) and _looks_like_abstract(text)
        )
        if is_abstract and style_name != "Abstract Title":
            try:
                para.style = doc.styles["Abstract Title"]
                style_name = "Abstract Title"
            except Exception:
                pass
        is_non_body = style_name in no_indent_styles
        is_correspondence = text.startswith("Correspondence:")
        is_formula_para = False
        has_omath_para = False
        try:
            for _ in para._element.iter(qn('m:oMath')):
                is_formula_para = True
                break
            for _ in para._element.iter(qn('m:oMathPara')):
                has_omath_para = True
                break
        except Exception:
            pass
        is_display_formula_para = is_formula_para and (
            has_omath_para
            or (
                style_name in ("", "Normal")
                and re.fullmatch(r"\(?[\w.\-]+\)?", text or "") is not None
            )
        )
        if text.lower() == "references":
            references_active = True
            is_heading = True
        is_bibliography = references_active and not is_heading and bool(text)

        # 对齐方式（从CLS模板动态提取）
        if style_name == "Title":
            if title_align is not None:
                para.alignment = title_align  # CLS: \raggedright → LEFT
        elif style_name == "Author":
            if author_align is not None:
                para.alignment = author_align  # CLS: \flushleft → LEFT
        elif is_abstract:
            # Abstract标题+正文在同一行，两端对齐，无缩进
            abstract_align = _align_map.get(style_spec.get("abstract_align") or "JUSTIFY")
            para.alignment = abstract_align or WD_ALIGN_PARAGRAPH.JUSTIFY
        elif not (is_heading or is_caption or is_non_body or is_display_formula_para or is_correspondence):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        fmt = para.paragraph_format
        fmt.line_spacing = abstract_line_spacing if is_abstract else line_spacing
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        # 标题、图例、参考文献、非正文、公式段落、Correspondence均不设首行缩进
        if is_heading or is_caption or is_bibliography or is_non_body or is_display_formula_para or is_correspondence:
            _set_direct_indent(fmt, 0)
        else:
            _set_direct_indent(fmt, parindent)
        if is_abstract:
            _set_direct_indent(
                fmt,
                0,
                style_spec.get("abstract_left_indent_pt") or 0,
                style_spec.get("abstract_right_indent_pt") or 0,
            )

        # Section间距: 标题段落使用CLS提取的beforeskip/afterskip
        if is_heading:
            sp_key = {"Heading 1": "section", "Heading 2": "subsection", "Heading 3": "subsubsection"}.get(style_name, "section")
            sp = section_spacing.get(sp_key, {})
            fmt.space_before = Pt(sp.get("before_pt", 0))
            fmt.space_after = Pt(sp.get("after_pt", 0))
        elif is_caption:
            before_pt = caption_spacing.get("above_pt")
            after_pt = caption_spacing.get("below_pt")
            if text.startswith(("Figure", "Fig.")) and caption_spacing.get("textfloat_pt") is not None:
                after_pt = max(after_pt or 0, caption_spacing["textfloat_pt"] * 0.35)
            fmt.space_before = Pt(before_pt or 0)
            fmt.space_after = Pt(after_pt or 0)
        elif is_abstract:
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(
                style_spec.get("abstract_after_pt", abstract_baseline)
                if _next_nonempty_is_heading(paragraphs, para_idx) else 0
            )
        else:
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)

        if style_name == "Title" and title_before_pt and not first_title_done:
            fmt.space_before = Pt(title_before_pt)
            first_title_done = True

        # 悬挂缩进: 参考文献区域设置 left=720, hanging=360 (Word XML层)
        if is_bibliography:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:left'), '720')
            ind.set(qn('w:hanging'), '360')

        run_font = heading_font if is_heading else main_font
        run_size = small_size if (is_caption or is_bibliography) else body_size
        if is_abstract:
            run_size = abstract_size
        # 标题/作者使用从CLS动态提取的字号
        if style_name == "Title":
            run_font = cjk_font or main_font
            run_size = title_size
        elif style_name == "Author":
            run_size = author_size
        for run in para.runs:
            _set_font(run, cjk_font, run_font, size=run_size)
            if is_heading:
                run.font.bold = style_spec["heading_bold"]
                # 标题颜色从CLS提取(heading_color)，未指定兜底黑(不继承Pandoc蓝)
                run.font.color.rgb = RGBColor(*style_spec.get("heading_color") or (0, 0, 0))
            elif is_bibliography:
                run.font.bold = False
            elif style_name == "Title":
                if title_bold is not None:
                    run.font.bold = title_bold
                if style_spec.get("heading_color"):
                    run.font.color.rgb = RGBColor(*style_spec["heading_color"])
            elif style_name == "Author":
                if author_bold is not None:
                    run.font.bold = author_bold
            elif style_name == "Abstract Title":
                # 仅"Abstract."标签run加粗(文本检测)，正文run不加粗（支持多段摘要不合并）
                if re.match(r"^(abstract|摘要)\b", run.text or "", flags=re.I):
                    run.font.bold = True
                    if style_spec.get("heading_color"):
                        run.font.color.rgb = RGBColor(*style_spec["heading_color"])
                else:
                    run.font.bold = False
            elif is_correspondence and re.match(r"^Correspondence:", run.text or ""):
                run.font.bold = True
                if style_spec.get("heading_color"):
                    run.font.color.rgb = RGBColor(*style_spec["heading_color"])

    for tbl in doc.tables:
        if _is_layout_floating_table_elem(tbl._element):
            _clear_table_borders(tbl._element)
            continue
        layout_table = (style_spec.get("layout_spec") or {}).get("table", {})
        tbl_font_size = _table_font_size_from_specs(
            style_spec, style_spec.get("layout_spec"), small_size)
        table_align = str(layout_table.get("alignment") or "").lower()
        table_jc = {
            "center": "center",
            "centering": "center",
            "left": "left",
            "right": "right",
        }.get(table_align)
        from docx.oxml import OxmlElement
        if table_jc:
            tblPr = tbl._tbl.find(qn('w:tblPr'))
            if tblPr is not None:
                jc = tblPr.find(qn('w:jc'))
                if jc is None:
                    jc = OxmlElement('w:jc')
                    _append_tbl_pr_child(tblPr, jc)
                jc.set(qn('w:val'), table_jc)
        for row_idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    if table_jc == "center":
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif table_jc == "left":
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif table_jc == "right":
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in para.runs:
                        _set_font(run, cjk_font, main_font, size=tbl_font_size)
        _keep_table_on_one_page(tbl._tbl)
    _enforce_caption_table_keep_together(doc)
    # Do not synthesize extra horizontal rules here. TikZ/PDF-derived
    # tables already carry exact line segments from \draw commands.


def _merge_page_geometry_with_layout(style_spec, layout_spec=None):
    layout_geo = (layout_spec or style_spec.get("layout_spec") or {}).get("page_geometry") or {}
    page_geo = style_spec.get("page_geometry") or {}
    if not layout_geo:
        return page_geo
    if not page_geo:
        page_geo = dict(layout_geo)
    else:
        page_geo = dict(page_geo)
        for key in ("column_count", "column_sep_mm"):
            if page_geo.get(key) in (None, "") and layout_geo.get(key) not in (None, ""):
                page_geo[key] = layout_geo[key]
    style_spec["page_geometry"] = page_geo
    style_spec["column_count"] = page_geo.get(
        "column_count", style_spec.get("column_count", 1))
    style_spec["column_sep_mm"] = page_geo.get(
        "column_sep_mm", style_spec.get("column_sep_mm"))
    return page_geo


def _merge_font_specs_with_layout(style_spec, layout_spec=None):
    layout = layout_spec or style_spec.get("layout_spec") or {}
    document = layout.get("document", {}) if isinstance(layout, dict) else {}
    fontspec = layout.get("fontspec_config", {}) if isinstance(layout, dict) else {}
    candidates = {
        "main_font": fontspec.get("main_font") or document.get("main_font"),
        "sans_font": fontspec.get("sans_font") or document.get("sans_font"),
        "mono_font": fontspec.get("mono_font") or document.get("mono_font"),
        "cjk_font": fontspec.get("setCJKmainfont") or document.get("cjk_main_font"),
    }
    for key, value in candidates.items():
        if not style_spec.get(key) and value:
            style_spec[key] = value
    if style_spec.get("heading_uses_sans"):
        style_spec["heading_font"] = style_spec.get("sans_font")
    elif not style_spec.get("heading_font"):
        style_spec["heading_font"] = style_spec.get("main_font")


def _merge_template_layout(style_spec, layout_spec):
    if layout_spec:
        style_spec["layout_spec"] = layout_spec
    _merge_page_geometry_with_layout(style_spec)
    _merge_font_specs_with_layout(style_spec)
    return style_spec


def _table_font_size_from_specs(style_spec=None, layout_spec=None, default=9):
    """Resolve editable Word table body size from extracted template specs."""
    spec_layout = layout_spec or (style_spec or {}).get("layout_spec") or {}
    table_spec = spec_layout.get("table", {})
    size = table_spec.get("body_size") or (style_spec or {}).get("small_size") or default
    if isinstance(size, str):
        match = re.match(r"([\d.]+)\s*pt$", size.strip())
        if match:
            return float(match.group(1))
        return _latex_size_token_to_pt(size, default)
    try:
        return float(size)
    except (TypeError, ValueError):
        return float(default)


def _caption_style_from_specs(style_spec=None, layout_spec=None):
    """Resolve caption formatting from extracted template specs."""
    style_spec = style_spec or {}
    spec_layout = layout_spec or style_spec.get("layout_spec") or {}
    caption_spec = spec_layout.get("caption", {}) if isinstance(spec_layout, dict) else {}
    body_size = float(style_spec.get("body_size") or 10)
    default_size = style_spec.get("small_size") or body_size
    size = (
        caption_spec.get("font_size_pt")
        or caption_spec.get("size_pt")
        or caption_spec.get("font_size")
        or default_size
    )
    if isinstance(size, str):
        match = re.match(r"([\d.]+)\s*pt$", size.strip())
        size_pt = float(match.group(1)) if match else _latex_size_token_to_pt(size, default_size)
    else:
        size_pt = float(size)
    align = str(caption_spec.get("alignment") or caption_spec.get("justification") or "left").lower()
    align = {
        "center": "center",
        "centering": "center",
        "raggedright": "left",
        "left": "left",
        "right": "right",
        "justified": "both",
        "justify": "both",
    }.get(align, align)
    cls_spacing = style_spec.get("caption_spacing") or {}
    before_twips = _latex_length_to_twips(caption_spec.get("above_skip"), body_size)
    after_twips = _latex_length_to_twips(caption_spec.get("below_skip"), body_size)
    if before_twips is None and cls_spacing.get("above_pt") is not None:
        before_twips = int(round(cls_spacing["above_pt"] * 20))
    if after_twips is None and cls_spacing.get("below_pt") is not None:
        after_twips = int(round(cls_spacing["below_pt"] * 20))
    if after_twips is None and cls_spacing.get("textfloat_pt") is not None:
        after_twips = int(round(cls_spacing["textfloat_pt"] * 0.35 * 20))
    return {
        "font_size_pt": size_pt,
        "font_half_points": str(int(round(size_pt * 2))),
        "alignment": align,
        "label_bold": caption_spec.get("label_weight", "bold") != "normal",
        "space_before_twips": before_twips,
        "space_after_twips": after_twips,
    }


def apply_template_word_layout(docx_path, tex_path=None, config_mode=None,
                               layout_spec=None):
    """Apply template-derived page geometry and columns without changing styles."""
    from docx import Document

    style_spec = _load_template_word_style(tex_path, config_mode=config_mode)
    if not style_spec:
        print("  [template layout] skipped: no template layout source found")
        return docx_path
    _merge_template_layout(style_spec, layout_spec)
    doc = Document(docx_path)
    page_geo = style_spec.get("page_geometry")
    if page_geo:
        _apply_page_setup_to_doc(
            doc, page_geo, footer_dims=style_spec.get("footer_dims"))
    else:
        _apply_column_setup_to_doc(
            doc,
            style_spec.get("column_count", 1),
            style_spec.get("column_sep_mm"),
        )
    doc.save(docx_path)
    print(
        "  [template layout] "
        f"cols={style_spec.get('column_count', 1)}, "
        f"mode={_word_config_mode(config_mode, []) or 'template-default'}"
    )
    return docx_path


def apply_template_word_styles(docx_path, tex_path=None, config_mode=None, layout_spec=None):
    """Apply template-derived fonts, sizes, and spacing to a converted DOCX.

    config_mode: 模板配置模式名 ('manuscript', 'final', 'discussions')
    layout_spec: 排版规格字典 (从 *_template_spec.json 加载)
    """
    from docx import Document

    style_spec = _load_template_word_style(tex_path, config_mode=config_mode)
    if not style_spec:
        print("  [template styles] skipped: no template style source found")
        return docx_path
    _merge_template_layout(style_spec, layout_spec)
    doc = Document(docx_path)
    _apply_template_styles_to_doc(doc, style_spec)
    _clean_formula_placeholders(doc)
    _prune_empty_math_placeholders(doc)
    _fix_overline_acc_to_bar(doc)
    _fix_empty_nary_body(doc)
    _clean_leading_table_bullets(doc)
    _remove_blank_paragraphs_around_tables(doc)
    _remove_blank_paragraphs_around_images(doc)
    _fix_table_continuation_indent(doc)
    doc.save(docx_path)
    print(
        "  [template styles] "
        f"font={style_spec['main_font']}, cjk={style_spec['cjk_font']}, "
        f"body={style_spec['body_size']}pt, small={style_spec['small_size']}pt, "
        f"title_align={style_spec.get('title_align', 'CENTER')}, "
        f"author_align={style_spec.get('author_align', 'CENTER')}, "
        f"line_spacing={style_spec['line_spacing']:.3g}, "
        f"parindent={style_spec.get('applied_parindent', style_spec.get('parindent', style_spec['body_size'])):.3g}pt"
    )
    return docx_path


def _effective_config_mode_from_tex(tex_path):
    """从 tex 的 documentclass 选项推断有效配置模式。

    copernicus: manuscript 选项 -> 单栏, final -> 双栏(对齐/颜色不同)。
    """
    if not tex_path:
        return None
    try:
        content = _read_text_file(tex_path)
    except Exception:
        return None
    if not content:
        return None
    m = re.search(r"\\documentclass(?:\[([^\]]*)\])?", content)
    if not m or not m.group(1):
        return None
    opts = [o.strip().lower() for o in m.group(1).split(",") if o.strip()]
    return _word_config_mode(None, opts)


def restore_front_matter_from_tex(docx_path, tex_path):
    """Restore front matter commands from the original .tex file into Word.

    Some journal commands are consumed by the class or removed during Pandoc
    preprocessing, so recover them from the compiled source.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not tex_path or not os.path.exists(tex_path):
        return

    tex_content = _read_text_file(tex_path)
    if not tex_content:
        return

    doc = Document(docx_path)
    cls_content = _read_cls_from_tex(tex_path)

    abstract_label, abstract_text = _extract_tex_abstract(tex_content, cls_content)
    if _insert_abstract_before_body(doc, abstract_label, abstract_text):
        print(f"  [front matter] restored Abstract: {abstract_text[:50]}...")

    corr_m = re.search(r"\\correspondence\{([^}]+)\}", tex_content)
    corr_text = corr_m.group(1).strip() if corr_m else ""

    if corr_text:
        for para in doc.paragraphs:
            style_name = (para.style.name or "") if para.style else ""
            if style_name in ("Author", "author"):
                corr_para = OxmlElement("w:p")
                pPr = OxmlElement("w:pPr")
                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), "BodyText")
                pPr.append(pStyle)
                # 对齐方式按 config_mode: final→居中, 其他→左对齐(与 CLS if@stage@final 一致)
                corr_mode = _effective_config_mode_from_tex(tex_path)
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center" if corr_mode == "final" else "left")
                pPr.append(jc)
                # 无首行缩进(前置内容不应继承正文缩进)
                corr_ind = OxmlElement("w:ind")
                corr_ind.set(qn("w:left"), "0")
                corr_ind.set(qn("w:right"), "0")
                corr_ind.set(qn("w:firstLine"), "0")
                pPr.append(corr_ind)
                corr_para.append(pPr)

                run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "18")  # 9pt
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "18")
                rPr.append(szCs)
                run.append(rPr)
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = f"Correspondence: {corr_text}"
                run.append(t)
                corr_para.append(run)

                para._element.addnext(corr_para)
                print(f"  [front matter] restored Correspondence: {corr_text[:50]}...")
                break

    _ensure_front_matter_single_column(doc)
    doc.save(docx_path)


def _read_text_file(path):
    """Read a text file with encoding fallback."""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, FileNotFoundError):
            continue
    return None


def _read_cls_from_tex(tex_path):
    """从tex文件路径推断并读取对应的.cls文件内容"""
    tex_path = Path(tex_path)
    tex_dir = tex_path.parent
    tex_content = _read_text_file(tex_path)
    if not tex_content:
        return None
    cls_m = re.search(r"\\documentclass(?:\[[^\]]*\])?\{([^}]+)\}", tex_content)
    cls_name = cls_m.group(1).strip() if cls_m else ""
    if cls_name:
        for candidate in [tex_dir / f"{cls_name}.cls", tex_dir / f"{cls_name}_paper" / f"{cls_name}.cls"]:
            if candidate.exists():
                return _read_text_file(candidate)
    # glob fallback
    for p in tex_dir.glob("*_paper/*.cls"):
        return _read_text_file(p)
    return None


def _balanced_group(text, brace_pos):
    if brace_pos < 0 or brace_pos >= len(text) or text[brace_pos] != "{":
        return ""
    depth = 0
    for idx in range(brace_pos, len(text)):
        ch = text[idx]
        prev = text[idx - 1] if idx else ""
        if ch == "{" and prev != "\\":
            depth += 1
        elif ch == "}" and prev != "\\":
            depth -= 1
            if depth == 0:
                return text[brace_pos + 1:idx]
    return ""


def _extract_tex_abstract(tex_content, cls_content=None):
    cmd_match = re.search(r"\\abstract(?:\[([^\]]*)\])?\s*\{", tex_content)
    if cmd_match:
        body = _balanced_group(tex_content, cmd_match.end() - 1)
        label = (cmd_match.group(1) or "").strip()
        return label, _clean_latex_text(body).strip()

    env_match = re.search(
        r"\\begin\{abstract\*?\}(.*?)\\end\{abstract\*?\}",
        tex_content,
        flags=re.DOTALL | re.IGNORECASE,
    )
    if not env_match:
        return "", ""
    label = "Abstract"
    if cls_content:
        name_match = re.search(r"\\(?:renewcommand|def)\\abstractname\{([^}]+)\}", cls_content)
        if name_match:
            label = _clean_latex_text(name_match.group(1)).strip() or label
    return label, _clean_latex_text(env_match.group(1)).strip()


def _doc_has_abstract(doc, abstract_text):
    front = []
    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style else ""
        text = para.text.strip()
        if style.startswith("Heading") or re.match(r"^\d+(?:\.\d+)*\s+\S+", text):
            break
        front.append(text)
    all_text = re.sub(r"\s+", " ", "\n".join(front)).strip()
    if re.search(r"\babstract\b|摘要", all_text, flags=re.IGNORECASE):
        return True
    sample = re.sub(r"\s+", " ", abstract_text).strip()[:40]
    return bool(sample and sample in all_text)


def _insert_abstract_before_body(doc, label, abstract_text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not abstract_text or _doc_has_abstract(doc, abstract_text):
        return False
    label = (label or "Abstract").strip()
    para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "AbstractTitle")
    pPr.append(pStyle)
    para.append(pPr)

    label_run = OxmlElement("w:r")
    label_pr = OxmlElement("w:rPr")
    label_pr.append(OxmlElement("w:b"))
    label_run.append(label_pr)
    label_text = OxmlElement("w:t")
    label_text.set(qn("xml:space"), "preserve")
    label_text.text = f"{label}. "
    label_run.append(label_text)
    para.append(label_run)

    body_run = OxmlElement("w:r")
    body_text = OxmlElement("w:t")
    body_text.set(qn("xml:space"), "preserve")
    body_text.text = abstract_text
    body_run.append(body_text)
    para.append(body_run)

    target = None
    for item in doc.paragraphs:
        style = (item.style.name or "") if item.style else ""
        text = item.text.strip()
        if style.startswith("Heading") or re.match(r"^\d+(?:\.\d+)*\s+\S+", text):
            target = item
            break
    if target is not None:
        _heading_anchor_element(target._element).addprevious(para)
    else:
        doc.element.body.append(para)
    return True


def postprocess_docx(docx_path, display_formula_data, tables, cite_map=None, cite_style='apa', bbl_path=None, layout_spec=None, style_spec=None, tex_path=None):
    """后处理Word文件: display公式OMML插入 + 表格边框重建 + 字体设置 + 标题样式修复 + 文献交叉引用"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)

    # 1. 用 latex_to_omml skill 替换 display 公式占位符
    _insert_display_formulas(doc, display_formula_data, layout_spec=style_spec)

    # 2. 修复标题样式
    # Keep heading styles from Pandoc/reference-doc instead of hardcoding them.

    # 3. 文献交叉引用
    if cite_map and _HAS_CROSSREF:
        if bbl_path and os.path.exists(bbl_path):
            bbl_items = parse_bbl_items(bbl_path)
            build_references_section(doc, bbl_items, cite_map=cite_map, cite_style=cite_style)
        insert_bib_cross_references(doc, cite_map, cite_style=cite_style)

    # 4. 公式统计
    omml_count = 0
    for elem in doc.element.body.iter(qn('m:oMath')):
        omml_count += 1
    print(f'  Found {omml_count} OMML formulas in Word')

    # 5. 表格边框重建
    doc_tables = doc.tables
    print(f'  Found {len(doc_tables)} tables in Word')
    print(f'  Extracted {len(tables)} tables from LaTeX')

    real_table_idx = 0
    for tbl in doc_tables:
        if len(tbl.rows) <= 1 and len(tbl.columns) <= 1:
            continue
        tex_tbl = None
        if real_table_idx < len(tables):
            tex_tbl = tables[real_table_idx]
        _add_table_borders(tbl, tex_tbl, layout_spec=layout_spec)
        # Do not synthesize extra horizontal rules for PDF-exact table output.
        real_table_idx += 1
    print(f'  Added borders to {real_table_idx} tables')
    # 6. Do not override document fonts; keep template/Pandoc styles.

    # 7. 清理公式中的〖〗占位符残留 + 空OMML槽 + 表格numPr
    # 从tex_path读取CLS内容用于Abstract标签格式提取
    cls_content = None
    if tex_path:
        cls_content = _read_cls_from_tex(tex_path)
    _normalize_abstract_paragraphs(doc, cls_content=cls_content)
    _clean_formula_placeholders(doc)
    _prune_empty_math_placeholders(doc)
    _fix_overline_acc_to_bar(doc)
    _fix_empty_nary_body(doc)
    _clean_leading_table_bullets(doc)
    _remove_blank_paragraphs_around_tables(doc)
    _remove_blank_paragraphs_around_images(doc)
    _fix_table_continuation_indent(doc)

    # 8. 恢复公式编号（从display_formula_data中的eq_num）
    doc.save(docx_path)
    return docx_path


def _resolve_tex_asset(path_value, tex_dir):
    path = Path(path_value)
    if path.is_absolute() and path.exists():
        return path
    candidate = Path(tex_dir) / path
    if candidate.exists():
        return candidate
    return path


def _latex_image_width_to_docx(width_spec, section, image_path=None, compiled_layout=None):
    from docx.shared import Cm, Inches
    from docx.oxml.ns import qn

    compiled_layout = compiled_layout or {}
    to_emu = lambda twips: int(twips * 635)
    def _length_emu(value):
        try:
            return int(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    if compiled_layout.get("textwidth"):
        max_width = to_emu(compiled_layout["textwidth"])
    else:
        page_width = _length_emu(getattr(section, "page_width", None))
        left_margin = _length_emu(getattr(section, "left_margin", None))
        right_margin = _length_emu(getattr(section, "right_margin", None))
        if all(value is not None for value in (page_width, left_margin, right_margin)):
            max_width = max(page_width - left_margin - right_margin, int(Inches(1.0)))
        else:
            max_width = int(Inches(6.5))
    page_height = _length_emu(getattr(section, "page_height", None))
    top_margin = _length_emu(getattr(section, "top_margin", None))
    bottom_margin = _length_emu(getattr(section, "bottom_margin", None))
    if all(value is not None for value in (page_height, top_margin, bottom_margin)):
        text_height = max(page_height - top_margin - bottom_margin, int(Inches(1.0)))
    else:
        text_height = int(Inches(9.0))

    # 检测section是否为多栏，从sectPr读取cols
    col_count = 1
    try:
        sectPr = section._sectPr
        cols_elem = sectPr.find(qn('w:cols'))
        if cols_elem is not None:
            num_attr = cols_elem.get(qn('w:num'))
            if num_attr:
                col_count = int(num_attr)
    except Exception:
        pass

    # 栏间距（默认360 twips ≈ 6.35mm）
    col_space = 360 * 635
    try:
        sectPr = section._sectPr
        cols_elem = sectPr.find(qn('w:cols'))
        if cols_elem is not None:
            space_attr = cols_elem.get(qn('w:space'))
            if space_attr:
                col_space = int(space_attr) * 635
    except Exception:
        pass

    # 单栏宽度 = (max_width - (col_count-1)*col_space) / col_count
    if compiled_layout.get("columnwidth"):
        col_width = to_emu(compiled_layout["columnwidth"])
    else:
        col_width = (
            max((max_width - (col_count - 1) * col_space) // col_count, int(Inches(1.0)))
            if col_count > 1 else max_width
        )

    if not width_spec:
        # 无宽度指定时: 单栏用max_width, 双栏用col_width(半栏)
        width = col_width if col_count > 1 else max_width
    else:
        spec = width_spec.replace(" ", "")
        match = re.search(r"width=([0-9.]+)\\textwidth", spec)
        if match:
            width = int(max_width * float(match.group(1)))
        elif "width=\\textwidth" in spec:
            width = max_width
        elif "width=\\columnwidth" in spec:
            # \columnwidth = 单栏宽度
            width = col_width
        elif re.search(r"width=([0-9.]+)\\columnwidth", spec):
            match = re.search(r"width=([0-9.]+)\\columnwidth", spec)
            width = int(col_width * float(match.group(1)))
        else:
            match = re.search(r"width=([0-9.]+)cm", spec)
            width = Cm(float(match.group(1))) if match else (col_width if col_count > 1 else max_width)

    # 限制图片高度不超过页面文本区域的75%（匹配PDF实际渲染比例）
    max_height = int(text_height * 0.75)
    if image_path:
        try:
            from PIL import Image
            with Image.open(image_path) as img:
                px_w, px_h = img.size
            if px_w and px_h:
                rendered_height = int(width * (px_h / px_w))
                if rendered_height > max_height:
                    width = int(max_height * (px_w / px_h))
        except Exception as exc:
            print(f"  Warning: cannot inspect image size for {image_path}: {exc}")
    return width


def _append_run_size_props(rPr, half_points):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not half_points:
        return
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(half_points))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(half_points))
    rPr.append(szCs)


def _caption_para_element(kind, number, caption, legend="", caption_style=None, keep_next=False):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    caption_style = caption_style or _caption_style_from_specs()
    cap_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    _append_caption_paragraph_props(
        pPr, caption_style=caption_style, keep_next=keep_next)
    cap_para.append(pPr)

    bold_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if caption_style.get("label_bold"):
        rPr.append(OxmlElement("w:b"))
    _append_run_size_props(rPr, caption_style.get("font_half_points"))
    bold_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = f"{kind} {number}. "
    bold_run.append(t)
    cap_para.append(bold_run)

    normal_run = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    _append_run_size_props(rPr2, caption_style.get("font_half_points"))
    normal_run.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.set(qn("xml:space"), "preserve")
    extra = f" {legend}" if legend else ""
    t2.text = f"{caption}{extra}"
    normal_run.append(t2)
    cap_para.append(normal_run)
    return cap_para


def embed_images_in_docx(docx_path, images, tex_dir, layout_spec=None, style_spec=None):
    """Insert LaTeX figures into Word at [FIGURE_N] placeholders."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document(docx_path)
    compiled_layout = _compiled_layout_twips(tex_dir)
    caption_style = _caption_style_from_specs(style_spec, layout_spec)
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    img_paras = []
    for para in doc.paragraphs:
        if para._element.findall('.//w:drawing', nsmap):
            img_paras.append(para)

    placeholder_map = {}
    for para in doc.paragraphs:
        match = re.search(r'\[FIGURE_(\d+)\]', para.text)
        if match:
            placeholder_map[int(match.group(1))] = para

    def _fallback_figure_anchor():
        previous = None
        for para in doc.paragraphs:
            text = (para.text or "").strip().lower()
            if text == "references" or text.startswith("references\n"):
                return previous or para
            previous = para
        return doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()

    print(
        f'  Found {len(img_paras)} image paragraphs in Word, '
        f'{len(placeholder_map)} figure placeholders, {len(images)} images from LaTeX'
    )

    inserted = 0
    failed = 0
    body = doc.element.body
    for i, img_info in enumerate(images):
        target_para = placeholder_map.get(i)
        if target_para is None and i < len(img_paras):
            target_para = img_paras[i]
        if target_para is None:
            target_para = _fallback_figure_anchor()

        image_path = _resolve_tex_asset(img_info.get("path", ""), tex_dir)
        if not image_path.exists():
            failed += 1
            print(f"  Warning: image not found: {image_path}")
            continue

        image_para = doc.add_paragraph()
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.paragraph_format.space_before = Pt(0)
        image_para.paragraph_format.space_after = Pt(0)
        if img_info.get('caption') or img_info.get('caption_full'):
            image_para.paragraph_format.keep_with_next = True
        try:
            width = _latex_image_width_to_docx(
                img_info.get("width", ""), doc.sections[0],
                image_path=image_path, compiled_layout=compiled_layout)
            image_para.add_run().add_picture(str(image_path), width=width)
        except Exception as exc:
            failed += 1
            body.remove(image_para._element)
            print(f"  Warning: failed to insert image {image_path}: {exc}")
            continue

        image_elem = image_para._element
        body.remove(image_elem)
        caption = _clean_latex_text(img_info.get('caption', ''))
        legend = _clean_latex_text(img_info.get('legend', ''))
        number = img_info.get("number") or str(i + 1)
        cap_elem = (
            _caption_para_element("Figure", number, caption, legend, caption_style=caption_style)
            if caption else None
        )

        is_full_width = bool(img_info.get("is_full_width")) or "\\textwidth" in (img_info.get("width") or "")
        block_elems = [image_elem, cap_elem]
        if is_full_width:
            block_elems = _full_width_block_elements(doc, block_elems)
        _keep_float_block_together(block_elems)
        insert_anchor = target_para._element
        if is_full_width:
            required_mm = _image_required_space_mm(img_info, image_path, width, style_spec)
            insert_anchor = _full_width_float_insert_anchor(
                insert_anchor, doc=doc, style_spec=style_spec,
                required_space_mm=required_mm,
                pdf_guidance=img_info.get("pdf_guidance"))
        _insert_after_in_order(insert_anchor, block_elems)
        if is_full_width:
            _apply_float_page_top(
                insert_anchor, block_elems,
                pdf_guidance=img_info.get("pdf_guidance"))
        # 如果段落仅含占位符 → 删除整段；否则只清理占位符文本（保留表格等其他内容）
        remaining = re.sub(r'\[FIGURE_\d+\]', '', target_para.text or '').strip()
        if remaining:
            # 段落还有其他内容(如合并的TIKZ_TABLE占位符)，只清除FIGURE占位符
            for run in target_para.runs:
                if run.text:
                    run.text = re.sub(r'\[FIGURE_\d+\]', '', run.text)
        else:
            target_para._element.getparent().remove(target_para._element)
        inserted += 1

    _remove_blank_paragraphs_around_images(doc)
    doc.save(docx_path)
    print(f'  Inserted images: {inserted}, failed: {failed}')
    return docx_path


def add_tikz_table_to_docx(docx_path, tikz_tables, tex_dir,
                           layout_spec=None, style_spec=None):
    """Replace TikZ table placeholders in Word using lossless metadata when available."""
    from copy import deepcopy
    from pathlib import Path
    from docx import Document
    from docx.table import Table
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    from _tex_extraction import parse_tikz_table

    table_skill_dir = Path(__file__).resolve().parent.parent / "table-lossless-extract"
    if str(table_skill_dir) not in sys.path:
        sys.path.insert(0, str(table_skill_dir))
    try:
        from _tikz_parser import tikz_to_json
        from gen_table_from_json import generate_docx as _generate_table_docx
    except Exception:
        tikz_to_json = None
        _generate_table_docx = None

    def _max_table_width_twips(full_width=False):
        compiled = _compiled_layout_twips(tex_dir)
        compiled_width = compiled.get("textwidth" if full_width else "columnwidth", 0)
        if compiled_width:
            return compiled_width
        if doc.sections:
            section_width = _section_content_width_twips(
                doc.sections[0], single_column=not full_width)
            if section_width:
                return section_width
        try:
            cls_files = list(Path(tex_dir).glob("*.cls"))
            if not cls_files:
                cls_files = list(Path(tex_dir).glob("*_paper/*.cls"))
            if cls_files:
                from shared.template_config import get_page_geometry_for_mode
                content = cls_files[0].read_text(encoding="utf-8", errors="ignore")
                geo = get_page_geometry_for_mode(content) or {}
                textwidth_mm = geo.get("textwidth_mm")
                if textwidth_mm:
                    return int(float(textwidth_mm) * 56.6929)
        except Exception:
            pass
        if doc.sections:
            section = doc.sections[0]
            dims = (section.page_width, section.left_margin, section.right_margin)
            if all(dim is not None for dim in dims):
                return int((dims[0] - dims[1] - dims[2]) / 635)
        return 0

    def _scale_table_json_to_doc_width(table_json, full_width=False):
        grid_cols = table_json.get("grid_cols") or table_json.get("grid_col") or []
        total_twips = sum(int(col.get("width_twips", 0) or 0) for col in grid_cols)
        if not grid_cols or total_twips <= 0:
            return
        max_twips = _max_table_width_twips(full_width=full_width)
        if max_twips <= 0:
            return
        scale = max_twips / total_twips
        for col in grid_cols:
            col["width_twips"] = max(120, int(int(col.get("width_twips", 0) or 0) * scale))

    def _lossless_table_element(tikz_body, table_idx, full_width=False):
        if tikz_to_json is None or _generate_table_docx is None:
            return None
        try:
            # The visible Word caption is inserted separately; keep structure meta,
            # but remove the internal caption-row hint to avoid duplicate captions.
            tikz_for_word = "\n".join(
                line for line in tikz_body.splitlines()
                if not re.match(r"\s*%\s*meta:(has_caption_row|caption_text)=", line)
            )
            table_json = tikz_to_json(tikz_for_word)
            if not table_json or not table_json.get("rows"):
                return None
            _scale_table_json_to_doc_width(table_json, full_width=full_width)
            tmp_path = Path(tex_dir) / f"_lossless_tikz_table_{table_idx}.docx"
            _generate_table_docx(table_json, str(tmp_path))
            tmp_doc = Document(str(tmp_path))
            if not tmp_doc.tables:
                return None
            return deepcopy(tmp_doc.tables[0]._element)
        except Exception as exc:
            print(f"  Warning: lossless TikZ table rebuild failed for table {table_idx + 1}: {exc}")
            return None

    def _caption_element(caption, table_idx, table_number=None):
        if not caption:
            return None
        caption = _clean_latex_text(caption)
        number = table_number or str(table_idx + 1)
        return _caption_para_element(
            "Table", number, caption, caption_style=caption_style, keep_next=True)

    doc = Document(docx_path)
    caption_style = _caption_style_from_specs(style_spec, layout_spec)
    fallback_font_size = _table_font_size_from_specs(
        style_spec, layout_spec, default=9)

    for i, tikz_tbl in enumerate(tikz_tables):
        caption = tikz_tbl.get("caption", "")
        table_number = tikz_tbl.get("number", "")
        tikz_body = tikz_tbl.get("tikz_body", "")
        placeholder = f"[TIKZ_TABLE_{i}]"
        placeholder_para = next((para for para in doc.paragraphs if placeholder in para.text), None)

        is_full_width = bool(tikz_tbl.get("is_full_width")) or "% meta:full_width=1" in tikz_body
        table_elem = _lossless_table_element(tikz_body, i, full_width=is_full_width)
        if table_elem is None:
            table_data = parse_tikz_table(tikz_body)
            if not table_data["headers"] and not table_data["rows"]:
                continue
            if table_data["rows"]:
                num_cols = max(
                    len(table_data["headers"]),
                    max((len(row) for row in table_data["rows"]), default=0),
                )
                num_rows = len(table_data["rows"]) + 1
            else:
                num_cols = len(table_data["headers"])
                num_rows = 1
            if num_cols == 0:
                continue
            table = doc.add_table(rows=num_rows, cols=num_cols)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            for j, header in enumerate(table_data["headers"]):
                if j < num_cols:
                    run = table.rows[0].cells[j].paragraphs[0].add_run(header)
                    run.bold = True
                    run.font.size = Pt(fallback_font_size)
            for row_idx, row_data in enumerate(table_data["rows"]):
                if row_idx + 1 < num_rows:
                    for j, cell_text in enumerate(row_data):
                        if j < num_cols:
                            run = table.rows[row_idx + 1].cells[j].paragraphs[0].add_run(cell_text)
                            run.font.size = Pt(fallback_font_size)
            _add_table_borders(table)
            table_elem = table._element

        _complete_table_bottom_line(Table(table_elem, doc))
        cap_para_elem = _caption_element(caption, i, table_number)
        if placeholder_para is not None:
            # 检查段落是否包含占位符以外的内容（合并后的情况）
            remaining = re.sub(r'\[TIKZ_TABLE_\d+\]', '', placeholder_para.text or '').strip()
            if remaining:
                # 段落被合并过，需要分割：前文 → 表格 → 后文
                from copy import deepcopy
                para_elem = placeholder_para._element
                body_elem = para_elem.getparent()
                all_runs = para_elem.findall(qn('w:r'))

                # 找到包含占位符的run
                ph_idx = None
                for idx, run in enumerate(all_runs):
                    run_text = ''.join(t.text or '' for t in run.findall(qn('w:t')))
                    if '[TIKZ_TABLE_' in run_text:
                        ph_idx = idx
                        break

                if ph_idx is not None:
                    # 创建"后文"段落（深拷贝后删除占位符run及之前的runs）
                    after_para = deepcopy(para_elem)
                    after_runs = after_para.findall(qn('w:r'))
                    for idx in range(ph_idx + 1):
                        after_para.remove(after_runs[idx])

                    # 从原文删除占位符run及其后的runs
                    for idx in range(ph_idx, len(all_runs)):
                        para_elem.remove(all_runs[idx])

                    # 插入: 原段落(前文) → caption → table → 后段落(后文)
                    # addnext 插在调用元素正后方，所以先加 table 再加 caption
                    # 才能使 caption 排在 table 前面
                    block_elems = [cap_para_elem, table_elem]
                    if is_full_width:
                        block_elems = _full_width_block_elements(doc, block_elems)
                    _keep_float_block_together(block_elems)

                    # 检查后段落是否有文本
                    after_text = ''.join(t.text or '' for t in after_para.iter(qn('w:t')))
                    if is_full_width and after_text.strip():
                        para_elem.addnext(after_para)
                        insert_anchor = _full_width_table_anchor(
                            after_para, doc, table_elem, style_spec,
                            pdf_guidance=tikz_tbl.get("pdf_guidance"))
                        last_inserted = _insert_after_in_order(insert_anchor, block_elems)
                        _apply_float_page_top(
                            insert_anchor, block_elems,
                            pdf_guidance=tikz_tbl.get("pdf_guidance"))
                    else:
                        insert_anchor = para_elem
                        if is_full_width:
                            insert_anchor = _full_width_table_anchor(
                                para_elem, doc, table_elem, style_spec,
                                pdf_guidance=tikz_tbl.get("pdf_guidance"))
                        last_inserted = _insert_after_in_order(insert_anchor, block_elems)
                        if is_full_width:
                            _apply_float_page_top(
                                insert_anchor, block_elems,
                                pdf_guidance=tikz_tbl.get("pdf_guidance"))
                    if (not is_full_width) and after_text.strip():
                        last_inserted.addnext(after_para)

                    # 检查原段落是否还有文本（前文为空则删除）
                    orig_text = ''.join(t.text or '' for t in para_elem.iter(qn('w:t')))
                    if not orig_text.strip():
                        body_elem.remove(para_elem)

                    print(f"  Split & replaced inline [{placeholder}] with TikZ table {i + 1}")
                else:
                    # 找不到具体run，fallback
                    block_elems = [cap_para_elem, table_elem]
                    if is_full_width:
                        block_elems = _full_width_block_elements(doc, block_elems)
                    _keep_float_block_together(block_elems)
                    insert_anchor = para_elem
                    if is_full_width:
                        insert_anchor = _full_width_table_anchor(
                            para_elem, doc, table_elem, style_spec,
                            pdf_guidance=tikz_tbl.get("pdf_guidance"))
                    _insert_after_in_order(insert_anchor, block_elems)
                    if is_full_width:
                        _apply_float_page_top(
                            insert_anchor, block_elems,
                            pdf_guidance=tikz_tbl.get("pdf_guidance"))
                    para_elem.getparent().remove(para_elem)
                    print(f"  Replaced [{placeholder}] with TikZ table {i + 1}")
            else:
                # 纯占位符段落，直接替换
                block_elems = [cap_para_elem, table_elem]
                if is_full_width:
                    block_elems = _full_width_block_elements(doc, block_elems)
                _keep_float_block_together(block_elems)
                insert_anchor = placeholder_para._element
                if is_full_width:
                    insert_anchor = _full_width_table_anchor(
                        insert_anchor, doc, table_elem, style_spec,
                        pdf_guidance=tikz_tbl.get("pdf_guidance"))
                _insert_after_in_order(insert_anchor, block_elems)
                if is_full_width:
                    _apply_float_page_top(
                        insert_anchor, block_elems,
                        pdf_guidance=tikz_tbl.get("pdf_guidance"))
                placeholder_para._element.getparent().remove(placeholder_para._element)
                print(f"  Replaced placeholder [{placeholder}] with TikZ table {i + 1}")
        else:
            block_elems = [cap_para_elem, table_elem]
            if is_full_width:
                block_elems = _full_width_block_elements(doc, block_elems)
            _keep_float_block_together(block_elems)
            for elem in (elem for elem in block_elems if elem is not None):
                doc.element.body.append(elem)
            print(f"  Warning: Placeholder [{placeholder}] not found, table {i + 1} appended at end")

        _keep_table_on_one_page(table_elem)
        for tr in table_elem.iter(qn("w:tr")):
            # 去除单元格内的 numPr（防止黑色方块/项目符号）
            for tc in tr.iter(qn("w:tc")):
                for p in tc.iter(qn("w:p")):
                    pp = p.find(qn("w:pPr"))
                    if pp is not None:
                        np = pp.find(qn("w:numPr"))
                        if np is not None:
                            pp.remove(np)

        # 补全表格最后一行底线（仅 bottom，不补中间横线）
        tbl_obj = None
        for tbl in doc.tables:
            if tbl._element is table_elem:
                tbl_obj = tbl
                break
        if tbl_obj is None:
            from docx.table import Table
            tbl_obj = Table(table_elem, doc)
        if tbl_obj:
            _complete_table_bottom_line(tbl_obj)

    _remove_blank_paragraphs_around_tables(doc)
    _remove_blank_paragraphs_around_images(doc)

    # 修复续接段落缩进（表格前后应连续的文字）
    _fix_table_continuation_indent(doc)

    # 彻底清理所有表格单元格的 numPr + 空段落 + 列表缩进残留
    # 防止黑色方块重新出现
    _deep_clean_table_cells(doc)
    _enforce_caption_table_keep_together(doc)

    # 最终防线：在XML层面彻底移除numbering.xml中所有列表定义
    # 防止Word模板中的残留编号定义与表格单元格numPr交互产生黑色方块
    _purge_numbering_xml(doc)

    doc.save(docx_path)
    return docx_path


def _restore_equation_numbers(doc, formulas):
    """恢复Word中OMML公式的编号

    兼容两种数据格式:
    - formulas格式: {'type': 'display', 'eq_num': '1', ...}
    - display_formula_data格式: {'latex': '...', 'eq_num': '1', 'env': 'equation'}
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 兼容两种格式: 有'type'键的用type过滤, 否则全视为display
    display_formulas = []
    for f in formulas:
        eq_num = f.get('eq_num', '')
        if not eq_num:
            continue
        # 有type字段的按type过滤
        if 'type' in f:
            if f['type'] == 'display':
                display_formulas.append(f)
        else:
            # display_formula_data格式(env字段), 全部是display
            display_formulas.append(f)

    if not display_formulas:
        print('  No display formulas with numbers found')
        return

    standalone_paras = []
    for para in doc.paragraphs:
        has_omml = False
        for elem in para._element.iter(qn('m:oMath')):
            has_omml = True
            break
        if not has_omml:
            continue
        text = para.text.strip()
        if len(text) <= 10:
            standalone_paras.append(para)

    print(f'  Standalone formula paragraphs: {len(standalone_paras)}, display formulas: {len(display_formulas)}')

    matched = min(len(standalone_paras), len(display_formulas))
    for i in range(matched):
        para = standalone_paras[i]
        eq_num = display_formulas[i]['eq_num']

        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._element.insert(0, pPr)

        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:right'), '0')
        ind.set(qn('w:firstLine'), '0')

        tabs = pPr.find(qn('w:tabs'))
        if tabs is None:
            tabs = OxmlElement('w:tabs')
            pPr.append(tabs)
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '8505')
        tabs.append(tab)

        run = para.add_run(f'\t{eq_num}')
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    print(f'  Added equation numbers to {matched} formulas')
