#!/usr/bin/env python3
"""Create a PDF-faithful DOCX by placing each PDF page as a full-page image.

This mode is for visual/layout fidelity. The generated DOCX is not a normal
editable-flow manuscript; text and tables are fixed as page images so Word's
layout engine cannot reflow them differently from LaTeX/PDF.
"""

from __future__ import annotations

import argparse
import tempfile
from pathlib import Path

import fitz
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


def _pt_to_inches(value: float) -> float:
    return float(value) / 72.0


def _anchor_page_image(inline):
    """Convert an inline drawing to a page-positioned anchor."""
    children = {child.tag: child for child in list(inline)}
    inline.clear()
    inline.tag = qn("wp:anchor")
    for key, value in {
        "simplePos": "0",
        "relativeHeight": "0",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
    }.items():
        inline.set(key, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    inline.append(simple_pos)
    inline.append(_page_position("wp:positionH", "left"))
    inline.append(_page_position("wp:positionV", "top"))
    for tag in ("wp:extent", "wp:effectExtent"):
        child = children.get(qn(tag))
        if child is not None:
            inline.append(child)
    inline.append(OxmlElement("wp:wrapNone"))
    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        child = children.get(qn(tag))
        if child is not None:
            inline.append(child)


def _page_position(tag, align_text):
    pos = OxmlElement(tag)
    pos.set("relativeFrom", "page")
    align = OxmlElement("wp:align")
    align.text = align_text
    pos.append(align)
    return pos


def pdf_to_exact_docx(
    pdf_path: str | Path,
    output_path: str | Path,
    dpi: int = 240,
    page_scale: float = 1.0,
    layout_spec: dict | None = None,
) -> str:
    pdf_path = Path(pdf_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    pdf = fitz.open(str(pdf_path))
    if pdf.page_count == 0:
        raise ValueError(f"PDF has no pages: {pdf_path}")

    doc = Document()
    section = doc.sections[0]
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(0))

    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)

    with tempfile.TemporaryDirectory(prefix="pdf_exact_docx_") as tmp:
        tmp_dir = Path(tmp)
        for page_index, page in enumerate(pdf):
            width_in = _pt_to_inches(page.rect.width)
            height_in = _pt_to_inches(page.rect.height)

            if page_index > 0:
                doc.add_page_break()

            section = doc.sections[-1]
            section.page_width = Inches(width_in)
            section.page_height = Inches(height_in)
            for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
                setattr(section, margin, Inches(0))

            image_path = tmp_dir / f"page_{page_index + 1:04d}.png"
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            pix.save(str(image_path))

            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.left_indent = Inches(0)
            paragraph.paragraph_format.right_indent = Inches(0)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.line_spacing = 1
            run = paragraph.add_run()
            fit_scale = page_scale
            drawing = run.add_picture(
                str(image_path),
                width=Inches(width_in * fit_scale),
                height=Inches(height_in * fit_scale),
            )
            _anchor_page_image(drawing._inline)

    doc.save(str(output_path))
    pdf.close()
    return str(output_path)


def _apply_layout_columns(section, layout_spec):
    page_spec = (layout_spec or {}).get("page_geometry", {})
    try:
        column_count = int(page_spec.get("column_count", 1) or 1)
    except (TypeError, ValueError):
        column_count = 1
    if column_count < 2:
        return
    cols = section._sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        section._sectPr.append(cols)
    cols.set(qn("w:num"), str(column_count))
    try:
        column_sep_mm = float(page_spec.get("column_sep_mm", 0) or 0)
    except (TypeError, ValueError):
        column_sep_mm = 0
    if column_sep_mm:
        cols.set(qn("w:space"), str(int(round(column_sep_mm * 56.6929))))


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a PDF-faithful DOCX facsimile.")
    parser.add_argument("pdf", help="Input PDF path")
    parser.add_argument("-o", "--output", required=True, help="Output DOCX path")
    parser.add_argument("--dpi", type=int, default=240, help="Render DPI, default: 240")
    parser.add_argument("--page-scale", type=float, default=1.0,
                        help="Image width/page width ratio, default: 1.0")
    args = parser.parse_args()
    print(pdf_to_exact_docx(args.pdf, args.output, dpi=args.dpi, page_scale=args.page_scale))


if __name__ == "__main__":
    main()
