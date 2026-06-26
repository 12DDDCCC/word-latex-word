from pathlib import Path
import subprocess
import sys

SKILL_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SKILL_DIR / 'convert-latex'))
sys.path.insert(0, str(SKILL_DIR / 'tex-to-word'))
sys.path.insert(0, str(SKILL_DIR / 'table-lossless-extract'))
sys.path.insert(0, str(SKILL_DIR / 'journal-template-extract'))
sys.path.insert(0, str(SKILL_DIR / 'image-rename'))
sys.path.insert(0, str(SKILL_DIR / 'omml-to-latex'))
sys.path.insert(0, str(SKILL_DIR / 'shared'))
sys.path.insert(0, str(SKILL_DIR / 'citation-extract'))

import orchestrator
import assemble_tex as assemble_tex_module
import _docx_insert as docx_insert_module
from assemble_tex import (
    _apply_spec_to_preamble,
    _ensure_cjk_min_parindent,
    _ensure_url_breaking,
    _normalize_class_options_for_spec,
    _uses_strip_full_width,
)
from spec_adapter import SpecAdapter
from template_spec_extract import _ensure_cls_files
from _docx_insert import (
    _apply_template_styles_to_doc,
    _body_parindent_pt,
    _caption_para_element,
    _caption_style_from_specs,
    _compiled_layout_twips,
    _complete_table_bottom_line,
    _ensure_front_matter_single_column,
    _full_width_float_insert_anchor,
    _full_width_block_elements,
    _latex_image_width_to_docx,
    _load_template_word_style,
    _merge_template_layout,
    _section_break_column_count,
    _table_font_size_from_specs,
    _template_line_numbering,
    _template_page_numbering,
    _template_title_before_space_pt,
    _word_page_geometry,
    _word_config_mode,
)
from pdf_exact_docx import _apply_layout_columns, pdf_to_exact_docx
from _docx_audit import audit_docx_structure, is_clean
from tex_postprocess import (
    _fallback_strip_fullwidth_to_floats,
    _fix_mathrm_wrappers,
    _sanitize_bbl_urls,
    postprocess_tex,
    validate_bibliography_keys,
)
from tikz_table_gen import process_table, table_requires_full_width
from supertabular_gen import (
    estimate_rendered_table_height_mm,
    rendered_table_output_height_mm,
    requires_multipage_table,
)
from _image_table_insert import (
    _append_full_width_block,
    _includegraphics_options,
    _insert_images_and_tables,
    _table_required_space_mm,
    build_image_map,
    image_requires_full_width,
)
from _paragraph_process import (
    _clean_formula_latex,
    _insert_label_before_outer_math_end,
    _process_paragraph_loop,
)
from _tex_extraction import _count_tabular_columns, _parse_tabular
from caption_utils import clean_caption
from image_order_match import natural_sort_key
from omml_to_latex import GREEK_MAP, _convert_accent, _convert_run
from _cls_cfg_parser import parse_cfg_file, parse_cls_file
from template_config import get_page_geometry_for_mode
from numbering_system import convert_numbering_references, renumber_sectioned_to_simple
import formula_roundtrip as formula_roundtrip_module


def _patch_light_pipeline(monkeypatch, tmp_path):
    monkeypatch.setattr(orchestrator, 'extract_docx_text',
                        lambda *a, **k: {'paragraphs': [{'text': 'Title'}]})
    monkeypatch.setattr(orchestrator, 'extract_all_images_with_position',
                        lambda *a, **k: [])
    monkeypatch.setattr(orchestrator, 'extract_tables',
                        lambda *a, **k: {'tables': []})
    monkeypatch.setattr(orchestrator, 'extract_citations',
                        lambda *a, **k: {'citations': []})
    monkeypatch.setattr(orchestrator, 'extract_template_spec',
                        lambda *a, **k: {'doc_options': [], 'journal': 'acp'})
    tex_path = tmp_path / 'acp_full.tex'
    tex_path.write_text(r'\documentclass{article}\begin{document}x\end{document}',
                        encoding='utf-8')
    monkeypatch.setattr(orchestrator, 'assemble_tex',
                        lambda *a, **k: (tex_path, {}))
    monkeypatch.setattr(orchestrator, 'copy_support_files',
                        lambda *a, **k: None)


def test_omml_greek_sigma_mapping_is_not_reversed():
    assert GREEK_MAP['σ'] == r'\sigma'
    assert GREEK_MAP['ς'] == r'\varsigma'


def test_omml_plain_math_run_keeps_underscore_semantics():
    from lxml import etree

    run = etree.fromstring(
        b'<m:r xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        b'<m:rPr><m:sty m:val="p"/></m:rPr><m:t>A_1</m:t></m:r>'
    )

    assert _convert_run(run) == r'\mathrm{A_1}'


def test_unknown_omml_accent_is_preserved_instead_of_becoming_hat():
    from lxml import etree

    accent = etree.fromstring(
        '<m:acc xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:accPr><m:chr m:val="\u20db"/></m:accPr>'
        '<m:e><m:r><m:t>x</m:t></m:r></m:e></m:acc>'.encode('utf-8')
    )

    assert _convert_accent(accent) == r'\overset{⃛}{x}'


def test_missing_latex_log_content_does_not_break_layout_read(monkeypatch, tmp_path):
    (tmp_path / 'main.log').write_text('', encoding='utf-8')
    monkeypatch.setattr(docx_insert_module, '_read_text_file', lambda _path: None)

    assert _compiled_layout_twips(tmp_path) == {}


def test_numbering_replacement_respects_math_command_and_word_boundaries():
    text = r'See (2-1), token(2-1)x, and \tag{2-1}.'
    mapped = convert_numbering_references(text, {'(2-1)': 'eq2-1'})

    assert mapped == r'See \eqref{eq2-1}, token(2-1)x, and \tag{2-1}.'

    renumbered = renumber_sectioned_to_simple(text)
    assert renumbered == r'See (1), token(2-1)x, and \tag{2-1}.'


def test_empty_formula_report_uses_zero_percentages(tmp_path):
    report = formula_roundtrip_module.generate_report([], [], tmp_path)
    content = Path(report).read_text(encoding='utf-8')

    assert '| OMML转换成功 | 0 | 0.0% |' in content


def test_tabular_caption_supports_nested_braces():
    table = _parse_tabular(
        r'\caption{A \textbf{nested} caption}'
        r'\begin{tabular}{lc}A & B\\1 & 2\end{tabular}',
        1,
    )

    assert table['caption'] == r'A \textbf{nested} caption'


def test_postprocess_keeps_nested_math_parentheses():
    tex = r'\begin{document}$((a+b)*c)$ ((Gui et al., 2024))\end{document}'

    processed = postprocess_tex(tex)

    assert '$((a+b)*c)$' in processed
    assert '(Gui et al., 2024)' in processed
    assert '((Gui et al., 2024))' not in processed


def test_nested_formula_label_goes_before_outer_end():
    latex = (
        r'\begin{align}'
        r'\begin{aligned}a&=b\\c&=d\end{aligned}'
        r'\end{align}'
    )

    labelled = _insert_label_before_outer_math_end(latex, r'\label{eq1}')

    assert r'\end{aligned}  \label{eq1}' in labelled
    assert labelled.endswith(r'\end{align}')


def test_image_map_keeps_multiple_images_in_same_paragraph():
    images = [
        {'para_index': 7, 'image_file': 'fig1.png'},
        {'para_index': 7, 'image_file': 'fig2.png'},
    ]

    mapped = build_image_map(images)

    assert [img['image_file'] for img in mapped[7]] == ['fig1.png', 'fig2.png']


def test_natural_sort_orders_fig10_after_fig2():
    assert sorted(['fig1.png', 'fig10.png', 'fig2.png'], key=natural_sort_key) == [
        'fig1.png', 'fig2.png', 'fig10.png'
    ]


def test_clean_caption_does_not_double_escape_textbackslash_braces():
    assert clean_caption(r'A\B_{2}') == r'A\textbackslash{}B\_\{2\}'


def test_tabular_column_count_handles_p_and_at_specs():
    assert _count_tabular_columns(r'|l|p{3cm}|@{}c@{}|X|') == 4


def test_cls_cfg_parser_matches_single_latex_backslash(tmp_path):
    cls = tmp_path / 'journal.cls'
    cls.write_text(
        '\n'.join([
            r'\ProvidesClass{journal}[2026/01/01 Test class]',
            r'\DeclareOption{acp}{}',
            r'\newcommand{\foo}[1]{#1}',
            r'\newcommand\bar{bar}',
            r'\DeclareMathOperator{\baz}{baz}',
        ]),
        encoding='utf-8',
    )
    cfg = tmp_path / 'journal.cfg'
    cfg.write_text(
        '\n'.join([
            r'\DeclareOption{acp}{}',
            r'\def\cop@journal@ACP@abbreviation{acp}',
            r'\def\cop@journal@ACP@name{Atmospheric Chemistry and Physics}',
        ]),
        encoding='utf-8',
    )

    cls_info = parse_cls_file(cls)
    cfg_info = parse_cfg_file(cfg)

    assert {'foo', 'bar', 'baz'}.issubset(set(cls_info['custom_commands']))
    assert 'acp' in cfg_info['journals']
    assert cfg_info['journals']['acp']['name'] == 'Atmospheric Chemistry and Physics'


def test_tex_roundtrip_summary_counts_semantic_matches_without_name_error(monkeypatch, tmp_path):
    tex = tmp_path / 'main.tex'
    tex.write_text(r'\begin{equation}a=b\end{equation}', encoding='utf-8')
    out = tmp_path / 'out'

    monkeypatch.setattr(
        formula_roundtrip_module, 'extract_formulas_from_tex',
        lambda _path: [{'latex': 'a=b', 'type': 'equation', 'env': 'equation'}],
    )
    monkeypatch.setattr(
        formula_roundtrip_module, 'roundtrip_single',
        lambda latex, _xslt=None: {
            'original_latex': latex,
            'roundtrip_latex': latex,
            'status': 'ok',
            'error': None,
        },
    )
    monkeypatch.setattr(
        formula_roundtrip_module, 'compare_latex',
        lambda _a, _b: {
            'exact_match': True,
            'normalized_match': True,
            'diff_note': '',
        },
    )
    monkeypatch.setattr(
        formula_roundtrip_module, 'generate_report',
        lambda *_args: str(out / 'report.md'),
    )
    monkeypatch.setattr(
        formula_roundtrip_module, 'generate_roundtrip_docx',
        lambda *_args: None,
    )

    result = formula_roundtrip_module.run_roundtrip_from_tex(tex, out)

    assert result['semantic_match'] == 0


def test_paragraph_tail_float_keeps_sectioned_label_context():
    paragraphs = [
        {
            'para_index': 1,
            'text': '1 Introduction',
            'latex': 'Introduction',
            'heading_level': 1,
            'semantic_type': 'body',
            'has_formula': False,
        },
        {
            'para_index': 2,
            'text': 'Body paragraph.',
            'latex': 'Body paragraph.',
            'heading_level': None,
            'semantic_type': 'body',
            'has_formula': False,
            'runs': [],
        },
    ]
    ref_map = {}
    img_insert_map = {
        2: [{
            'image_file': 'fig1.png',
            'caption': 'Figure 1.1: Tail figure',
            'caption_full': 'Figure 1.1: Tail figure',
            'width_pt': 120,
        }],
    }

    body, *_rest = _process_paragraph_loop(
        paragraphs, None, {}, img_insert_map, {}, {}, ref_map,
        {
            'conclusions_cmd': None,
            'introduction_cmd': None,
            'statement_cmds': {},
        }, {},
        set(), {}, 'sectioned', False,
    )

    assert r'\label{fig1-1}' in '\n'.join(body)
    assert ref_map['Figure 1.1'] == 'fig1-1'


def test_clean_formula_latex_removes_mojibake_dead_replacements():
    cleaned = _clean_formula_latex(r'\mathrm{CO₂}+\mathrm{X}\mathrm{CO₂}')

    assert r'\mathrm{CO}_{2}' in cleaned
    assert '鈧' not in cleaned


def test_mathrm_fix_respects_nested_braces():
    tex = r'$\mathrm{+\lambda_{1}\times X}+\mathrm{CO_{2}}$'

    fixed = _fix_mathrm_wrappers(tex)

    assert r'\mathrm{+\lambda' not in fixed
    assert r'+\lambda_{1}\times X' in fixed
    assert r'\mathrm{CO_{2}}' in fixed


def test_caption_escapes_tilde_and_caret():
    assert clean_caption('A~B^2') == (
        r'A\textasciitilde{}B\textasciicircum{}2'
    )


def test_word_conversion_failure_marks_pipeline_failed(monkeypatch, tmp_path):
    _patch_light_pipeline(monkeypatch, tmp_path)
    monkeypatch.setattr(orchestrator, 'tex_to_word', lambda *a, **k: None)

    result = orchestrator.run_pipeline(
        'input.docx', 'template', 'refs.bib', 'acp', tmp_path,
        compile_pdf=False, convert_word=True, verify=False)

    assert result.success is False
    assert any(w.step_name == 'tex-to-word' and w.severity == 'error'
               for w in result.warnings)


def test_word_conversion_falls_back_to_standalone_entrypoint(monkeypatch, tmp_path):
    _patch_light_pipeline(monkeypatch, tmp_path)
    seen = {}

    def imported_converter(*args, **kwargs):
        raise RuntimeError('Pandoc executable not found')

    def cli_converter(tex_path, output_path, bib_path, config_mode=None,
                      pdf_float_wrap=False, pdf_float_reflow=False):
        seen['fallback'] = (Path(tex_path).name, Path(output_path).name, config_mode)
        Path(output_path).write_text('docx placeholder', encoding='utf-8')
        return str(output_path)

    monkeypatch.setattr(orchestrator, 'tex_to_word', imported_converter)
    monkeypatch.setattr(orchestrator, '_tex_to_word_cli', cli_converter)

    result = orchestrator.run_pipeline(
        'input.docx', 'template', 'refs.bib', 'acp', tmp_path,
        compile_pdf=False, convert_word=True, verify=False,
        config_mode='final')

    assert result.success is True
    assert seen['fallback'] == ('acp_full.tex', 'acp_converted.docx', 'final')


def test_config_mode_is_forwarded_to_template_and_word(monkeypatch, tmp_path):
    seen = {}
    _patch_light_pipeline(monkeypatch, tmp_path)

    def template_stub(*args, **kwargs):
        seen['template_mode'] = kwargs.get('config_mode')
        return {'doc_options': [], 'journal': 'acp'}

    def word_stub(*args, **kwargs):
        seen['word_mode'] = kwargs.get('config_mode')
        out = Path(kwargs['output_path'])
        out.write_text('docx placeholder', encoding='utf-8')
        return str(out)

    monkeypatch.setattr(orchestrator, 'extract_template_spec', template_stub)
    monkeypatch.setattr(orchestrator, 'tex_to_word', word_stub)

    result = orchestrator.run_pipeline(
        'input.docx', 'template', 'refs.bib', 'acp', tmp_path,
        compile_pdf=False, convert_word=True, verify=False,
        config_mode='final')

    assert result.success is True
    assert seen == {'template_mode': 'final', 'word_mode': 'final'}


def test_pdf_float_wrap_is_forwarded_to_word_converter(monkeypatch, tmp_path):
    seen = {}
    _patch_light_pipeline(monkeypatch, tmp_path)

    def word_stub(*args, **kwargs):
        seen['wrap'] = kwargs.get('use_pdf_float_wrap')
        out = Path(kwargs['output_path'])
        out.write_text('docx placeholder', encoding='utf-8')
        return str(out)

    monkeypatch.setattr(orchestrator, 'tex_to_word', word_stub)

    result = orchestrator.run_pipeline(
        'input.docx', 'template', 'refs.bib', 'acp', tmp_path,
        compile_pdf=False, convert_word=True, verify=False,
        pdf_float_wrap=True)

    assert result.success is True
    assert seen['wrap'] is True


def test_final_mode_enables_pdf_float_wrap_by_default(monkeypatch, tmp_path):
    seen = {}
    _patch_light_pipeline(monkeypatch, tmp_path)

    def word_stub(*args, **kwargs):
        seen['wrap'] = kwargs.get('use_pdf_float_wrap')
        out = Path(kwargs['output_path'])
        out.write_text('docx placeholder', encoding='utf-8')
        return str(out)

    monkeypatch.setattr(orchestrator, 'tex_to_word', word_stub)

    result = orchestrator.run_pipeline(
        'input.docx', 'template', 'refs.bib', 'acp', tmp_path,
        compile_pdf=False, convert_word=True, verify=False,
        config_mode='final')

    assert result.success is True
    assert seen['wrap'] is True


def test_pdf_float_wrap_can_be_disabled_for_final_mode(monkeypatch, tmp_path):
    seen = {}
    _patch_light_pipeline(monkeypatch, tmp_path)

    def word_stub(*args, **kwargs):
        seen['wrap'] = kwargs.get('use_pdf_float_wrap')
        out = Path(kwargs['output_path'])
        out.write_text('docx placeholder', encoding='utf-8')
        return str(out)

    monkeypatch.setattr(orchestrator, 'tex_to_word', word_stub)

    result = orchestrator.run_pipeline(
        'input.docx', 'template', 'refs.bib', 'acp', tmp_path,
        compile_pdf=False, convert_word=True, verify=False,
        config_mode='final', pdf_float_wrap=False)

    assert result.success is True
    assert seen['wrap'] is False


def test_standalone_word_conversion_infers_final_mode_from_generated_tex():
    assert _word_config_mode(None, ['acp', 'final']) == 'final'
    assert _word_config_mode(None, ['acp', 'manuscript']) == 'manuscript'
    assert _word_config_mode('classic', ['acp', 'final']) == 'classic'


def test_layout_without_column_count_keeps_cls_detected_columns():
    style_spec = {
        'page_geometry': {'column_count': 2, 'textwidth_mm': 140},
        'column_count': 2,
    }
    layout_spec = {'page_geometry': {'textheight_mm': 220}}

    merged = _merge_template_layout(style_spec, layout_spec)

    assert merged['page_geometry']['column_count'] == 2
    assert merged['column_count'] == 2


def test_reference_doc_layout_applies_columns_without_overwriting_fonts(
        monkeypatch, tmp_path):
    from docx import Document
    from docx.oxml.ns import qn

    docx_path = tmp_path / 'reference.docx'
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.add_paragraph('Body')
    doc.save(docx_path)

    monkeypatch.setattr(
        docx_insert_module,
        '_load_template_word_style',
        lambda *_args, **_kwargs: {
            'page_geometry': {
                'paperwidth_mm': 210,
                'paperheight_mm': 297,
                'textwidth_mm': 170,
                'oddsidemargin_mm': 20,
                'right_margin_mm': 20,
                'topmargin_mm': 20,
                'column_count': 2,
                'column_sep_mm': 6,
            },
            'column_count': 2,
            'column_sep_mm': 6,
            'footer_dims': {'bottom_margin_mm': 20, 'footskip_mm': 10},
        },
    )

    docx_insert_module.apply_template_word_layout(
        docx_path, 'paper.tex', config_mode='final')

    converted = Document(docx_path)
    cols = converted.sections[-1]._sectPr.find(qn('w:cols'))
    assert cols.get(qn('w:num')) == '2'
    assert converted.styles['Normal'].font.name == 'Arial'


def test_tikz_fallback_table_font_size_uses_template_specs():
    layout = {'table': {'body_size': '10pt'}}
    style = {'small_size': 8, 'layout_spec': {'table': {'body_size': '7pt'}}}

    assert _table_font_size_from_specs(style, layout) == 10
    assert _table_font_size_from_specs({'small_size': 8}, None) == 8
    assert _table_font_size_from_specs({}, {'table': {'body_size': '\\small'}}) == 9


def test_template_line_and_page_numbering_are_derived_from_cls():
    cls = (
        r'\RequirePackage[mathlines,modulo]{lineno}'
        r'\linenumbersep3\p@'
        r'\linenumbers'
        r'\def\@oddfoot{\reset@font\bfseries\hfil\thepage\hfil}'
    )

    line_spec = _template_line_numbering(cls, ['manuscript'])
    page_spec = _template_page_numbering(cls)

    assert line_spec == {'count_by': 5, 'distance_twips': 60}
    assert page_spec == {'footer': 'center'}


def test_manuscript_title_spacing_is_derived_from_maketitle_block():
    cls = (
        r'\def\@@maketitlemanuscript{'
        r'\parbox[t]{\textwidth}{\@manuscriptInfo}\par'
        r'\vskip\baselineskip'
        r'\raggedright{\LARGE\@title\par}}'
    )

    assert _template_title_before_space_pt(cls, ['manuscript'], 15.4) == 46.2


def test_final_title_spacing_converts_template_vspace_mm_to_pt():
    cls = (
        r'\def\@@maketitlefinal{'
        r'\vspace*{31mm}'
        r'{\LARGE\@title\par}}'
    )

    assert _template_title_before_space_pt(cls, [], 12) == 87.87


def test_final_title_spacing_accepts_template_vskip_before_title():
    cls = (
        r'\def\@@maketitlefinal{'
        r'\vskip2\baselineskip'
        r'{\LARGE\@title\par}}'
    )

    assert _template_title_before_space_pt(cls, [], 12) == 24.0


def test_generic_maketitle_spacing_sums_numeric_vskips_before_title():
    cls = (
        r'\def\@maketitle{'
        r'\newpage\null\vskip 18.5pt'
        r'\colorbox{dsubhdcolor}{\revsptitlefont \@subhead\par}\vskip9pt'
        r'{\titlefont \@title \par}}'
    )

    assert _template_title_before_space_pt(cls, [], 14) == 27.5


def test_twocolumn_wide_table_uses_template_full_width():
    layout_spec = {
        'document': {
            'is_twocolumn': True,
            'supports_double_column_floats': True,
        },
        'page_geometry': {
            'textwidth_mm': 178,
            'column_count': 1,
            'column_sep_mm': 7,
        },
    }
    cells = [
        {'text': f'H{i}', 'gridSpan': 1, 'borders': {}, 'paragraphs': []}
        for i in range(6)
    ]
    table = {
        'grid_cols': [{'width_twips': 1400} for _ in range(6)],
        'rows': [{'row_height': 400, 'cells': cells}],
    }

    assert table_requires_full_width(table, layout_spec) is True
    assert r'\resizebox{\textwidth}' in process_table(table, 1, layout_spec)


def test_wrapped_tall_table_uses_multipage_template_environment():
    layout_spec = {
        'table': {'multipage_environment': 'supertabular'},
        'page_geometry': {'textwidth_mm': 177, 'textheight_mm': 170},
    }
    rows = []
    for idx in range(16):
        rows.append({
            'row_height': 400,
            'cells': [
                {
                    'text': 'Warm-temperate semi-humid climate class',
                    'gridSpan': 1,
                    'borders': {},
                    'paragraphs': [],
                    'bold': idx == 0,
                },
                *[
                    {
                        'text': str(idx),
                        'gridSpan': 1,
                        'borders': {},
                        'paragraphs': [],
                    }
                    for _ in range(6)
                ],
            ],
        })
    table = {
        'grid_cols': [{'width_twips': 1400} for _ in range(7)],
        'rows': rows,
    }

    assert sum(row['row_height'] for row in rows) / 56.6929 < 170
    assert estimate_rendered_table_height_mm(table, layout_spec) > 170
    assert requires_multipage_table(table, layout_spec) is True


def test_non_multipage_tall_table_preserves_text_size():
    layout_spec = {
        'document': {'is_twocolumn': True, 'supports_double_column_floats': True},
        'page_geometry': {
            'textwidth_mm': 164.83,
            'textheight_mm': 190,
            'column_count': 2,
            'column_sep_mm': 3.5,
        },
    }
    rows = []
    for idx in range(18):
        rows.append({
            'row_height': 400,
            'cells': [
                {
                    'text': 'Long land-cover and climate description',
                    'gridSpan': 1,
                    'borders': {},
                    'paragraphs': [],
                },
                *[
                    {
                        'text': str(idx),
                        'gridSpan': 1,
                        'borders': {},
                        'paragraphs': [],
                    }
                    for _ in range(6)
                ],
            ],
        })
    table = {
        'grid_cols': [{'width_twips': 1400} for _ in range(7)],
        'rows': rows,
    }

    rendered = process_table(table, 2, layout_spec)

    assert r'\resizebox{\textwidth}{!}' in rendered
    assert r'\resizebox{!}' not in rendered


def test_equation_wrap_depends_on_template_width():
    formula = (
        r"\begin{equation}" "\n"
        r"\tag{6}" "\n"
        r"X_{i}^{b}=X_{\mathbf{NEE}}^{b}\times \lambda_{\mathbf{1}}\times "
        r"\delta_{i}+X_{\mathbf{OCN}}^{b}\times \lambda_{\mathbf{2}}\times "
        r"\delta_{i}+X_{\mathbf{FIRE}}^{b}+X_{\mathbf{FOSSIL}}^{b},i=1,2,...,N" "\n"
        r"\end{equation}"
    )
    single = {
        "document": {"is_twocolumn": False},
        "page_geometry": {"column_count": 1, "textwidth_mm": 166.0},
    }
    double = {
        "document": {"is_twocolumn": True},
        "page_geometry": {"column_count": 2, "textwidth_mm": 166.0, "column_sep_mm": 7.0},
    }
    assert r"\begin{split}" not in postprocess_tex(formula, layout_spec=single)
    assert r"\begin{split}" in postprocess_tex(formula, layout_spec=double)


def test_twocolumn_equation_wraps_inline_tag_without_crossing_equations():
    tex = (
        r"\begin{equation}" "\n"
        r"a=b \tag{1}" "\n"
        r"\label{eq1}" "\n"
        r"\end{equation}" "\n"
        "text between\n"
        r"\begin{equation}" "\n"
        r"\boldsymbol{X}_{\boldsymbol{i}}^{\boldsymbol{b}}="
        r"\boldsymbol{X}_{\boldsymbol{NEE}}^{\boldsymbol{b}}\times\lambda_{\boldsymbol{1}}\times\delta_{\boldsymbol{i}}+"
        r"\boldsymbol{X}_{\boldsymbol{OCN}}^{\boldsymbol{b}}\times \lambda_{\boldsymbol{2}}\times\delta_{\boldsymbol{i}}+"
        r"\boldsymbol{X}_{\boldsymbol{FIRE}}^{\boldsymbol{b}}+\boldsymbol{X}_{\boldsymbol{FOSSIL}}^{\boldsymbol{b}} , i = 1, 2, ... , N \tag{6}" "\n"
        r"\label{eq6}" "\n"
        r"\end{equation}"
    )
    double = {
        "document": {"is_twocolumn": True},
        "page_geometry": {
            "column_count": 2,
            "textwidth_mm": 177.0,
            "column_sep_mm": 7.0,
        },
    }

    fixed = postprocess_tex(tex, layout_spec=double)

    assert fixed.count(r"\begin{equation}") == 2
    assert fixed.count(r"\begin{split}") == 1
    assert r"a=b \tag{1}" in fixed
    assert fixed.index(r"\end{split}") < fixed.index(r"\tag{6}\label{eq6}")
    assert r"FIRE" in fixed and r"FOSSIL" in fixed


def test_twocolumn_image_star_requires_template_figure_support():
    base_layout = {
        'document': {
            'is_twocolumn': True,
            'supports_double_column_tables': True,
            'supports_double_column_figures': False,
        },
        'figure': {'allow_full_width': True},
        'page_geometry': {
            'textwidth_mm': 178,
            'column_count': 2,
            'column_sep_mm': 7,
        },
    }
    image = {'layout_width': 'textwidth', 'width_pt': 400}

    assert image_requires_full_width(image, base_layout) is False

    with_figure_star = dict(base_layout)
    with_figure_star['document'] = dict(base_layout['document'])
    with_figure_star['document']['supports_double_column_figures'] = True
    assert image_requires_full_width(image, with_figure_star) is True


def test_twocolumn_image_star_uses_source_width_when_template_allows_it():
    layout = {
        'document': {
            'is_twocolumn': True,
            'supports_double_column_figures': True,
        },
        'figure': {'allow_full_width': True},
        'page_geometry': {
            'textwidth_mm': 164.83,
            'column_count': 2,
            'column_sep_mm': 3.51,
        },
    }

    assert image_requires_full_width({'width_pt': 260}, layout) is True
    assert image_requires_full_width({'width_pt': 200}, layout) is False


def test_textwidth_image_hint_does_not_override_template_column_fit():
    layout = {
        'document': {
            'is_twocolumn': True,
            'supports_double_column_figures': True,
        },
        'figure': {'allow_full_width': True},
        'page_geometry': {
            'textwidth_mm': 164.83,
            'column_count': 2,
            'column_sep_mm': 3.51,
        },
    }

    assert image_requires_full_width(
        {'layout_width': 'textwidth', 'width_pt': 200}, layout
    ) is False


def test_figure_star_support_is_written_to_layout_spec(tmp_path):
    cls = tmp_path / 'journal.cls'
    cls.write_text(r'\newenvironment{figure*}{\@dblfloat{figure}}{\end@dblfloat}',
                   encoding='utf-8')
    spec = {'document_class': {'class_name': 'journal', 'declared_options': []}}

    layout = SpecAdapter(spec, {}, cls_path=cls, template_dir=tmp_path).to_layout_spec()

    assert layout['figure']['allow_full_width'] is True
    assert layout['figure']['full_width_source'] == 'explicit-class-figure-star'


def test_full_width_float_position_is_extracted_from_template_examples(tmp_path):
    cls = tmp_path / 'journal.cls'
    cls.write_text(r'\LoadClass[twocolumn]{article}', encoding='utf-8')
    sample = tmp_path / 'sample.tex'
    sample.write_text(
        r'\begin{figure*}[t]\end{figure*}' '\n'
        r'\begin{table*}[t]\end{table*}',
        encoding='utf-8',
    )
    spec = {'document_class': {'class_name': 'journal', 'declared_options': []}}

    layout = SpecAdapter(spec, {}, cls_path=cls, template_dir=tmp_path).to_layout_spec()

    assert layout['figure']['full_width_float_position'] == 't'
    assert layout['table']['full_width_float_position'] == 't'
    assert layout['figure']['full_width_container'] == 'native-star'
    assert layout['table']['full_width_container'] == 'native-star'


def test_float_policy_is_extracted_from_effective_template_values(tmp_path):
    cls = tmp_path / 'journal.cls'
    cls.write_text(
        '\\LoadClass[twocolumn]{article}\n'
        '\\renewcommand\\topfraction{.8}\n'
        '\\renewcommand\\dbltopfraction{.65}\n',
        encoding='utf-8',
    )
    spec = {'document_class': {'class_name': 'journal', 'declared_options': []}}

    layout = SpecAdapter(spec, {}, cls_path=cls, template_dir=tmp_path).to_layout_spec()

    assert layout['float_policy']['topfraction'] == 0.8
    assert layout['float_policy']['dbltopfraction'] == 0.65


def test_complete_static_float_policy_does_not_run_probe(tmp_path, monkeypatch):
    cls = tmp_path / 'journal.cls'
    cls.write_text(
        '\\def\\topfraction{1}\n'
        '\\def\\textfraction{0}\n'
        '\\def\\floatpagefraction{.7}\n'
        '\\def\\dbltopfraction{1}\n'
        '\\def\\dblfloatpagefraction{.7}\n',
        encoding='utf-8',
    )
    adapter = SpecAdapter({}, {}, cls_path=cls, template_dir=tmp_path)
    monkeypatch.setattr(
        adapter, '_class_probe_log',
        lambda: (_ for _ in ()).throw(AssertionError('probe should not run')),
    )

    policy = adapter._derive_float_policy()

    assert policy['topfraction'] == 1.0
    assert policy['dblfloatpagefraction'] == 0.7


def test_cross_column_image_height_uses_template_dbltopfraction():
    layout = {
        'document': {'is_twocolumn': True},
        'page_geometry': {'textheight_mm': 200},
        'float_policy': {'topfraction': 0.9, 'dbltopfraction': 0.7},
    }

    options = _includegraphics_options('\\textwidth', layout)

    assert 'height=122.0mm' in options
    assert 'keepaspectratio' in options


def test_explicit_strip_example_still_prefers_native_star_float(tmp_path):
    cls = tmp_path / 'journal.cls'
    cls.write_text(r'\LoadClass[twocolumn]{article}', encoding='utf-8')
    sample = tmp_path / 'sample.tex'
    sample.write_text(r'\begin{strip}wide content\end{strip}', encoding='utf-8')
    spec = {'document_class': {'class_name': 'journal', 'declared_options': []}}

    layout = SpecAdapter(spec, {}, cls_path=cls, template_dir=tmp_path).to_layout_spec()

    assert layout['figure']['full_width_container'] == 'native-star'
    assert layout['table']['full_width_container'] == 'native-star'


def test_full_width_blocks_use_strip_when_template_supports_cross_column():
    layout = {
        'document': {'is_twocolumn': True},
        'figure': {'full_width_container': 'strip'},
    }
    body_lines = []

    _append_full_width_block(body_lines, 'figure', ['\\centering', 'BODY'], layout)

    assert body_lines[0] == '% SKILL-FULLWIDTH-FLOAT figure pos=!t'
    assert body_lines[1] == '\\begin{strip}'
    assert body_lines[2] == r'\noindent\begin{minipage}{\textwidth}'
    assert '\\def\\@captype{figure}' in body_lines[4]
    assert body_lines[-2] == r'\end{minipage}'
    assert body_lines[-1] == '\\end{strip}'
    assert _uses_strip_full_width(layout) is True


def test_strip_full_width_block_keeps_content_together_without_needspace():
    layout = {
        'document': {'is_twocolumn': True},
        'figure': {'full_width_container': 'strip'},
    }
    body_lines = []

    _append_full_width_block(
        body_lines, 'figure', ['\\centering', 'BODY'], layout,
        required_space_mm=92.5)

    assert r'\skillneedspace{92.5mm}' not in body_lines
    assert body_lines[0] == '% SKILL-FULLWIDTH-FLOAT figure pos=!t'
    assert body_lines[2] == r'\noindent\begin{minipage}{\textwidth}'


def test_full_width_table_reserves_scaled_render_height():
    rows = [
        {
            'row_height': 900,
            'cells': [{'gridSpan': 1, 'text': f'row {index}'}],
        }
        for index in range(20)
    ]
    table = {'rows': rows, 'grid_cols': [{'width_twips': 4000}]}
    layout = {'page_geometry': {'textheight_mm': 193.3, 'textwidth_mm': 164.8}}

    assert estimate_rendered_table_height_mm(table, layout) > 193.3
    assert rendered_table_output_height_mm(table, layout) == 193.3 * 0.82
    assert _table_required_space_mm(table, layout) == 193.3 * 0.82 + 18.0


def test_full_width_table_is_emitted_as_one_indivisible_float():
    header = {
        'row_height': 400,
        'cells': [
            {'text': 'Climate region', 'gridSpan': 1, 'bold': True},
            {'text': 'Value', 'gridSpan': 1, 'bold': True},
        ],
    }
    rows = [header] + [
        {
            'row_height': 900,
            'cells': [
                {'text': f'Warm-temperate semi-humid region {idx}', 'gridSpan': 1},
                {'text': str(idx), 'gridSpan': 1},
            ],
        }
        for idx in range(30)
    ]
    table = {
        'table_index': 1,
        'position': {'caption': 'Table 1. Site compare'},
        'rows': rows,
        'grid_cols': [{'width_twips': 7000}, {'width_twips': 7000}],
    }
    layout = {
        'page_geometry': {
            'textheight_mm': 120,
            'textwidth_mm': 160,
            'column_count': 2,
            'column_sep_mm': 4,
            'columnwidth_mm': 78,
        },
        'document': {
            'is_twocolumn': True,
            'supports_double_column_tables': True,
        },
    }
    body_lines = []

    _insert_images_and_tables(
        0, {}, {0: [table]}, body_lines, set(), set(),
        layout_spec=layout, tab_counter=[0])
    tex = '\n'.join(body_lines)

    assert tex.count(r'\begin{table*}') == 1
    assert tex.count(r'\end{table*}') == 1
    assert 'Warm-temperate semi-humid region 29' in tex


def test_strip_needspace_macro_is_inserted_with_cuted_support():
    layout = {
        'document': {'is_twocolumn': True},
        'figure': {'full_width_container': 'strip'},
    }

    preamble = _apply_spec_to_preamble(
        r'\documentclass{article}', {}, layout, set())

    assert r'\usepackage{cuted}' in preamble
    assert r'\usepackage{needspace}' in preamble
    assert r'\newcommand{\skillneedspace}' in preamble


def test_native_star_preamble_enables_bottom_double_column_floats():
    layout = {
        'document': {'is_twocolumn': True},
        'figure': {'full_width_container': 'native-star'},
        'table': {'full_width_container': 'native-star'},
    }

    preamble = _apply_spec_to_preamble(
        r'\documentclass{article}', {}, layout, set())

    assert 'stfloats.sty' in preamble
    assert 'dblfloatfix.sty' in preamble


def test_native_star_float_does_not_force_float_only_page():
    layout = {
        'document': {'is_twocolumn': True},
        'figure': {
            'full_width_container': 'native-star',
            'full_width_float_position': 'tp',
        },
    }
    body_lines = []

    _append_full_width_block(body_lines, 'figure', ['\\centering', 'BODY'], layout)

    assert '\\FloatBarrier' not in body_lines
    assert body_lines[0] == '\\begin{figure*}[!t]'
    assert body_lines[1] == r'\noindent\begin{minipage}{\textwidth}'
    assert body_lines[-2] == r'\end{minipage}'
    assert body_lines[-1] == '\\end{figure*}'


def test_assembly_does_not_force_final_float_only_page():
    source = Path(assemble_tex_module.__file__).read_text(encoding='utf-8')
    bibliography_block = source[source.index('bib_spec_d ='):source.index('full_tex = preamble')]

    assert "body_lines.append('\\\\FloatBarrier')" not in bibliography_block


def test_metadata_author_mathrm_superscripts_do_not_corrupt_author_command():
    title_para = {
        'text': 'Title',
        'latex': 'Title',
        'semantic_type': 'title',
        'para_index': 0,
    }
    author_para = {
        'text': 'aa b,Fb J',
        'latex': r'aa $ \mathrm{b}^{1} $,Fb $ \mathrm{J}^{1,2,3} $',
        'semantic_type': 'author',
        'para_index': 1,
    }

    result = assemble_tex_module._assemble_metadata(
        r'\begin{document}' + '\n' + r'\maketitle' + '\n' + r'\end{document}',
        {'metadata_block': '', 'spec': {}, 'journal': 'elsarticle'},
        '',
        [title_para, author_para],
        title_para,
        [author_para],
        [],
        {},
        {},
        {},
    )

    author_lines = [line for line in result.splitlines() if line.startswith(r'\author{')]
    assert author_lines == [r'\author{aa b$^{1}$,Fb J$^{1,2,3}$}']
    assert r'\author{aa $ \mathrm{b}^{1} $,Fb $ \mathrm{J}^{1,2,3} $}^{1}' not in result


def test_metadata_affiliation_commands_follow_journal_templates():
    title_para = {'text': 'Title', 'latex': 'Title', 'semantic_type': 'title', 'para_index': 0}
    author_para = {'text': 'Author', 'latex': r'Author$^{1}$', 'semantic_type': 'author', 'para_index': 1}
    affil_para = {
        'text': 'Jiangsu Provincial Key Laboratory, China',
        'latex': r'$ ^{1}\mathrm{Jiangsu} $ Provincial Key Laboratory, China',
        'semantic_type': 'affiliation',
        'para_index': 2,
    }
    base_tex = r'\begin{document}' + '\n' + r'\maketitle' + '\n' + r'\end{document}'

    els = assemble_tex_module._assemble_metadata(
        base_tex,
        {'metadata_block': '', 'spec': {'document_class': {'class_name': 'elsarticle'}}, 'journal': 'elsarticle'},
        '',
        [title_para, author_para, affil_para],
        title_para,
        [author_para],
        [affil_para],
        {},
        {},
        {},
    )
    nsr = assemble_tex_module._assemble_metadata(
        base_tex,
        {'metadata_block': '', 'spec': {'document_class': {'class_name': 'nsr'}}, 'journal': 'nsr'},
        '',
        [title_para, author_para, affil_para],
        title_para,
        [author_para],
        [affil_para],
        {},
        {},
        {},
    )

    assert r'\affiliation[1]{organization={Jiangsu Provincial Key Laboratory, China}}' in els
    assert r'\affil[]{' not in els
    assert r'\affil{$^{1}$Jiangsu Provincial Key Laboratory, China}' in nsr
    assert r'\affil[]{' not in nsr


def test_restored_affiliations_follow_title_author_order():
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    doc.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph('Title').style = 'Title'
    doc.add_paragraph('Author').style = 'Author'
    doc.add_paragraph('Abstract. Body')
    heading = doc.add_paragraph('1 Introduction')
    heading.style = 'Heading 1'

    tex = '\n'.join([
        r'\title{Title}',
        r'\author{Author}',
        r'\affiliation[1]{organization={Jiangsu Provincial Key Laboratory, China}}',
        r'\affil{$^{2}$Jiangsu China}',
        r'\affil{$^{3}$Frontiers Science Center, China}',
    ])

    docx_insert_module._restore_title_author_affil(doc, tex)

    texts = [p.text for p in doc.paragraphs[:7]]
    assert texts[:6] == [
        'Title',
        'Author',
        '1 Jiangsu Provincial Key Laboratory, China',
        '2 Jiangsu China',
        '3 Frontiers Science Center, China',
        'Abstract. Body',
    ]
    assert all('[]' not in text and ']1' not in text for text in texts)


def test_elsarticle_manuscript_twocolumn_normalizes_to_3p():
    spec = {
        'document_class': {
            'class_name': 'elsarticle',
            'declared_options': ['preprint', 'final', 'review', '1p', '3p', '5p'],
            'default_options': ['twocolumn', 'manuscript'],
        }
    }

    options = _normalize_class_options_for_spec(
        spec, 'elsarticle', doc_options=['twocolumn', 'manuscript'])

    assert options == ['3p']


def test_class_option_normalization_does_not_map_other_templates_to_3p():
    spec = {
        'document_class': {
            'class_name': 'copernicus',
            'declared_options': ['copernicus', 'manuscript', 'final'],
            'default_options': ['copernicus'],
        }
    }

    options = _normalize_class_options_for_spec(
        spec, 'copernicus', doc_options=['twocolumn', 'manuscript'])

    assert options == ['copernicus', 'manuscript']
    assert '3p' not in options
    assert 'twocolumn' not in options


def test_class_option_normalization_preserves_options_without_declared_list():
    spec = {
        'document_class': {
            'class_name': 'custom',
            'declared_options': [],
            'default_options': ['twocolumn', 'manuscript'],
        }
    }

    options = _normalize_class_options_for_spec(
        spec, 'custom', doc_options=['twocolumn', 'manuscript'])

    assert options == ['twocolumn', 'manuscript']


def test_tex_to_word_extracts_strip_figure_and_table(tmp_path):
    sys.path.insert(0, str(SKILL_DIR / 'tex-to-word'))
    from _tex_extraction import extract_images_from_tex, extract_tikz_tables_from_tex

    tex = tmp_path / 'main.tex'
    tex.write_text(
        r'''
\begin{strip}
\begingroup
\makeatletter\def\@captype{figure}\makeatother
\includegraphics[width=\textwidth]{fig/a.png}
\caption{Wide figure}
\endgroup
\end{strip}
\begin{strip}
\begingroup
\makeatletter\def\@captype{table}\makeatother
\caption{Wide table}
\begin{tikzpicture}
\node[anchor=center] at (0,0) {A};
\end{tikzpicture}
\endgroup
\end{strip}
''',
        encoding='utf-8',
    )

    images = extract_images_from_tex(tex)
    tables = extract_tikz_tables_from_tex(tex)

    assert images[0]['is_full_width'] is True
    assert images[0]['caption'] == 'Wide figure'
    assert tables[0]['is_full_width'] is True
    assert tables[0]['caption'] == 'Wide table'


def test_compile_fallback_rewrites_strip_to_star_float(tmp_path):
    tex = tmp_path / 'main.tex'
    tex.write_text(
        r'''
% SKILL-FULLWIDTH-FLOAT table pos=t
\begin{strip}
\begingroup
\makeatletter\def\@captype{table}\makeatother
\centering
\caption{Wide table}
\begin{tikzpicture}\end{tikzpicture}
\endgroup
\end{strip}
''',
        encoding='utf-8',
    )

    assert _fallback_strip_fullwidth_to_floats(tex) is True
    fixed = tex.read_text(encoding='utf-8')
    assert r'\begin{table*}[t]' in fixed
    assert r'\@captype' not in fixed
    assert r'\begin{strip}' not in fixed


def test_compiled_image_width_twips_are_converted_to_docx_emu():
    class Section:
        page_width = None
        page_height = None
        left_margin = None
        right_margin = None
        top_margin = None
        bottom_margin = None

    width = _latex_image_width_to_docx(
        r'width=\textwidth',
        Section(),
        compiled_layout={'textwidth': 9380, 'columnwidth': 4590})

    assert width == 9380 * 635


def test_word_full_width_block_uses_native_continuous_sections():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    cols = doc.sections[0]._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        doc.sections[0]._sectPr.append(cols)
    cols.set(qn('w:num'), '2')

    block = _full_width_block_elements(doc, [OxmlElement('w:tbl')])

    assert len(block) == 3
    assert block[0].tag == qn('w:p')
    assert block[1].tag == qn('w:tbl')
    assert block[2].tag == qn('w:p')
    assert _section_break_column_count(block[0]) == 2
    assert _section_break_column_count(block[2]) == 1
    assert block[1].find(f'{qn("w:tblPr")}/{qn("w:tblpPr")}') is None


def test_full_width_float_anchor_lets_following_body_fill_gap():
    from docx import Document

    doc = Document()
    placeholder = doc.add_paragraph('[FIGURE_0]')
    body = doc.add_paragraph('后续正文应该先填充图表前的空白。')

    assert _full_width_float_insert_anchor(placeholder._element) is body._element


def test_full_width_float_anchor_treats_table_reference_as_body_text():
    from docx import Document

    doc = Document()
    placeholder = doc.add_paragraph('[TIKZ_TABLE_0]')
    body = doc.add_paragraph('表1是站点评估的内容，可以看到正文应填充图表前空白。')

    assert _full_width_float_insert_anchor(placeholder._element) is body._element


def test_full_width_float_anchor_uses_pdf_delay_after_text():
    from docx import Document

    doc = Document()
    placeholder = doc.add_paragraph('[FIGURE_0]')
    first = doc.add_paragraph('Short bridge text before the delayed float.')
    target = doc.add_paragraph(
        'This paragraph is confirmed by PDF rendering before the top float.')
    later = doc.add_paragraph('This paragraph should stay after the float anchor.')

    anchor = _full_width_float_insert_anchor(
        placeholder._element,
        doc=doc,
        pdf_guidance={
            'position': 'top',
            'delay_after': {
                'text': 'This paragraph is confirmed by PDF rendering before the top float.'
            },
        },
    )

    assert anchor is target._element
    assert anchor is not first._element
    assert anchor is not later._element


def test_full_width_float_anchor_keeps_pdf_inline_position():
    from docx import Document

    doc = Document()
    placeholder = doc.add_paragraph('[TIKZ_TABLE_0]')
    body = doc.add_paragraph('Following text should not be used as a delayed anchor.')

    anchor = _full_width_float_insert_anchor(
        placeholder._element,
        doc=doc,
        pdf_guidance={'position': 'inline', 'y0_ratio': 0.48},
    )

    assert anchor is placeholder._element
    assert anchor is not body._element


def test_pdf_top_float_anchor_can_cross_heading_to_fill_page():
    from docx import Document

    doc = Document()
    placeholder = doc.add_paragraph('[FIGURE_0]')
    heading = doc.add_heading('2 Methods', level=1)
    body = doc.add_paragraph('Body text after a heading may fill the prior page.')

    plain_anchor = _full_width_float_insert_anchor(placeholder._element, doc=doc)
    guided_anchor = _full_width_float_insert_anchor(
        placeholder._element,
        doc=doc,
        pdf_guidance={'position': 'top', 'page': 3},
    )

    assert plain_anchor is placeholder._element
    assert guided_anchor is body._element
    assert guided_anchor is not heading._element


def test_table_alignment_keeps_tblpr_schema_order():
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_table(rows=1, cols=1)

    _apply_template_styles_to_doc(doc, {
        'body_size': 10,
        'small_size': 9,
        'main_font': None,
        'cjk_font': None,
        'heading_font': None,
        'heading_bold': False,
        'mono_font': None,
        'line_spacing': None,
        'layout_spec': {'table': {'alignment': 'center'}},
    })

    tbl_pr = doc.tables[0]._tbl.find(qn('w:tblPr'))
    order = [node.tag.rsplit('}', 1)[-1] for node in tbl_pr]
    assert order.index('jc') < order.index('tblLook')


def test_cell_border_keeps_tcpr_schema_order():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(OxmlElement('w:vAlign'))

    docx_insert_module._set_cell_border(
        cell, bottom=('single', '8', '000000'))

    order = [node.tag.rsplit('}', 1)[-1] for node in tc_pr]
    assert order.index('tcBorders') < order.index('vAlign')


def test_table_bottom_line_is_completed_for_all_last_row_cells():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_bottom(tc, sz='12', color='336699'):
        tc_pr = tc.find(qn('w:tcPr'))
        if tc_pr is None:
            tc_pr = OxmlElement('w:tcPr')
            tc.insert(0, tc_pr)
        borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), sz)
        bottom.set(qn('w:color'), color)
        borders.append(bottom)
        tc_pr.append(borders)

    def bottom_spec(tc):
        borders = tc.find(f'{qn("w:tcPr")}/{qn("w:tcBorders")}')
        bottom = borders.find(qn('w:bottom')) if borders is not None else None
        return (
            bottom.get(qn('w:val')),
            bottom.get(qn('w:sz')),
            bottom.get(qn('w:color')),
        )

    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    last_tcs = table._tbl.findall(qn('w:tr'))[-1].findall(qn('w:tc'))
    set_bottom(last_tcs[0])

    _complete_table_bottom_line(table)

    assert [bottom_spec(tc) for tc in last_tcs] == [
        ('single', '12', '336699'),
        ('single', '12', '336699'),
        ('single', '12', '336699'),
    ]


def test_table_bottom_line_uses_physical_last_row_cells_and_tcpr_order():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    last_tr = table._tbl.findall(qn('w:tr'))[-1]
    physical_cells = last_tr.findall(qn('w:tc'))
    tc_pr = physical_cells[1].find(qn('w:tcPr'))
    shd = OxmlElement('w:shd')
    tc_pr.append(shd)

    _complete_table_bottom_line(table)

    for tc in physical_cells:
        bottom = tc.find(f'{qn("w:tcPr")}/{qn("w:tcBorders")}/{qn("w:bottom")}')
        assert bottom is not None
        assert bottom.get(qn('w:val')) == 'single'

    order = [node.tag.rsplit('}', 1)[-1] for node in tc_pr]
    assert order.index('tcBorders') < order.index('shd')


def test_full_width_block_tables_remain_plain_editable_tables():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.table import Table

    doc = Document()
    cols = doc.sections[0]._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        doc.sections[0]._sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    table = doc.add_table(rows=1, cols=1)
    table_elem = table._element
    table_elem.getparent().remove(table_elem)
    block = _full_width_block_elements(doc, [table_elem])
    table = Table(block[1], doc)

    _complete_table_bottom_line(table)

    assert table._tbl.find(f'{qn("w:tblPr")}/{qn("w:tblpPr")}') is None
    bottom = table._tbl.find(
        f'.//{qn("w:tcPr")}/{qn("w:tcBorders")}/{qn("w:bottom")}')
    assert bottom is not None
    assert bottom.get(qn('w:val')) == 'single'


def test_full_width_float_anchor_lets_next_body_paragraph_pass():
    from docx import Document

    doc = Document()
    placeholder = doc.add_paragraph('[FIGURE_0]')
    body_para = doc.add_paragraph('Body text that should fill the current page.')

    anchor = docx_insert_module._full_width_float_insert_anchor(
        placeholder._element)

    assert anchor is body_para._element


def test_apply_float_page_top_breaks_before_new_float_group():
    """新组（前一个是正文）→ block[0] 加 pageBreakBefore，强制到新页顶。"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph('前面的正文段落。')  # 正文在前 → 新组
    placeholder = doc.add_paragraph('[FIGURE_0]')
    block = [doc.add_paragraph()._element]  # 模拟图表块首段

    docx_insert_module._apply_float_page_top(placeholder._element, block)

    pPr = block[0].find(qn('w:pPr'))
    assert pPr is not None
    assert pPr.find(qn('w:pageBreakBefore')) is not None


def test_apply_float_page_top_respects_pdf_inline_guidance():
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph('Body before float.')
    placeholder = doc.add_paragraph('[FIGURE_0]')
    block = [doc.add_paragraph('Figure body')._element]

    docx_insert_module._apply_float_page_top(
        placeholder._element,
        block,
        pdf_guidance={'position': 'inline', 'y0_ratio': 0.45},
    )

    pPr = block[0].find(qn('w:pPr'))
    assert pPr is None or pPr.find(qn('w:pageBreakBefore')) is None


def test_apply_float_page_top_keeps_pdf_top_guidance():
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph('Body before float.')
    placeholder = doc.add_paragraph('[FIGURE_0]')
    block = [doc.add_paragraph('Figure body')._element]

    docx_insert_module._apply_float_page_top(
        placeholder._element,
        block,
        pdf_guidance={'position': 'top', 'y0_ratio': 0.05},
    )

    pPr = block[0].find(qn('w:pPr'))
    assert pPr is not None
    assert pPr.find(qn('w:pageBreakBefore')) is not None


def test_protect_tikz_marks_table_star_full_width():
    from _pandoc_prep import _protect_tikz

    tex = r"""
\begin{table*}
\caption{Wide TikZ table}
\begin{tikzpicture}
\node {cell};
\end{tikzpicture}
\end{table*}
"""

    result, tables = _protect_tikz(tex)

    assert r'\textbf{[TIKZ_TABLE_0]}' in result
    assert len(tables) == 1
    assert tables[0]['is_full_width'] is True


def test_pdf_guidance_includes_meta_full_width_tikz_tables():
    from _pdf_float_guidance import _iter_guided_items

    records = list(_iter_guided_items(
        [],
        [{'tikz_body': '% meta:full_width=1\n\\node {cell};'}],
    ))

    assert len(records) == 1
    assert records[0]['kind'] == 'table'


def test_apply_float_page_top_skips_break_for_consecutive_float():
    """连续浮动体（前一个是表格+其 caption）→ 不加 pageBreakBefore，靠 keepWithNext 同页堆叠。"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_table(rows=1, cols=1)  # 前一个表格
    doc.add_paragraph('Table 1. 上一张表的标题。')  # 真 caption 紧跟表格
    placeholder = doc.add_paragraph('[TABLE_1]')
    block = [doc.add_paragraph()._element]

    docx_insert_module._apply_float_page_top(placeholder._element, block)

    pPr = block[0].find(qn('w:pPr'))
    assert pPr is None or pPr.find(qn('w:pageBreakBefore')) is None


def test_ensure_page_break_before_keeps_sectPr_last():
    """分节符段（pPr 含 sectPr）加 pageBreakBefore 时，sectPr 仍须在 pPr 最末。"""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    sect_para = doc.add_paragraph()
    pPr = sect_para._element.get_or_add_pPr()
    pPr.append(OxmlElement('w:sectPr'))  # 模拟分节符段

    docx_insert_module._ensure_page_break_before(sect_para._element)

    children = list(pPr)
    assert children[-1].tag == qn('w:sectPr')  # sectPr 仍是最末
    assert pPr.find(qn('w:pageBreakBefore')) is not None  # pageBreakBefore 已加


def test_ensure_page_break_before_is_idempotent():
    """重复调用不重复添加 pageBreakBefore。"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    para = doc.add_paragraph()

    docx_insert_module._ensure_page_break_before(para._element)
    docx_insert_module._ensure_page_break_before(para._element)  # 再调一次

    pPr = para._element.find(qn('w:pPr'))
    assert len(pPr.findall(qn('w:pageBreakBefore'))) == 1


def test_prev_sibling_is_float_treats_narrative_reference_as_new_group():
    """叙事引用"Figure 4.1 shows…"（前是正文）应为新组，不是连续浮动体。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph('Some body text before.')
    doc.add_paragraph('Figure 4.1 shows the result.')  # 叙事引用
    placeholder = doc.add_paragraph('[FIGURE_1]')

    assert docx_insert_module._prev_sibling_is_float(placeholder._element) is False


def test_prev_sibling_is_float_treats_real_caption_as_consecutive():
    """真 caption（紧跟表格）应为连续浮动体。"""
    from docx import Document

    doc = Document()
    doc.add_table(rows=1, cols=1)  # 前一个表格
    doc.add_paragraph('Table 1. Real caption.')  # 真 caption 紧跟表格
    placeholder = doc.add_paragraph('[TABLE_1]')

    assert docx_insert_module._prev_sibling_is_float(placeholder._element) is True


def test_is_caption_like_text_accepts_real_caption():
    """真 caption（编号后标点）应判为 caption（初步筛选通过，再由结构确认）。"""
    assert docx_insert_module._is_caption_like_text("Figure 4.1. The overview of model.")
    assert docx_insert_module._is_caption_like_text("图4.1. 模型概览。")
    assert docx_insert_module._is_caption_like_text("Table 1. Site comparison.")


def test_word_front_matter_is_single_column_before_twocolumn_body():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    cols = doc.sections[0]._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        doc.sections[0]._sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    doc.add_paragraph('Title').style = 'Title'
    doc.add_paragraph('Abstract. Body text')
    heading = doc.add_paragraph('1 Introduction')
    heading.style = 'Heading 1'

    assert _ensure_front_matter_single_column(doc) is True

    previous = heading._element.getprevious()
    assert previous.tag == qn('w:p')
    assert _section_break_column_count(previous) == 1
    heading_index = list(doc.element.body).index(heading._element)
    before_heading = list(doc.element.body)[:heading_index]
    assert all(
        elem.find(f'{qn("w:tblPr")}/{qn("w:tblpPr")}') is None
        for elem in before_heading
        if elem.tag == qn('w:tbl')
    )
    body_cols = doc.element.body.find(qn('w:sectPr')).find(qn('w:cols'))
    assert body_cols.get(qn('w:num')) == '2'


def test_copernicus_geometry_keeps_pdf_bleed_size():
    cls_content = r"""
    \if@stage@final
      \if@manuscript
        \oddsidemargin16.4mm
        \textwidth177mm
        \topmargin10mm
      \else
        \oddsidemargin16.4mm
        \textwidth177mm
        \topmargin\z@
      \fi
    \fi
    \newdimen\bleed \bleed3mm\relax
    \if@stage@final
      \if@manuscript
        \paperheight\dimexpr240mm+2\bleed\relax
      \else
        \paperheight\dimexpr277mm+2\bleed\relax
      \fi
      \paperwidth\dimexpr210mm+2\bleed\relax
    \fi
    """

    classic = get_page_geometry_for_mode(cls_content, config_mode='classic')
    final = get_page_geometry_for_mode(cls_content, config_mode='final')

    assert classic['paperwidth_mm'] == 216.0
    assert classic['paperheight_mm'] == 246.0
    assert final['paperwidth_mm'] == 216.0
    assert final['paperheight_mm'] == 283.0


def test_layout_spec_does_not_override_mode_page_geometry():
    style_spec = {
        'page_geometry': {
            'paperwidth_mm': 216.0,
            'paperheight_mm': 283.0,
            'textwidth_mm': 177.0,
            'oddsidemargin_mm': 16.4,
            'right_margin_mm': 22.6,
            'topmargin_mm': 0.0,
            'column_count': 2,
            'column_sep_mm': 7.0,
        },
        'column_count': 2,
        'column_sep_mm': 7.0,
    }
    layout_spec = {
        'page_geometry': {
            'paperwidth_mm': 166.0,
            'paperheight_mm': 283.0,
            'textwidth_mm': 177.0,
            'right_margin_mm': -27.4,
            'column_count': 2,
            'column_sep_mm': 7.0,
        }
    }

    merged = _merge_template_layout(style_spec, layout_spec)

    assert merged['page_geometry']['paperwidth_mm'] == 216.0
    assert merged['page_geometry']['right_margin_mm'] == 22.6


def test_word_style_does_not_invent_fonts_without_template_signal(tmp_path):
    tex_path = tmp_path / 'paper.tex'
    tex_path.write_text(
        r'\documentclass{article}'
        '\n'
        r'\usepackage[UTF8,fontset=windows]{ctex}'
        '\n'
        r'\begin{document}正文\end{document}',
        encoding='utf-8',
    )
    (tmp_path / 'article.cls').write_text(
        r'\renewcommand\normalsize{\@setfontsize\normalsize\@xpt\@xiipt}',
        encoding='utf-8',
    )

    style = _load_template_word_style(tex_path)

    assert style['main_font'] is None
    assert style['cjk_font'] is None


def test_layout_spec_fonts_fill_missing_word_style_fonts():
    style = {
        'main_font': None,
        'sans_font': None,
        'mono_font': None,
        'cjk_font': None,
        'heading_font': None,
        'heading_uses_sans': True,
    }
    layout_spec = {
        'document': {
            'main_font': 'Template Serif',
            'sans_font': 'Template Sans',
            'mono_font': 'Template Mono',
        }
    }

    merged = _merge_template_layout(style, layout_spec)

    assert merged['main_font'] == 'Template Serif'
    assert merged['sans_font'] == 'Template Sans'
    assert merged['mono_font'] == 'Template Mono'
    assert merged['heading_font'] == 'Template Sans'


def test_word_geometry_converts_raw_latex_zero_left_margin():
    geo = _word_page_geometry({
        'paperwidth_mm': 215.9,
        'paperheight_mm': 279.4,
        'textwidth_mm': 164.83,
        'textheight_mm': 193.3,
        'oddsidemargin_mm': 0.0,
        'right_margin_mm': 51.07,
        'topmargin_mm': 5.62,
    })

    assert 25.0 < geo['oddsidemargin_mm'] < 25.8
    assert 25.0 < geo['right_margin_mm'] < 25.8


def test_word_geometry_uses_template_header_for_zero_top_margin():
    geo = _word_page_geometry(
        {
            'paperwidth_mm': 166.0,
            'paperheight_mm': 271.0,
            'textwidth_mm': 177.0,
            'textheight_mm': 140.0,
            'oddsidemargin_mm': 16.4,
            'right_margin_mm': -27.4,
            'topmargin_mm': 0.0,
        },
        footer_dims={'headheight_mm': 16.4, 'headsep_mm': 5.0},
    )

    assert geo['paperwidth_mm'] == 209.8
    assert geo['topmargin_mm'] == 21.4


def test_word_table_xml_marks_rows_and_paragraphs_keep_together():
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    for idx, row in enumerate(table.rows):
        row.cells[0].text = f'row {idx}'

    docx_insert_module._keep_table_on_one_page(table._tbl)
    rows = list(table._tbl.iter(qn('w:tr')))

    for row_idx, row in enumerate(rows):
        trPr = row.find(qn('w:trPr'))
        assert trPr.find(qn('w:cantSplit')) is not None
        para = next(row.iter(qn('w:p')))
        pPr = para.find(qn('w:pPr'))
        assert pPr.find(qn('w:keepLines')) is not None
        keep_next = pPr.find(qn('w:keepNext'))
        if row_idx < len(rows) - 1:
            assert keep_next is not None
        else:
            assert keep_next is None


def test_word_figure_caption_keeps_lines_without_chaining_to_body():
    from docx.oxml.ns import qn

    elem = docx_insert_module._caption_para_element('Figure', '1', 'Caption')
    pPr = elem.find(qn('w:pPr'))

    assert pPr.find(qn('w:keepLines')) is not None
    assert pPr.find(qn('w:keepNext')) is None


def test_pdf_exact_docx_can_preserve_template_column_metadata():
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    layout_spec = {
        'page_geometry': {'column_count': 2, 'column_sep_mm': 3.5},
    }

    _apply_layout_columns(doc.sections[0], layout_spec)

    cols = doc.sections[0]._sectPr.find(qn('w:cols'))
    assert cols.get(qn('w:num')) == '2'


def test_pdf_exact_docx_page_image_mode_does_not_apply_columns():
    source = Path(pdf_to_exact_docx.__code__.co_filename).read_text(encoding='utf-8')
    body = source[source.index('def pdf_to_exact_docx'):source.index('def _apply_layout_columns')]

    assert '_apply_layout_columns(' not in body


def test_url_breaking_is_inserted_after_documentclass():
    preamble = '\\documentclass{article}\n\\usepackage{url}\n'

    fixed = _ensure_url_breaking(preamble)

    assert fixed.index('SKILL-URL-BREAKING') > fixed.index('\\documentclass')
    assert '\\IfFileExists{xurl.sty}' in fixed
    assert '\\UrlBreaks' in fixed


def test_cjk_body_parindent_floor_is_inserted_after_documentclass():
    preamble = '\\documentclass{article}\n'
    paragraphs = [{'semantic_type': 'body', 'text': '这是一段中文正文。'}]

    fixed = _ensure_cjk_min_parindent(preamble, paragraphs)

    assert fixed.index('SKILL-CJK-PARINDENT') > fixed.index('\\documentclass')
    assert r'\ifdim\parindent<2em' in fixed


def test_english_body_keeps_template_parindent():
    preamble = '\\documentclass{article}\n'
    paragraphs = [{'semantic_type': 'body', 'text': 'This is an English paragraph.'}]

    assert _ensure_cjk_min_parindent(preamble, paragraphs) == preamble
    assert _body_parindent_pt({'body_size': 10, 'parindent': 12}, False) == 12
    assert _body_parindent_pt({'body_size': 10, 'parindent': 12}, True) == 20


def test_bbl_arxiv_url_field_is_sanitized(tmp_path):
    bbl = tmp_path / 'main.bbl'
    bbl.write_text(
        r'\href{http://arxiv.org/abs/https://example.org/a-b}'
        r'{{\tt arXiv:https://example.org/a-b}}',
        encoding='utf-8',
    )

    assert _sanitize_bbl_urls(bbl) is True
    assert bbl.read_text(encoding='utf-8') == r'\url{https://example.org/a-b}'


def test_bbl_doi_url_is_normalized(tmp_path):
    bbl = tmp_path / 'main.bbl'
    bbl.write_text(r'\DOIprefix\doi{https://doi.org/10.1111/gcb.16412}', encoding='utf-8')

    assert _sanitize_bbl_urls(bbl) is True
    assert bbl.read_text(encoding='utf-8') == r'\DOIprefix\doi{10.1111/gcb.16412}'


def test_bibliography_key_validation_detects_mismatched_bib(tmp_path):
    tex = tmp_path / 'paper.tex'
    tex.write_text(
        r'\documentclass{article}'
        '\n'
        r'\begin{document}Text \cite{1,2}.\bibliography{references}\end{document}',
        encoding='utf-8',
    )
    (tmp_path / 'references.bib').write_text(
        '@Article{bs-1629,\n  title={Wrong sample entry}\n}\n',
        encoding='utf-8',
    )

    result = validate_bibliography_keys(tex, tmp_path)

    assert result['ok'] is False
    assert result['reason'] == 'missing-bib-keys'
    assert result['missing'] == ['1', '2']


def test_bibliography_key_validation_accepts_matching_bib(tmp_path):
    tex = tmp_path / 'paper.tex'
    tex.write_text(
        r'\documentclass{article}'
        '\n'
        r'\begin{document}Text \citep{1} and \citet{2}.'
        r'\bibliography{references}\end{document}',
        encoding='utf-8',
    )
    (tmp_path / 'references.bib').write_text(
        '@Article{1,\n  title={One}\n}\n'
        '@Article{2,\n  title={Two}\n}\n',
        encoding='utf-8',
    )

    result = validate_bibliography_keys(tex, tmp_path)

    assert result['ok'] is True
    assert result['missing'] == []


def test_template_explicit_table_star_support_is_extracted(tmp_path):
    cls = tmp_path / 'journal.cls'
    cls.write_text(r'\newenvironment{table*}{\@dblfloat{table}}{\end@dblfloat}',
                   encoding='utf-8')
    spec = {'document_class': {'class_name': 'journal', 'declared_options': []}}

    adapter = SpecAdapter(spec, {}, cls_path=cls, template_dir=tmp_path)
    support = adapter._derive_double_column_float_support('table')
    figure_support = adapter._derive_double_column_float_support('figure')

    assert support == {'supports': True, 'source': 'explicit-class-table-star'}
    assert figure_support == {'supports': False, 'source': 'not-detected'}


def test_article_twocolumn_table_star_support_is_marked_inherited(tmp_path):
    cls = tmp_path / 'journal.cls'
    cls.write_text(r'\LoadClass{article}', encoding='utf-8')
    spec = {
        'document_class': {
            'class_name': 'journal',
            'declared_options': ['twocolumn'],
        }
    }

    adapter = SpecAdapter(spec, {}, cls_path=cls, template_dir=tmp_path)
    support = adapter._derive_double_column_float_support('table')
    figure_support = adapter._derive_double_column_float_support('figure')

    assert support == {'supports': True, 'source': 'inherited-article-twocolumn'}
    assert figure_support == {'supports': True, 'source': 'inherited-article-twocolumn'}


def test_template_ins_sources_generate_cls_in_output(monkeypatch, tmp_path):
    import template_spec_extract

    template_dir = tmp_path / 'template'
    output_dir = tmp_path / 'out'
    template_dir.mkdir()
    (template_dir / 'journal.ins').write_text('% installer', encoding='utf-8')
    (template_dir / 'journal.dtx').write_text('% source', encoding='utf-8')

    def fake_run(cmd, cwd, **_kwargs):
        Path(cwd, 'journal.cls').write_text(
            r'\LoadClass{article}', encoding='utf-8')
        Path(cwd, 'journal.bst').write_text('% bst', encoding='utf-8')
        return subprocess.CompletedProcess(cmd, 0)

    monkeypatch.setattr(template_spec_extract.subprocess, 'run', fake_run)

    cls_files = _ensure_cls_files(template_dir, output_dir)

    assert cls_files == [output_dir / 'journal.cls']
    assert (output_dir / 'journal.bst').exists()
    assert not (template_dir / 'journal.cls').exists()


def test_geometry_probe_runs_outside_template_dir(monkeypatch, tmp_path):
    import spec_adapter

    template_dir = tmp_path / 'template'
    template_dir.mkdir()
    cls = template_dir / 'journal.cls'
    cls.write_text(r'\LoadClass{article}', encoding='utf-8')
    seen = {}

    def fake_run(cmd, cwd, **_kwargs):
        seen['cwd'] = Path(cwd)
        Path(cwd, '_skill_geometry_probe.log').write_text(
            '\n'.join([
                'SKILL-PROBE-TEXTWIDTH=100.0pt',
                'SKILL-PROBE-PAPERWIDTH=200.0pt',
                'SKILL-PROBE-COLUMNWIDTH=100.0pt',
                'SKILL-PROBE-COLUMNSEP=0.0pt',
            ]),
            encoding='utf-8',
        )
        return subprocess.CompletedProcess(cmd, 0)

    monkeypatch.setattr(spec_adapter.subprocess, 'run', fake_run)
    adapter = SpecAdapter(
        {'document_class': {'class_name': 'journal'}},
        {},
        cls_path=cls,
        template_dir=template_dir,
    )

    geo = adapter._probe_page_geometry()

    assert geo['textwidth_mm'] > 0
    assert seen['cwd'] != template_dir
    assert not list(template_dir.glob('_skill_geometry_probe*'))


def test_valid_static_single_column_geometry_skips_probe():
    adapter = SpecAdapter(
        {'document_class': {'class_name': 'journal'}},
        {},
        config_mode='classic',
    )
    geo = {
        'paperwidth_mm': 210.0,
        'textwidth_mm': 160.0,
        'right_margin_mm': 25.0,
        'column_count': 1,
    }

    assert adapter._page_geometry_needs_probe(geo, r'\LoadClass{article}') is False


def test_final_single_column_geometry_probes_when_cls_can_switch_twocolumn():
    adapter = SpecAdapter(
        {'document_class': {'class_name': 'journal', 'declared_options': ['final']}},
        {},
        config_mode='final',
    )
    geo = {
        'paperwidth_mm': 210.0,
        'textwidth_mm': 160.0,
        'right_margin_mm': 25.0,
        'column_count': 1,
    }

    assert adapter._page_geometry_needs_probe(geo, r'\@twocolumntrue') is True


def test_geometry_probe_copies_extended_template_support_files(tmp_path):
    template_dir = tmp_path / 'template'
    work_dir = tmp_path / 'probe'
    template_dir.mkdir()
    work_dir.mkdir()
    for name in (
        'journal.cls', 'journal.sty', 'journal.cfg', 'journal.clo',
        'local.ldf', 'style.fd', 'style.def', 'refs.bst',
        'biblatex.bbx', 'biblatex.cbx', 'biblatex.lbx', 'extra.bbd',
    ):
        (template_dir / name).write_text('% support', encoding='utf-8')

    SpecAdapter({}, {}, cls_path=template_dir / 'journal.cls')._copy_probe_support_files(
        template_dir, work_dir)

    for src in template_dir.iterdir():
        assert (work_dir / src.name).exists()


def _build_docx_with_floating_table(path, description=None):
    """构造一个含 w:tblpPr 浮动表格容器的 docx（模拟旧 _floating_full_width_block 输出）。"""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    tblPr = table._element.find(qn('w:tblPr'))
    tblpPr = OxmlElement('w:tblpPr')
    tblpPr.set(qn('w:vertAnchor'), 'paragraph')
    tblPr.append(tblpPr)
    if description:
        desc = OxmlElement('w:tblDescription')
        desc.set(qn('w:val'), description)
        tblPr.append(desc)
    doc.save(path)
    return path


def test_audit_clean_docx_has_no_floating_containers(tmp_path):
    from docx import Document
    path = tmp_path / 'clean.docx'
    Document().save(path)
    report = audit_docx_structure(path)
    assert report['xml_valid'] is True
    assert report['tblpPr_count'] == 0
    assert report['docpr_duplicate_ids'] == []
    assert report['dangling_rids_in_doc'] == []
    assert is_clean(path) is True


def test_audit_detects_floating_table_container(tmp_path):
    """审计函数必须能检出 tblpPr（验证审计本身有效，否则护返回恒真无意义）。"""
    path = tmp_path / 'bad.docx'
    _build_docx_with_floating_table(path)
    report = audit_docx_structure(path)
    assert report['tblpPr_count'] >= 1
    assert is_clean(path) is False


def test_audit_can_allow_skill_pdf_float_wrap_container(tmp_path):
    path = tmp_path / 'wrapped.docx'
    _build_docx_with_floating_table(path, description='skill-pdf-float-wrap-p2')
    report = audit_docx_structure(path)
    assert report['tblpPr_count'] == 1
    assert report['skill_floating_tblpPr_count'] == 1
    assert report['non_skill_tblpPr_count'] == 0
    assert is_clean(path) is False
    assert is_clean(path, allow_skill_floating=True) is True


def _minimal_drawing_para_element():
    from lxml import etree

    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    p = etree.Element(f'{{{w_ns}}}p')
    r = etree.SubElement(p, f'{{{w_ns}}}r')
    drawing = etree.SubElement(r, f'{{{w_ns}}}drawing')
    inline = etree.SubElement(drawing, f'{{{wp_ns}}}inline')
    extent = etree.SubElement(inline, f'{{{wp_ns}}}extent')
    extent.set('cx', '914400')
    extent.set('cy', '457200')
    etree.SubElement(inline, f'{{{wp_ns}}}effectExtent')
    etree.SubElement(inline, f'{{{wp_ns}}}docPr')
    etree.SubElement(inline, f'{{{wp_ns}}}cNvGraphicFramePr')
    etree.SubElement(inline, f'{{{a_ns}}}graphic')
    return p


def test_pdf_float_wrap_converts_section_wrapped_float_to_skill_container(tmp_path):
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from _pdf_float_reflow import _collect_float_blocks, _docx_to_workdir, _parse_document_xml
    from _pdf_float_wrap import wrap_cross_column_floats

    path = tmp_path / 'wrapped-float.docx'
    doc = Document()
    cols = doc.sections[0]._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        doc.sections[0]._sectPr.append(cols)
    cols.set(qn('w:num'), '2')

    doc.add_paragraph('1 Introduction')
    caption = _caption_para_element('Figure', '1', 'A wrapped caption')
    for elem in _full_width_block_elements(doc, [_minimal_drawing_para_element(), caption]):
        doc.element.body.append(elem)
    doc.save(path)

    stats = wrap_cross_column_floats(
        path,
        guidance_items=[{
            'caption': 'A wrapped caption',
            'number': '1',
            'pdf_guidance': {'page': 1, 'y0_pt': 72.0},
        }],
        verify_render=False,
        force_all=True,
    )
    report = audit_docx_structure(path)

    assert stats['enabled'] is True
    assert len(stats['accepted']) == 1
    assert report['tblpPr_count'] == 1
    assert report['skill_floating_tblpPr_count'] == 1
    assert is_clean(path, allow_skill_floating=True) is True

    work = tmp_path / 'unzipped'
    _docx_to_workdir(path, work)
    _, body = _parse_document_xml(work)
    assert _collect_float_blocks(body) == []

    tblp_pr = body.find(f'.//{qn("w:tblPr")}/{qn("w:tblpPr")}')
    assert tblp_pr is not None
    assert tblp_pr.get(qn('w:vertAnchor')) == 'margin'
    assert tblp_pr.get(qn('w:tblpYSpec')) == 'top'
    assert tblp_pr.get(qn('w:tblpY')) is None

    tc_bottom = body.find(f'.//{qn("w:tcPr")}/{qn("w:tcBorders")}/{qn("w:bottom")}')
    assert tc_bottom is not None
    assert tc_bottom.get(qn('w:val')) == 'nil'


def test_pdf_float_wrap_converts_inline_drawing_to_square_anchor():
    from lxml import etree
    from _pdf_float_wrap import WP_NS, A_NS, _convert_inline_to_square_anchor

    def tag(ns, local):
        return f'{{{ns}}}{local}'

    inline = etree.Element(tag(WP_NS, 'inline'))
    etree.SubElement(inline, tag(WP_NS, 'extent'))
    etree.SubElement(inline, tag(WP_NS, 'effectExtent'))
    etree.SubElement(inline, tag(WP_NS, 'docPr'))
    etree.SubElement(inline, tag(WP_NS, 'cNvGraphicFramePr'))
    etree.SubElement(inline, tag(A_NS, 'graphic'))

    assert _convert_inline_to_square_anchor(inline) is True
    assert inline.tag == tag(WP_NS, 'anchor')
    assert inline.find(tag(WP_NS, 'wrapSquare')) is not None
    assert inline.find(tag(WP_NS, 'positionH')).get('relativeFrom') == 'column'
    assert inline.find(tag(WP_NS, 'positionV')).get('relativeFrom') == 'paragraph'


def test_pdf_float_wrap_one_pass_converts_all_candidates(tmp_path):
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from _pdf_float_wrap import wrap_cross_column_floats

    path = tmp_path / 'multi-wrapped-float.docx'
    doc = Document()
    cols = doc.sections[0]._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        doc.sections[0]._sectPr.append(cols)
    cols.set(qn('w:num'), '2')

    doc.add_paragraph('1 Introduction')
    first = _caption_para_element('Figure', '1', 'First wrapped caption')
    for elem in _full_width_block_elements(doc, [_minimal_drawing_para_element(), first]):
        doc.element.body.append(elem)
    doc.add_paragraph('body between floats')
    second = _caption_para_element('Figure', '2', 'Second wrapped caption')
    for elem in _full_width_block_elements(doc, [_minimal_drawing_para_element(), second]):
        doc.element.body.append(elem)
    doc.save(path)

    stats = wrap_cross_column_floats(
        path,
        guidance_items=[
            {
                'caption': 'First wrapped caption',
                'number': '1',
                'pdf_guidance': {'page': 1, 'y0_pt': 72.0},
            },
            {
                'caption': 'Second wrapped caption',
                'number': '2',
                'pdf_guidance': {'page': 2, 'y0_pt': 96.0},
            },
        ],
        verify_render=False,
        force_all=True,
    )
    report = audit_docx_structure(path)

    assert stats['enabled'] is True
    assert stats['one_pass'] is True
    assert len(stats['accepted']) == 2
    assert report['skill_floating_tblpPr_count'] == 2


def test_pdf_float_wrap_skips_table_only_floats(tmp_path):
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from _pdf_float_reflow import _collect_float_blocks, _docx_to_workdir, _parse_document_xml
    from _pdf_float_wrap import wrap_cross_column_floats

    path = tmp_path / 'table-only-float.docx'
    doc = Document()
    cols = doc.sections[0]._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        doc.sections[0]._sectPr.append(cols)
    cols.set(qn('w:num'), '2')

    doc.add_paragraph('1 Introduction')
    caption = _caption_para_element('Table', '2', 'Editable table caption')
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'editable table cell'
    table_elem = table._element
    table_elem.getparent().remove(table_elem)
    for elem in _full_width_block_elements(doc, [caption, table_elem]):
        doc.element.body.append(elem)
    doc.save(path)

    stats = wrap_cross_column_floats(
        path,
        guidance_items=[{
            'caption': 'Editable table caption',
            'number': '2',
            'pdf_guidance': {'page': 2, 'y0_pt': 72.0},
        }],
        verify_render=False,
        force_all=True,
    )

    assert stats['enabled'] is True
    assert stats['accepted'] == []
    assert audit_docx_structure(path)['skill_floating_tblpPr_count'] == 0
    work = tmp_path / 'table_unzipped'
    _docx_to_workdir(path, work)
    _, body = _parse_document_xml(work)
    blocks = _collect_float_blocks(body)
    assert len(blocks) == 1
    assert 'Editable table caption' in blocks[0].caption


def test_generated_references_keep_numeric_labels():
    from docx import Document
    from _bbl_item_parser import BibItem
    from _ref_section_builder import build_references_section

    doc = Document()
    doc.add_paragraph('[REFERENCES_PLACEHOLDER]')
    result = build_references_section(
        doc,
        [BibItem(key='19', plain_text='Friedlingstein et al. Global Carbon Budget.')],
        cite_map={'19': '1'},
        cite_style='numbered',
    )

    text = '\n'.join(para.text for para in doc.paragraphs)
    assert result.bookmarks_added == 1
    assert '[1] Friedlingstein et al. Global Carbon Budget.' in text


def test_pdf_float_wrap_force_all_still_rejects_worse_table_only_trial():
    from _pdf_float_wrap import _accepts_trial

    current = (267.0, 3, 148.0)
    better = (100.0, 1, 80.0)
    worse = (681.0, 4, 413.0)

    assert _accepts_trial(better, current, force_all=True, converted_drawings=0) is False
    assert _accepts_trial(worse, current, force_all=True, converted_drawings=0) is False
    assert _accepts_trial(worse, current, force_all=True, converted_drawings=1) is True


def test_full_width_block_saved_to_docx_has_no_tblpPr(tmp_path):
    """端到端锁定：双栏文档经 _full_width_block_elements 包装写盘后，审计 tblpPr=0。

    这是 P0 修复（主链路用连续分节符，不再产 w:tblpPr）的回归护栏。
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document()
    sectPr = doc.sections[0]._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), '2')
    body_para = doc.add_paragraph('[FIGURE_0]')
    block = docx_insert_module._full_width_block_elements(doc, [body_para._element])
    for elem in block:
        doc.element.body.append(elem)
    path = tmp_path / 'float.docx'
    doc.save(path)
    report = audit_docx_structure(path)
    assert report['tblpPr_count'] == 0
    assert is_clean(path) is True


def test_extract_heading_color_named_definecolor_and_none():
    """heading_color 从 CLS 提取：命名色 + definecolor rgb，无色/未知名返回 None。"""
    extract = docx_insert_module._extract_heading_color
    assert extract("") is None
    assert extract(r"\def\@maketitle{\bfseries\@title}") is None
    assert extract(r"\def\@maketitle{\color{blue}\@title}") == (0, 0, 255)
    cls = (r"\definecolor{titleblue}{rgb}{0.1,0.2,0.8}"
           "\n"
           r"\def\@maketitle{\color{titleblue}\@title}")
    assert extract(cls) == (26, 51, 204)
    assert extract(r"\def\@maketitle{\color{unknownx}\@title}") is None


def test_normalize_abstract_paragraphs_keeps_multiple_paragraphs():
    """摘要多段不被合并：首段加 bold 'Abstract.' 标签，其余段保留独立。"""
    from docx import Document
    doc = Document()
    title = doc.add_paragraph("Abstract")
    try:
        title.style = doc.styles["Abstract Title"]
    except KeyError:
        pass
    doc.add_paragraph("First body paragraph of the abstract content.")
    doc.add_paragraph("Second body paragraph of the abstract content.")

    docx_insert_module._normalize_abstract_paragraphs(doc, cls_content=None)

    paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(paras) == 2, f"期望保留2段 实际{len(paras)}: {[p.text for p in paras]}"
    assert paras[0].runs[0].text.startswith("Abstract"), "首段首run应为Abstract标签"
    assert paras[0].runs[0].bold is True, "首段首run应加粗"
    assert "First body paragraph" in paras[0].text
    assert "Second body paragraph" in paras[1].text
    assert not paras[1].text.startswith("Abstract"), "第二段不应有Abstract标签"
