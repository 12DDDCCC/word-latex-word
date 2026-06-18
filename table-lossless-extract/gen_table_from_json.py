#!/usr/bin/env python3
"""JSON → Word 文档生成

将 tikz_to_json / tabular_to_json 的输出还原为 .docx 文件。
支持 gridSpan / vMerge / booktabs 边框 / shading / vAlign。

v3.3.1: 修复杂线问题——用"主 tc 集合"方案清除非主 tc 边框，替代 vMerge/hMerge continue 检查。
"""
import os
from docx import Document
from docx.shared import Pt, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _set_tc_border(tc, edge, val, sz, color='000000', space='0'):
    """直接在 <w:tc> 元素上设置单个边的边框"""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    elem = tcBorders.find(qn(f'w:{edge}'))
    if elem is None:
        elem = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
        tcBorders.append(elem)
    elem.set(qn('w:val'), str(val))
    elem.set(qn('w:sz'), str(sz))
    elem.set(qn('w:color'), str(color))
    elem.set(qn('w:space'), str(space))


def _set_tc_all_nil(tc):
    """将一个 <w:tc> 的所有边框设为 nil"""
    for edge in ('top', 'bottom', 'left', 'right'):
        _set_tc_border(tc, edge, 'nil', '0', 'auto', '0')


def _get_row_tcs(table, ri):
    """获取第 ri 行的所有底层 <w:tc> 元素列表"""
    tr = table.rows[ri]._tr
    return [tc for tc in tr.findall(qn('w:tc'))]


def _build_tc_grid(table, rows, num_cols):
    """Map logical grid positions to the underlying main <w:tc> after merges."""
    tc_grid = {}
    for ri, row in enumerate(rows):
        logical_col = 0
        row_tcs = _get_row_tcs(table, ri)
        tc_idx = 0
        for cell_data in row.get('cells', []):
            cs = cell_data.get('col_start', logical_col)
            while logical_col < cs:
                logical_col += 1
                tc_idx += 1
            if tc_idx < len(row_tcs):
                tc_grid[(ri, cs)] = row_tcs[tc_idx]
            else:
                try:
                    tc_grid[(ri, cs)] = table.cell(ri, cs)._tc
                except Exception:
                    pass
            logical_col = cs + cell_data.get('gridSpan', 1)
            tc_idx += 1
        while logical_col < num_cols:
            logical_col += 1
            tc_idx += 1
    return tc_grid


def _set_table_grid_widths(table, grid_cols):
    """Apply requested widths to both the table grid and table properties."""
    from docx.oxml import OxmlElement

    widths = [int(col.get('width_twips', 0) or 0) for col in grid_cols]
    table.autofit = False
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(sum(widths)))


def generate_docx(table_data, output_path):
    """从 JSON 数据生成 Word 表格文档"""
    grid_cols = table_data.get('grid_cols', table_data.get('grid_col', []))
    rows = table_data.get('rows', [])
    num_cols = len(grid_cols)
    num_rows = len(rows)

    if num_rows == 0 or num_cols == 0:
        return

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    _set_table_grid_widths(table, grid_cols)

    # Step 1: 设置列宽
    for ci, col in enumerate(grid_cols):
        w = col.get('width_twips', 1700)
        for ri in range(num_rows):
            cell = table.cell(ri, ci)
            cell.width = Twips(w)

    # Step 2: 设置行高
    for ri, row in enumerate(rows):
        rh = row.get('row_height', 0)
        if rh and int(rh) > 0:
            tr = table.rows[ri]._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
                tr.insert(0, trPr)
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is None:
                trHeight = parse_xml(f'<w:trHeight {nsdecls("w")}/>')
                trPr.append(trHeight)
            trHeight.set(qn('w:val'), str(rh))
            trHeight.set(qn('w:hRule'), row.get('row_height_rule') or 'atLeast')

    # Step 3: gridSpan merge
    for ri, row in enumerate(rows):
        for cell_data in row.get('cells', []):
            cs = cell_data.get('col_start', 0)
            gs = cell_data.get('gridSpan', 1)
            if gs > 1 and cs + gs <= num_cols:
                try:
                    table.cell(ri, cs).merge(table.cell(ri, cs + gs - 1))
                except Exception:
                    pass

    # Step 4: vMerge merge
    for ri, row in enumerate(rows):
        for cell_data in row.get('cells', []):
            cs = cell_data.get('col_start', 0)
            gs = cell_data.get('gridSpan', 1)
            vm = cell_data.get('vMerge', '')
            if vm == 'restart':
                span = 1
                for next_ri in range(ri + 1, num_rows):
                    found_continue = False
                    for cd in rows[next_ri].get('cells', []):
                        ncs = cd.get('col_start', 0)
                        ngs = cd.get('gridSpan', 1)
                        nvm = cd.get('vMerge', '')
                        if ncs == cs and ngs == gs and nvm == 'continue':
                            span += 1
                            found_continue = True
                            break
                    if not found_continue:
                        break
                if span > 1:
                    try:
                        table.cell(ri, cs).merge(table.cell(ri + span - 1, cs + gs - 1))
                    except Exception:
                        pass

    # Step 5: 清除所有底层 <w:tc> 的边框为 nil
    for ri in range(num_rows):
        for tc in _get_row_tcs(table, ri):
            _set_tc_all_nil(tc)

    tc_grid = _build_tc_grid(table, rows, num_cols)

    # Step 6: 按行遍历 JSON cells，设置边框
    # 使用 table.cell(ri, cs)._tc 获取正确的底层 <w:tc> 元素
    for ri, row in enumerate(rows):
        for cell_data in row.get('cells', []):
            cs = cell_data.get('col_start', 0)
            gs = cell_data.get('gridSpan', 1)
            vm = cell_data.get('vMerge', '')

            # vMerge continue 行不设置边框
            if vm == 'continue':
                continue

            tc = tc_grid.get((ri, cs))
            if tc is None:
                try:
                    tc = table.cell(ri, cs)._tc
                except Exception:
                    continue

            borders = cell_data.get('borders', {})
            for edge in ('top', 'bottom', 'left', 'right'):
                bd = borders.get(edge, {})
                val = bd.get('val', 'nil')
                if val and val not in ('none', 'nil', ''):
                    _set_tc_border(tc, edge, val,
                                   str(bd.get('sz', '4')),
                                   bd.get('color', '000000'),
                                   str(bd.get('space', '0')))

    # Step 6b: 清除非主 tc 的边框，防止合并区域内部出现杂线
    # 收集所有"主 tc"（每个 JSON cell 对应的 table.cell()._tc）
    main_tcs = set()
    for ri, row in enumerate(rows):
        for cell_data in row.get('cells', []):
            cs = cell_data.get('col_start', 0)
            vm = cell_data.get('vMerge', '')
            if vm == 'continue':
                continue
            tc = tc_grid.get((ri, cs))
            if tc is not None:
                main_tcs.add(id(tc))

    # 遍历所有底层 tc，非主 tc 的边框全部设为 nil
    for ri in range(num_rows):
        tr = table.rows[ri]._tr
        for tc in tr.findall(qn('w:tc')):
            if id(tc) not in main_tcs:
                _set_tc_all_nil(tc)

    # Step 7: 设置文本内容
    for ri, row in enumerate(rows):
        for cell_data in row.get('cells', []):
            cs = cell_data.get('col_start', 0)
            gs = cell_data.get('gridSpan', 1)
            vm = cell_data.get('vMerge', '')

            if vm == 'continue':
                continue

            try:
                cell = table.cell(ri, cs)
            except Exception:
                continue

            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ''

            paragraphs = cell_data.get('paragraphs', [])
            if not paragraphs:
                text = cell_data.get('text', '')
                if text and cell.paragraphs:
                    cell.paragraphs[0].text = text
                continue

            first_para = True
            for pi, para in enumerate(paragraphs):
                if first_para:
                    p = cell.paragraphs[0]
                    first_para = False
                else:
                    p = cell.add_paragraph()

                align = para.get('alignment', '')
                if align == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                runs = para.get('runs', [])
                if not runs:
                    text = cell_data.get('text', '')
                    if text:
                        run = p.add_run(text)
                    continue

                for run_data in runs:
                    run = p.add_run(run_data.get('text', ''))
                    fmt = run_data.get('format', {})
                    if fmt.get('bold'):
                        run.bold = True
                    if fmt.get('italic'):
                        run.italic = True
                    if fmt.get('size_pt'):
                        run.font.size = Pt(fmt['size_pt'])

            v_align = cell_data.get('vAlign', '')
            if v_align == 'center':
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            shading = cell_data.get('shading', {})
            if shading:
                fill = shading.get('fill', '')
                if fill:
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                        tc.insert(0, tcPr)
                    shd = tcPr.find(qn('w:shd'))
                    if shd is None:
                        shd = parse_xml(f'<w:shd {nsdecls("w")}/>')
                        tcPr.append(shd)
                    shd.set(qn('w:fill'), fill)
                    shd.set(qn('w:val'), 'clear')

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
